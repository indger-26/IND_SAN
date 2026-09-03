# ----- BIBLIOTECAS -----

import pandas as pd
import itertools
from datetime import datetime
from supabase import create_client, Client
import os
import streamlit as st
import matplotlib.pyplot as plt
import numpy as np
from pandas.plotting import table
import dataframe_image as dfi
from PIL import Image
from matplotlib.font_manager import FontProperties
from docx.shared import Inches
import textwrap
import unicodedata
from unidecode import unidecode
import gdown
from IPython.display import display, Javascript
import time
import random
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
import streamlit as st
from dotenv import load_dotenv

url = st.secrets["SUPABASE_URL"]
key = st.secrets["SUPABASE_KEY"]

supabase: Client = create_client(url, key)


#############################################################################################################################################################################################################################

# IMPORTAÇÃO E TRATAMENTO DOS DADOS FORNECIDOS

def ler_ajustar_arquivo ():

  infos_gerais = pd.read_excel("ATT_SQL.xlsx", sheet_name="GERAIS")
  iqa_detalhado = pd.read_excel("ATT_SQL.xlsx", sheet_name="IQA_DETALHADO")
  plano = pd.read_excel("ATT_SQL.xlsx", sheet_name="IQA_PLANO")

  iqe_detalhado = pd.read_excel("ATT_SQL.xlsx", sheet_name="IQE_DETALHADO")
  plano_iqe = pd.read_excel("ATT_SQL.xlsx", sheet_name="IQE_PLANO")

  iqa_detalhado = iqa_detalhado.drop('id_pond', axis=1)
  iqa_detalhado = iqa_detalhado.drop('conformidade_vi', axis=1)
  iqe_detalhado = iqe_detalhado.drop('id_pond', axis=1)
  iqe_detalhado = iqe_detalhado.drop('conformidade_vi', axis=1)

  prestadora_sigla = infos_gerais.iloc[0, 1]

  if prestadora_sigla == "BRK":
    prestadora = "BRK Ambiental"
  elif prestadora_sigla == "ADS":
    prestadora = "Conasa Águas do Sertão"
  else:
    prestadora = "Verde Ambiental Alagoas"

  ano = infos_gerais.iloc[1, 1]

  ano_contratual = infos_gerais.iloc[2, 1]

  mes = infos_gerais.iloc[3, 1]

  trimestre = infos_gerais.iloc[4, 1]

  iqa_prest = infos_gerais.iloc[7, 1]
  iqa_vi = infos_gerais.iloc[10, 1]
  iqa_meta = infos_gerais.iloc[11, 1]

  if mes == "November" or mes == "February" or mes == "May" or mes == "August":

    iqe_prest = float(infos_gerais.iloc[18, 1])
    iqe_vi = float(infos_gerais.iloc[21, 1])
    iqe_meta = float(infos_gerais.iloc[22, 1])

  
  
  # TRATAMENTO - IQA

  pos = iqa_detalhado.columns.get_loc("data_coleta")
  iqa_detalhado.insert(pos + 1, "mes", mes)
  iqa_detalhado.insert(pos + 2, "ano", ano)

  iqa_detalhado.to_csv('iqa_detalhado.csv', index=False)

  plano["Parâmetros"] = plano["Parâmetros"].replace({
      "Cor": "Cor Aparente",
      "Cor aparente": "Cor Aparente",
      "Coliformes totais": "Coliformes Totais",
      "Escherichia Coli": "Escherichia coli",
      "Cloro Residual Livre": "Cloro residual livre",
      "Residual de Cloro Livre": "Cloro residual livre",
      "Residual de cloro livre": "Cloro residual livre"
  })

  parametros_principais = [
      "Turbidez",
      "Cor Aparente",
      "pH",
      "Cloro residual livre",
      "Coliformes Totais",
      "Escherichia coli"
  ]

  plano["Parâmetros"] = plano["Parâmetros"].apply(
      lambda x: x if x in parametros_principais else "Demais Parâmetros"
  )

  colunas_chave = ["Cidade", "Parâmetros", "SAA"]
  plano_agrupado = plano.groupby(colunas_chave, as_index=False)[["PLANO", "DESCONSIDERACOES - ARSAL"]].sum()

  plano_agrupado.insert(0, "MES", mes)
  plano_agrupado.insert(1, "ANO", ano)

  municipios = plano_agrupado["Cidade"].unique()
  saas = plano_agrupado["SAA"].unique()

  combinacoes_principais = pd.DataFrame(
      list(itertools.product([mes], [ano], municipios, parametros_principais, saas)),
      columns=["MES", "ANO", "Cidade", "Parâmetros", "SAA"]
  )

  plano_principais = pd.merge(combinacoes_principais, plano_agrupado,
                          on=["MES", "ANO", "Cidade", "Parâmetros", "SAA"],
                          how="left")

  plano_principais["PLANO"] = plano_principais["PLANO"].fillna(0)
  plano_principais["DESCONSIDERACOES - ARSAL"] = plano_principais["DESCONSIDERACOES - ARSAL"].fillna(0)  # NOVO

  plano_demais = plano_agrupado[plano_agrupado["Parâmetros"] == "Demais Parâmetros"]

  plano_completo = pd.concat([plano_principais, plano_demais], ignore_index=True)

  plano_dados_tratados = plano_completo.copy()

  plano_dados_tratados = plano_dados_tratados[
      ~(
          plano_dados_tratados["PLANO"].fillna(0).eq(0)
          &
          plano_dados_tratados["DESCONSIDERACOES - ARSAL"].fillna(0).eq(0)
      )
  ]

  plano_dados_tratados["PLANO"] = plano_dados_tratados["PLANO"].astype("Int64")
  plano_dados_tratados["DESCONSIDERACOES - ARSAL"] = plano_dados_tratados["DESCONSIDERACOES - ARSAL"].astype("Int64")

  plano_totais_municipio = (plano_completo.groupby("Cidade", as_index=False)[["PLANO", "DESCONSIDERACOES - ARSAL"]].sum())

  plano_totais_municipio["PLANO"] = plano_totais_municipio["PLANO"].astype("Int64")
  plano_totais_municipio["DESCONSIDERACOES - ARSAL"] = plano_totais_municipio["DESCONSIDERACOES - ARSAL"].astype("Int64")

  plano_totais_parametros = (plano_completo.groupby("Parâmetros", as_index=False)[["PLANO", "DESCONSIDERACOES - ARSAL"]].sum())

  plano_totais_parametros["PLANO"] = plano_totais_parametros["PLANO"].astype("Int64")
  plano_totais_parametros["DESCONSIDERACOES - ARSAL"] = plano_totais_parametros["DESCONSIDERACOES - ARSAL"].astype("Int64")

  plano_totais_saa = (plano_completo.groupby("SAA", as_index=False)[["PLANO", "DESCONSIDERACOES - ARSAL"]].sum())

  iqe_prest = "x"
  iqe_vi = "x"
  iqe_meta = "x"
  iqe_detalhado = "x"
  plano_iqe_completo = "x"
  plano_iqe_totais_municipio = "x"
  plano_iqe_totais_parametros = "x"
  iqe_detalhado2 = "x"



  # TRATAMENTO - IQE

  if mes == "November" or mes == "February" or mes == "May" or mes == "August":

    pos_iqe = iqe_detalhado.columns.get_loc("id_amostra")

    iqe_detalhado.insert(pos_iqe + 1, "trimestre", trimestre)
    iqe_detalhado.insert(pos_iqe + 2, "ano_contratual", ano_contratual)

    plano_iqe["parametros"] = plano_iqe["parametros"].replace({
      "DBO - 5 dias": "DBO",
      "Óleos e Graxas Totais": "Óleos e Graxas",
      "Oxigênio Dissolvido": "OD"
    })

    parametros_principais_iqe = [
      "DBO",
      "DQO",
      "Materiais Flutuantes",
      "Temperatura",
      "pH",
      "Óleos e Graxas",
      "OD"
    ]

    plano_iqe["parametros"] = plano_iqe["parametros"].apply(
        lambda x: x if x in parametros_principais_iqe else "Demais Parâmetros"
    )

    iqe_detalhado2 = iqe_detalhado.copy()

    iqe_detalhado2["analise"] = iqe_detalhado2["analise"].apply(
        lambda x: x if x in parametros_principais_iqe else "Demais Parâmetros"
    )

    colunas_chave_iqe = ["cidade", "parametros", "ETE"]
    plano_iqe_agrupado = plano_iqe.groupby(colunas_chave_iqe, as_index=False)[["plano", "desconsideracoes_arsal"]].sum()

    plano_iqe_agrupado.insert(0, "trimestre", trimestre)
    plano_iqe_agrupado.insert(1, "ano_contratual", ano_contratual)

    municipios = plano_iqe_agrupado["cidade"].unique()
    sess = plano_iqe_agrupado["ETE"].unique()

    combinacoes_principais_iqe = pd.DataFrame(
        list(itertools.product([trimestre], [ano_contratual], municipios, parametros_principais_iqe, sess)),
        columns=["trimestre", "ano_contratual", "cidade", "parametros", "ETE"]
    )

    plano_iqe_principais = pd.merge(combinacoes_principais_iqe, plano_iqe_agrupado,
                            on=["trimestre", "ano_contratual", "cidade", "parametros", "ETE"],
                            how="left")

    plano_iqe_principais[["plano", "desconsideracoes_arsal"]] = plano_iqe_principais[["plano", "desconsideracoes_arsal"]].fillna(0)

    plano_iqe_demais = plano_iqe_agrupado[plano_iqe_agrupado["parametros"] == "Demais Parâmetros"]
    plano_iqe_completo = pd.concat([plano_iqe_principais, plano_iqe_demais], ignore_index=True)

    plano_iqe_totais_municipio = (plano_iqe_completo.groupby("cidade", as_index=False)[["plano", "desconsideracoes_arsal"]].sum())

    plano_iqe_totais_municipio["plano"] = plano_iqe_totais_municipio["plano"].astype("Int64")
    plano_iqe_totais_municipio["desconsideracoes_arsal"] = plano_iqe_totais_municipio["desconsideracoes_arsal"].astype("Int64")

    plano_iqe_totais_parametros = (plano_iqe_completo.groupby("parametros", as_index=False)[["plano", "desconsideracoes_arsal"]].sum())

    plano_iqe_totais_parametros["plano"] = plano_iqe_totais_parametros["plano"].astype("Int64")
    plano_iqe_totais_parametros["desconsideracoes_arsal"] = plano_iqe_totais_parametros["desconsideracoes_arsal"].astype("Int64")

    plano_iqe_totais_ete = (plano_iqe_completo.groupby("ETE", as_index=False)[["plano", "desconsideracoes_arsal"]].sum())


    return prestadora_sigla, ano, ano_contratual, mes, iqa_prest, iqa_vi, iqa_meta, trimestre, iqa_detalhado, plano_dados_tratados, plano_totais_municipio, plano_totais_parametros, iqe_prest, iqe_vi, iqe_meta, iqe_detalhado, plano_iqe_completo, plano_iqe_totais_municipio, plano_iqe_totais_parametros, iqe_detalhado2

  else:

    return prestadora_sigla, ano, ano_contratual, mes, iqa_prest, iqa_vi, iqa_meta, trimestre, iqa_detalhado, plano_dados_tratados, plano_totais_municipio, plano_totais_parametros, iqe_prest, iqe_vi, iqe_meta, iqe_detalhado, plano_iqe_completo, plano_iqe_totais_municipio, plano_iqe_totais_parametros, iqe_detalhado2



###############################################################################################################################################################################################################################

# ELABORAÇÃO DAS TABELAS ALIMENTADAS AO BANCO DE DADOS

def criar_tabelas (prestadora_sigla, ano, ano_contratual, mes, iqa_prest, iqa_vi, iqa_meta, trimestre, iqa_detalhado, plano_dados_tratados, plano_totais_municipio, plano_totais_parametros, iqe_prest, iqe_vi, iqe_meta, iqe_detalhado, plano_iqe_completo, plano_iqe_totais_municipio, plano_iqe_totais_parametros, iqe_detalhado2):


  # IQA

      # TABELA REALIZADOS

  valor_cloro = plano_totais_parametros[plano_totais_parametros['Parâmetros'] == 'Cloro residual livre']['PLANO'].values[0]
  plano_totais_parametros.loc[plano_totais_parametros['Parâmetros'] == 'Demais Parâmetros', 'PLANO'] += valor_cloro

  plano_totais_parametros = plano_totais_parametros[plano_totais_parametros['Parâmetros'] != 'Cloro residual livre'].reset_index(drop=True)

  realizados = plano_totais_parametros

  meses_en = {
      'January': 1, 'February': 2, 'March': 3, 'April': 4,
      'May': 5, 'June': 6, 'July': 7, 'August': 8,
      'September': 9, 'October': 10, 'November': 11, 'December': 12
  }

  mes_numero = meses_en[mes]
  data = datetime(ano, mes_numero, 1)

  mes_ano = data.strftime('%d/%m/%Y')

  realizados.insert(0, 'ano', ano)
  realizados.insert(1, 'mes', mes)
  realizados.insert(2, 'mes_ano', mes_ano)

  realizados = realizados.drop('DESCONSIDERACOES - ARSAL', axis=1)

  realizados = realizados.rename(columns={'Parâmetros': 'parametros'})
  realizados = realizados.rename(columns={'PLANO': 'plano_de_amostragem'})

  realizados['analises_realizadas'] = realizados.apply(
    lambda row: (
        "" if (
            iqa_detalhado[
                (iqa_detalhado['ano'] == row['ano']) &
                (iqa_detalhado['mes'] == row['mes']) &
                (
                    ~iqa_detalhado['analise'].isin(["Turbidez", "Cor Aparente", "pH", "Coliformes Totais", "Escherichia coli"])
                    if row['parametros'] == "Demais Parâmetros"
                    else iqa_detalhado['analise'] == row['parametros']
                )
            ]['mes'].notna().sum()
        ) == 1 else (
            iqa_detalhado[
                (iqa_detalhado['mes'] == row['mes']) &
                (iqa_detalhado['ano'] == row['ano']) &
                (
                    ~iqa_detalhado['analise'].isin(["Turbidez", "Cor Aparente", "pH", "Coliformes Totais", "Escherichia coli"])
                    if row['parametros'] == "Demais Parâmetros"
                    else iqa_detalhado['analise'] == row['parametros']
                )
            ]['mes'].notna().sum()
        )
    ),
    axis=1
  )

  realizados['expurgos'] = realizados.apply(
  lambda row: (
      iqa_detalhado[
          (iqa_detalhado['mes'] == row['mes']) &
          (iqa_detalhado['ano'] == row['ano']) &
          (iqa_detalhado['expurgos'] == "EXPURGAR") &
          (
              ~iqa_detalhado['analise'].isin(["Turbidez", "Cor Aparente", "pH", "Coliformes Totais", "Escherichia coli"])
              if row['parametros'] == "Demais Parâmetros"
              else iqa_detalhado['analise'] == row['parametros']
          )
      ].shape[0]
  ),
  axis=1
  )

  realizados['nam_realiz'] = realizados.apply(
  lambda row: max(row['plano_de_amostragem'], row['analises_realizadas']) - (row['expurgos'] if pd.notna(row['expurgos']) else 0),
  axis=1
  )

  print(realizados)

  realizados.to_csv('realizados.csv', index=False)


    
      # TABELA PARAMETROS

  parametros = realizados.drop('plano_de_amostragem', axis=1)
  
  parametros['expurgos_conformes'] = parametros.apply(
    lambda row: (
        iqa_detalhado[
            (iqa_detalhado['mes'] == row['mes']) &
            (iqa_detalhado['ano'] == row['ano']) &
            (iqa_detalhado['expurgos'] == "EXPURGAR") &
            (iqa_detalhado['resultado'] == "Conforme") &
            (
                ~iqa_detalhado['analise'].isin(["Turbidez", "Cor Aparente", "pH", "Coliformes Totais", "Escherichia coli"])
                if row['parametros'] == "Demais Parâmetros"
                else iqa_detalhado['analise'] == row['parametros']
            )
        ].shape[0]
    ),
    axis=1
    )

  parametros['nam_conf'] = parametros.apply(
  lambda row: (
      "" if (
          iqa_detalhado[
              (iqa_detalhado['mes'] == row['mes']) &
              (iqa_detalhado['ano'] == row['ano']) &
              (iqa_detalhado['resultado'] == "Conforme") &
              (
                  ~iqa_detalhado['analise'].isin(["Turbidez", "Cor Aparente", "pH", "Coliformes Totais", "Escherichia coli"])
                  if row['parametros'] == "Demais Parâmetros"
                  else iqa_detalhado['analise'] == row['parametros']
              )
          ]['mes'].notna().sum() - (row['expurgos_conformes'] if pd.notna(row['expurgos_conformes']) else 0)
      ) == 1 else (
          iqa_detalhado[
              (iqa_detalhado['mes'] == row['mes']) &
              (iqa_detalhado['ano'] == row['ano']) &
              (iqa_detalhado['resultado'] == "Conforme") &
              (
                  ~iqa_detalhado['analise'].isin(["Turbidez", "Cor Aparente", "pH", "Coliformes Totais", "Escherichia coli"])
                  if row['parametros'] == "Demais Parâmetros"
                  else iqa_detalhado['analise'] == row['parametros']
              )
          ]['mes'].notna().sum() - (row['expurgos_conformes'] if pd.notna(row['expurgos_conformes']) else 0)
      )
  ),
  axis=1
  )

  parametros['expurgos_totais'] = parametros.apply(
  lambda row: (
      iqa_detalhado[
          (iqa_detalhado['mes'] == row['mes']) &
          (iqa_detalhado['ano'] == row['ano']) &
          (iqa_detalhado['expurgos'] == "EXPURGAR") &
          ~iqa_detalhado['analise'].isin(["Turbidez", "Cor Aparente", "pH", "Coliformes Totais", "Escherichia coli"])
      ].shape[0]
  ) if row['parametros'] == "Demais Parâmetros" else (
      iqa_detalhado[
          (iqa_detalhado['mes'] == row['mes']) &
          (iqa_detalhado['ano'] == row['ano']) &
          (iqa_detalhado['expurgos'] == "EXPURGAR") &
          (iqa_detalhado['analise'] == row['parametros'])
      ].shape[0]
  ),
  axis=1
  )

  parametros['nam_realiz'] = parametros.apply(
    lambda row: realizados[
        (realizados['mes'] == row['mes']) &
        (realizados['ano'] == row['ano']) &
        (realizados['parametros'] == row['parametros'])
    ]['nam_realiz'].values[0] if any(
        (realizados['mes'] == row['mes']) &
        (realizados['ano'] == row['ano']) &
        (realizados['parametros'] == row['parametros'])
    ) else None,
    axis=1
  )


  parametros['iqa'] = parametros.apply(
    lambda row: row['nam_conf'] / row['nam_realiz'] if row['nam_realiz'] != 0 else None,
    axis=1
  )

  parametros.to_csv('parametros.csv', index=False)


      # TABELA MUN

  mun = plano_totais_municipio

  mun.insert(0, 'ano', ano)
  mun.insert(1, 'mes', mes)
  mun.insert(2, 'ano_mes', mes_ano)

  mun = mun.rename(columns={'Cidade': 'cidade'})
  mun = mun.rename(columns={'PLANO': 'plano_de_amostragem'})
  mun = mun.rename(columns={'DESCONSIDERACOES - ARSAL': 'desconsideracoes_plano'})

  mun['analises_realizadas'] = mun.apply(
      lambda row: (
          "" if (
              iqa_detalhado[
                  (iqa_detalhado['mes'] == row['mes']) &
                  (iqa_detalhado['ano'] == row['ano']) &
                  (iqa_detalhado['municipio'] == row['cidade'])
              ]['mes'].notna().sum()
          ) == 1 else (
              iqa_detalhado[
                  (iqa_detalhado['mes'] == row['mes']) &
                  (iqa_detalhado['ano'] == row['ano']) &
                  (iqa_detalhado['municipio'] == row['cidade'])
              ]['mes'].notna().sum()
          )
      ),
      axis=1
    )

  mun['expurgos'] = mun.apply(
  lambda row: (
      iqa_detalhado[
          (iqa_detalhado['mes'] == row['mes']) &
          (iqa_detalhado['ano'] == row['ano']) &
          (iqa_detalhado['municipio'] == row['cidade']) &
          (iqa_detalhado['expurgos'] == "EXPURGAR")
      ].shape[0]
  ),
  axis=1
  )

  mun['nam_realiz'] = mun.apply(
  lambda row: max(
      row['plano_de_amostragem'] if pd.notna(row['plano_de_amostragem']) else 0,
      row['analises_realizadas'] if row['analises_realizadas'] != "" and pd.notna(row['analises_realizadas']) else 0
  ) - (row['expurgos'] if pd.notna(row['expurgos']) else 0),
  axis=1
  )

  mun['expurgos_conformes'] = mun.apply(
  lambda row: (
      iqa_detalhado[
          (iqa_detalhado['mes'] == row['mes']) &
          (iqa_detalhado['ano'] == row['ano']) &
          (iqa_detalhado['municipio'] == row['cidade']) &
          (iqa_detalhado['resultado'] == "Conforme") &
          (iqa_detalhado['expurgos'] == "EXPURGAR")
      ].shape[0]
  ),
  axis=1
  )

  mun['nam_conf'] = mun.apply(
  lambda row: (
      "" if (
          iqa_detalhado[
              (iqa_detalhado['mes'] == row['mes']) &
              (iqa_detalhado['ano'] == row['ano']) &
              (iqa_detalhado['municipio'] == row['cidade']) &
              (iqa_detalhado['resultado'] == "Conforme")
          ]['mes'].notna().sum() - (row['expurgos_conformes'] if pd.notna(row['expurgos_conformes']) else 0)
      ) == 1 else (
          iqa_detalhado[
              (iqa_detalhado['mes'] == row['mes']) &
              (iqa_detalhado['ano'] == row['ano']) &
              (iqa_detalhado['municipio'] == row['cidade']) &
              (iqa_detalhado['resultado'] == "Conforme")
          ]['mes'].notna().sum() - (row['expurgos_conformes'] if pd.notna(row['expurgos_conformes']) else 0)
      )
  ),
  axis=1
  )

  mun['iqa'] = mun.apply(
    lambda row: row['nam_conf'] / row['nam_realiz'] if row['nam_realiz'] != 0 else None,
    axis=1
  )

  lat_long = pd.read_excel('lat_long.xlsx')

  mun['latitude'] = mun.apply(
    lambda row: lat_long[lat_long['cidade'] == row['cidade']]['latitude'].values[0] if any(lat_long['cidade'] == row['cidade']) else None,
    axis=1
  )

  mun['longitude'] = mun.apply(
      lambda row: lat_long[lat_long['cidade'] == row['cidade']]['longitude'].values[0] if any(lat_long['cidade'] == row['cidade']) else None,
      axis=1
  )

  mun['coordenada'] = mun['latitude'].astype(str) + ',' + mun['longitude'].astype(str)

  mun.to_csv('mun.csv', index=False)


      # TABELA PLANO

  plano_final = plano_dados_tratados

  cols = plano_final.columns.tolist()
  idx_ano = cols.index('ANO')
  idx_mes = cols.index('MES')
  cols[idx_ano], cols[idx_mes] = cols[idx_mes], cols[idx_ano]
  plano_final = plano_final[cols]

  plano_final = plano_final.rename(columns={'ANO': 'ano'})
  plano_final = plano_final.rename(columns={'MES': 'mes'})
  plano_final = plano_final.rename(columns={'Cidade': 'cidade'})
  plano_final = plano_final.rename(columns={'Parâmetros': 'parametros'})
  plano_final = plano_final.rename(columns={'PLANO': 'plano_de_amostragem'})
  plano_final = plano_final.rename(columns={'DESCONSIDERACOES - ARSAL': 'desconsideracoes_plano'})

  plano_final['amostras_realizadas'] = plano_final.apply(
      lambda row: (
          iqa_detalhado[
              (iqa_detalhado['mes'] == row['mes']) &
              (iqa_detalhado['ano'] == row['ano']) &
              (iqa_detalhado['municipio'] == row['cidade']) &
              (iqa_detalhado['saa'] == row['SAA']) &
              ~iqa_detalhado['analise'].isin(["Turbidez", "Cor Aparente", "pH", "Coliformes Totais", "Escherichia coli"])
          ]['mes'].notna().sum()
      ) if row['parametros'] == "Demais Parâmetros" else (
          iqa_detalhado[
              (iqa_detalhado['mes'] == row['mes']) &
              (iqa_detalhado['ano'] == row['ano']) &
              (iqa_detalhado['analise'] == row['parametros']) &
              (iqa_detalhado['municipio'] == row['cidade']) &
              (iqa_detalhado['saa'] == row['SAA'])
          ]['mes'].notna().sum()
      ),
      axis=1
    )

  plano_final['expurgos'] = plano_final.apply(
    lambda row: (
        iqa_detalhado[
            (iqa_detalhado['mes'] == row['mes']) &
            (iqa_detalhado['ano'] == row['ano']) &
            (iqa_detalhado['municipio'] == row['cidade']) &
            (iqa_detalhado['saa'] == row['SAA']) &
            (iqa_detalhado['expurgos'] == "EXPURGAR") &
            ~iqa_detalhado['analise'].isin(["Turbidez", "Cor Aparente", "pH", "Coliformes Totais", "Escherichia coli"])
        ]['mes'].notna().sum()
    ) if row['parametros'] == "Demais Parâmetros" else (
        iqa_detalhado[
            (iqa_detalhado['mes'] == row['mes']) &
            (iqa_detalhado['ano'] == row['ano']) &
            (iqa_detalhado['analise'] == row['parametros']) &
            (iqa_detalhado['municipio'] == row['cidade']) &
            (iqa_detalhado['expurgos'] == "EXPURGAR") &
            (iqa_detalhado['saa'] == row['SAA'])
        ]['mes'].notna().sum()
    ),
    axis=1
  )

  plano_final['nam_realiz'] = plano_final[['plano_de_amostragem', 'amostras_realizadas']].max(axis=1) - plano_final['expurgos'].fillna(0)

  plano_final['expurgos_conf'] = plano_final.apply(
    lambda row: (
        iqa_detalhado[
            (iqa_detalhado['mes'] == row['mes']) &
            (iqa_detalhado['ano'] == row['ano']) &
            (iqa_detalhado['municipio'] == row['cidade']) &
            (iqa_detalhado['saa'] == row['SAA']) &
            (iqa_detalhado['expurgos'] == "EXPURGAR") &
            (iqa_detalhado['resultado'] == "Conforme") &
            ~iqa_detalhado['analise'].isin(["Turbidez", "Cor Aparente", "pH", "Coliformes Totais", "Escherichia coli"])
        ]['mes'].notna().sum()
    ) if row['parametros'] == "Demais Parâmetros" else (
        iqa_detalhado[
            (iqa_detalhado['mes'] == row['mes']) &
            (iqa_detalhado['ano'] == row['ano']) &
            (iqa_detalhado['analise'] == row['parametros']) &
            (iqa_detalhado['municipio'] == row['cidade']) &
            (iqa_detalhado['saa'] == row['SAA']) &
            (iqa_detalhado['expurgos'] == "EXPURGAR") &
            (iqa_detalhado['resultado'] == "Conforme")
        ]['mes'].notna().sum()
    ),
    axis=1
  )

  plano_final['nam_conf'] = plano_final.apply(
    lambda row: (
        iqa_detalhado[
            (iqa_detalhado['mes'] == row['mes']) &
            (iqa_detalhado['ano'] == row['ano']) &
            (iqa_detalhado['municipio'] == row['cidade']) &
            (iqa_detalhado['saa'] == row['SAA']) &
            (iqa_detalhado['resultado'] == "Conforme") &
            ~iqa_detalhado['analise'].isin(["Turbidez", "Cor Aparente", "pH", "Coliformes Totais", "Escherichia coli"])
        ]['mes'].notna().sum() - (row['expurgos_conf'] if pd.notna(row['expurgos_conf']) else 0)
    ) if row['parametros'] == "Demais Parâmetros" else (
        iqa_detalhado[
            (iqa_detalhado['mes'] == row['mes']) &
            (iqa_detalhado['ano'] == row['ano']) &
            (iqa_detalhado['analise'] == row['parametros']) &
            (iqa_detalhado['municipio'] == row['cidade']) &
            (iqa_detalhado['saa'] == row['SAA']) &
            (iqa_detalhado['resultado'] == "Conforme")
        ]['mes'].notna().sum() - (row['expurgos_conf'] if pd.notna(row['expurgos_conf']) else 0)
    ),
    axis=1
  )

  plano_final['latitude'] = plano_final.apply(
    lambda row: lat_long[lat_long['cidade'] == row['cidade']]['latitude'].values[0] if any(lat_long['cidade'] == row['cidade']) else None,
    axis=1
  )

  plano_final['longitude'] = plano_final.apply(
      lambda row: lat_long[lat_long['cidade'] == row['cidade']]['longitude'].values[0] if any(lat_long['cidade'] == row['cidade']) else None,
      axis=1
  )

  plano_final['coordenada'] = plano_final['latitude'].astype(str) + ',' + plano_final['longitude'].astype(str)

  plano_final = plano_final.rename(columns={'SAA': 'saa'})

  plano_final.to_csv('plano_final.csv', index=False)


      # TABELA IQA FINAL

  prest_sigla_lower = prestadora_sigla.lower()

  coluna_iqa = f'iqa_{prest_sigla_lower}'

  iqa_final = pd.DataFrame({
      'ano': [ano],
      'ano_contratual': [ano_contratual],
      'trimestre': [trimestre],
      'mes': [mes],
      coluna_iqa: [iqa_prest],
      'iqa_vi': [iqa_vi],
      'meta': [iqa_meta]
  })

  iqa_final['nam_conf'] = iqa_final.apply(
    lambda row: (
        mun[
            (mun['mes'] == row['mes']) &
            (mun['ano'] == row['ano'])
        ]['nam_conf'].sum()
    ),
    axis=1
    )

  iqa_final['nam_realiz'] = iqa_final.apply(
  lambda row: (
      mun[
          (mun['mes'] == row['mes']) &
          (mun['ano'] == row['ano'])
      ]['nam_realiz'].sum()
  ),
  axis=1
  )

  iqa_final['iqa'] = iqa_final.apply(
    lambda row: row['nam_conf'] / row['nam_realiz'] if row['nam_realiz'] != 0 else None,
    axis=1
  )

  iqa_final['expurgos_conformes'] = iqa_final.apply(
  lambda row: (
      mun[
          (mun['mes'] == row['mes']) &
          (mun['ano'] == row['ano'])
      ]['expurgos_conformes'].sum()
  ),
  axis=1
  )

  iqa_final['expurgos_totais'] = iqa_final.apply(
  lambda row: (
      mun[
          (mun['mes'] == row['mes']) &
          (mun['ano'] == row['ano'])
      ]['expurgos'].sum()
  ),
  axis=1
  )

  iqa_final.to_csv('iqa_final.csv', index=False)


    
  # IQE

  if mes == "November" or mes == "February" or mes == "May" or mes == "August":

    # TABELA REALIZADOS

    realizados_iqe = plano_iqe_totais_parametros
    realizados_iqe.insert(0, "ano_contratual", ano_contratual)
    realizados_iqe.insert(0, "trimestre", trimestre)
    realizados_iqe = realizados_iqe.drop('desconsideracoes_arsal', axis=1)
    realizados_iqe['analises_realizadas'] = realizados_iqe.apply(
      lambda row: (
          "" if (
              iqe_detalhado2[
                  (iqe_detalhado2['trimestre'] == row['trimestre']) &
                  (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
                  (
                      ~iqe_detalhado2['analise'].isin(['DBO', 'DQO', 'Materiais Flutuantes', 'Óleos e Graxas', 'pH', 'Temperatura'])
                      if row['parametros'] == "Demais parâmetros"
                      else iqe_detalhado2['analise'] == row['parametros']
                  )
              ]['trimestre'].notna().sum()
          ) == 1 else (
              iqe_detalhado2[
                  (iqe_detalhado2['trimestre'] == row['trimestre']) &
                  (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
                  (
                      ~iqe_detalhado2['analise'].isin(['DBO', 'DQO', 'Materiais Flutuantes', 'Óleos e Graxas', 'pH', 'Temperatura'])
                      if row['parametros'] == "Demais parâmetros"
                      else iqe_detalhado2['analise'] == row['parametros']
                  )
              ]['trimestre'].notna().sum()
          )
      ),
      axis=1
    )

    realizados_iqe['expurgos'] = realizados_iqe.apply(
    lambda row: (
        iqe_detalhado2[
            (iqe_detalhado2['trimestre'] == row['trimestre']) &
            (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
            (iqe_detalhado2['expurgos'] == "EXPURGAR") &
            (
                ~iqe_detalhado2['analise'].isin(['DBO', 'DQO', 'Materiais Flutuantes', 'Óleos e Graxas', 'pH', 'Temperatura'])
                if row['parametros'] == "Demais Parâmetros"
                else iqe_detalhado2['analise'] == row['parametros']
            )
        ].shape[0]
    ),
    axis=1
    )

    realizados_iqe['nam_realiz'] = realizados_iqe.apply(
    lambda row: max(row['plano'], row['analises_realizadas']) - (row['expurgos'] if pd.notna(row['expurgos']) else 0),
    axis=1
    )

    print(realizados_iqe)


      
    # TABELA PARAMETROS

    parametros_iqe = realizados_iqe.drop('plano', axis=1)

    parametros_iqe['expurgos_conformes'] = parametros_iqe.apply(
    lambda row: (
        iqe_detalhado2[
            (iqe_detalhado2['trimestre'] == row['trimestre']) &
            (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
            (iqe_detalhado2['expurgos'] == "EXPURGAR") &
            (iqe_detalhado2['resultado'] == "Conforme") &
            (
                ~iqe_detalhado2['analise'].isin(['DBO', 'DQO', 'Materiais Flutuantes', 'Óleos e Graxas', 'pH', 'Temperatura'])
                if row['parametros'] == "Demais Parâmetros"
                else iqe_detalhado2['analise'] == row['parametros']
            )
        ].shape[0]
    ),
    axis=1
    )

    parametros_iqe['nam_conf'] = parametros_iqe.apply(
    lambda row: (
        "" if (
            iqe_detalhado2[
                (iqe_detalhado2['trimestre'] == row['trimestre']) &
                (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
                (iqe_detalhado2['resultado'] == "Conforme") &
                (
                    ~iqe_detalhado2['analise'].isin(['DBO', 'DQO', 'Materiais Flutuantes', 'Óleos e Graxas', 'pH', 'Temperatura'])
                    if row['parametros'] == "Demais parâmetros"
                    else iqe_detalhado2['analise'] == row['parametros']
                )
            ]['trimestre'].notna().sum() - (row['expurgos_conformes'] if pd.notna(row['expurgos_conformes']) else 0)
        ) == 1 else (
            iqe_detalhado2[
                (iqe_detalhado2['trimestre'] == row['trimestre']) &
                (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
                (iqe_detalhado2['resultado'] == "Conforme") &
                (
                    ~iqe_detalhado2['analise'].isin(['DBO', 'DQO', 'Materiais Flutuantes', 'Óleos e Graxas', 'pH', 'Temperatura'])
                    if row['parametros'] == "Demais parâmetros"
                    else iqe_detalhado2['analise'] == row['parametros']
                )
            ]['trimestre'].notna().sum() - (row['expurgos_conformes'] if pd.notna(row['expurgos_conformes']) else 0)
        )
    ),
    axis=1
    )

    parametros_iqe['expurgos_totais'] = parametros_iqe.apply(
    lambda row: (
        iqe_detalhado2[
            (iqe_detalhado2['trimestre'] == row['trimestre']) &
            (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
            (iqe_detalhado2['expurgos'] == "EXPURGAR") &
            ~iqe_detalhado2['analise'].isin(['DBO', 'DQO', 'Materiais Flutuantes', 'Óleos e Graxas', 'pH', 'Temperatura'])
        ].shape[0]
    ) if row['parametros'] == "Demais parâmetros" else (
        iqe_detalhado2[
            (iqe_detalhado2['trimestre'] == row['trimestre']) &
            (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
            (iqe_detalhado2['expurgos'] == "EXPURGAR") &
            (iqe_detalhado2['analise'] == row['parametros'])
        ].shape[0]
    ),
    axis=1
    )

    parametros_iqe['nam_realiz'] = parametros_iqe.apply(
      lambda row: realizados_iqe[
          (realizados_iqe['trimestre'] == row['trimestre']) &
          (realizados_iqe['ano_contratual'] == row['ano_contratual']) &
          (realizados_iqe['parametros'] == row['parametros'])
      ]['nam_realiz'].values[0] if any(
          (realizados_iqe['trimestre'] == row['trimestre']) &
          (realizados_iqe['ano_contratual'] == row['ano_contratual']) &
          (realizados_iqe['parametros'] == row['parametros'])
      ) else None,
      axis=1
    )

    parametros_iqe['iqe'] = parametros_iqe.apply(
      lambda row: row['nam_conf'] / row['nam_realiz'] if row['nam_realiz'] != 0 else None,
      axis=1
    )

      

    # TABELA MUN

    mun_iqe = plano_iqe_totais_municipio

    mun_iqe.insert(0, 'trimestre', trimestre)
    mun_iqe.insert(1, 'ano_contratual', ano_contratual)

    mun_iqe['analises_realizadas'] = mun_iqe.apply(
      lambda row: (
          "" if (
              iqe_detalhado2[
                  (iqe_detalhado2['trimestre'] == row['trimestre']) &
                  (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
                  (iqe_detalhado2['cidade'] == row['cidade'])
              ]['trimestre'].notna().sum()
          ) == 1 else (
              iqe_detalhado2[
                  (iqe_detalhado2['trimestre'] == row['trimestre']) &
                  (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
                  (iqe_detalhado2['cidade'] == row['cidade'])
              ]['trimestre'].notna().sum()
          )
      ),
      axis=1
    )

    mun_iqe['expurgos'] = mun_iqe.apply(
    lambda row: (
        iqe_detalhado2[
            (iqe_detalhado2['trimestre'] == row['trimestre']) &
            (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
            (iqe_detalhado2['cidade'] == row['cidade']) &
            (iqe_detalhado2['expurgos'] == "EXPURGAR")
        ].shape[0]
    ),
    axis=1
    )

    mun_iqe['nam_realiz'] = mun_iqe.apply(
    lambda row: max(
        row['plano'] if pd.notna(row['plano']) else 0,
        row['analises_realizadas'] if row['analises_realizadas'] != "" and pd.notna(row['analises_realizadas']) else 0
    ) - (row['expurgos'] if pd.notna(row['expurgos']) else 0),
    axis=1
    )

    mun_iqe['expurgos_conformes'] = mun_iqe.apply(
    lambda row: (
        iqe_detalhado2[
            (iqe_detalhado2['trimestre'] == row['trimestre']) &
            (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
            (iqe_detalhado2['cidade'] == row['cidade']) &
            (iqe_detalhado2['resultado'] == "Conforme") &
            (iqe_detalhado2['expurgos'] == "EXPURGAR")
        ].shape[0]
    ),
    axis=1
    )

    mun_iqe['nam_conf'] = mun_iqe.apply(
    lambda row: (
        "" if (
            iqe_detalhado2[
                (iqe_detalhado2['trimestre'] == row['trimestre']) &
                (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
                (iqe_detalhado2['cidade'] == row['cidade']) &
                (iqe_detalhado2['resultado'] == "Conforme")
            ]['trimestre'].notna().sum() - (row['expurgos_conformes'] if pd.notna(row['expurgos_conformes']) else 0)
        ) == 1 else (
            iqe_detalhado2[
                (iqe_detalhado2['trimestre'] == row['trimestre']) &
                (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
                (iqe_detalhado2['cidade'] == row['cidade']) &
                (iqe_detalhado2['resultado'] == "Conforme")
            ]['trimestre'].notna().sum() - (row['expurgos_conformes'] if pd.notna(row['expurgos_conformes']) else 0)
        )
    ),
    axis=1
    )

    mun_iqe['iqe'] = mun_iqe.apply(
      lambda row: row['nam_conf'] / row['nam_realiz'] if row['nam_realiz'] != 0 else None,
      axis=1
    )

    lat_long = pd.read_excel('lat_long.xlsx')

    mun_iqe['latitude'] = mun_iqe.apply(
      lambda row: lat_long[lat_long['cidade'] == row['cidade']]['latitude'].values[0] if any(lat_long['cidade'] == row['cidade']) else None,
      axis=1
    )

    mun_iqe['longitude'] = mun_iqe.apply(
        lambda row: lat_long[lat_long['cidade'] == row['cidade']]['longitude'].values[0] if any(lat_long['cidade'] == row['cidade']) else None,
        axis=1
    )

    mun_iqe['coordenada'] = mun_iqe['latitude'].astype(str) + ',' + mun_iqe['longitude'].astype(str)



    # TABELA PLANO

    plano_final_iqe = plano_iqe_completo

    plano_final_iqe['amostras_realizadas'] = plano_final_iqe.apply(
      lambda row: (
          iqe_detalhado2[
              (iqe_detalhado2['trimestre'] == row['trimestre']) &
              (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
              (iqe_detalhado2['cidade'] == row['cidade']) &
              (iqe_detalhado2['ETE'] == row['ETE']) &
              ~iqe_detalhado2['analise'].isin(['DBO', 'DQO', 'pH', 'Óleos e Graxas', 'OD'])
          ]['trimestre'].notna().sum()
      ) if row['parametros'] == "Demais Parâmetros" else (
          iqe_detalhado2[
              (iqe_detalhado2['trimestre'] == row['trimestre']) &
              (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
              (iqe_detalhado2['analise'] == row['parametros']) &
              (iqe_detalhado2['cidade'] == row['cidade']) &
              (iqe_detalhado2['ETE'] == row['ETE'])
          ]['trimestre'].notna().sum()
      ),
      axis=1
    )

    plano_final_iqe['expurgos'] = plano_final_iqe.apply(
      lambda row: (
          iqe_detalhado[
              (iqe_detalhado2['trimestre'] == row['trimestre']) &
              (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
              (iqe_detalhado2['cidade'] == row['cidade']) &
              (iqe_detalhado2['ETE'] == row['ETE']) &
              (iqe_detalhado2['expurgos'] == "EXPURGAR") &
              ~iqe_detalhado2['analise'].isin(['DBO', 'DQO', 'pH', 'Óleos e Graxas', 'OD'])
          ]['trimestre'].notna().sum()
      ) if row['parametros'] == "Demais Parâmetros" else (
          iqe_detalhado2[
              (iqe_detalhado2['trimestre'] == row['trimestre']) &
              (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
              (iqe_detalhado2['analise'] == row['parametros']) &
              (iqe_detalhado2['cidade'] == row['cidade']) &
              (iqe_detalhado2['expurgos'] == "EXPURGAR") &
              (iqe_detalhado2['ETE'] == row['ETE'])
          ]['trimestre'].notna().sum()
      ),
      axis=1
    )

    plano_final_iqe['nam_realiz'] = plano_final_iqe[['plano', 'amostras_realizadas']].max(axis=1) - plano_final_iqe['expurgos'].fillna(0)

    plano_final_iqe['expurgos_conf'] = plano_final_iqe.apply(
      lambda row: (
          iqe_detalhado[
              (iqe_detalhado2['trimestre'] == row['trimestre']) &
              (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
              (iqe_detalhado2['cidade'] == row['cidade']) &
              (iqe_detalhado2['ETE'] == row['ETE']) &
              (iqe_detalhado2['expurgos'] == "EXPURGAR") &
              (iqe_detalhado2['resultado'] == "Conforme") &
              ~iqe_detalhado2['analise'].isin(['DBO', 'DQO', 'pH', 'Óleos e Graxas', 'OD'])
          ]['trimestre'].notna().sum()
      ) if row['parametros'] == "Demais Parâmetros" else (
          iqe_detalhado2[
              (iqe_detalhado2['trimestre'] == row['trimestre']) &
              (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
              (iqe_detalhado2['analise'] == row['parametros']) &
              (iqe_detalhado2['cidade'] == row['cidade']) &
              (iqe_detalhado2['ETE'] == row['ETE']) &
              (iqe_detalhado2['expurgos'] == "EXPURGAR") &
              (iqe_detalhado2['resultado'] == "Conforme")
          ]['trimestre'].notna().sum()
      ),
      axis=1
    )

    plano_final_iqe['nam_conf'] = plano_final_iqe.apply(
      lambda row: (
          iqe_detalhado2[
              (iqe_detalhado2['trimestre'] == row['trimestre']) &
              (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
              (iqe_detalhado2['cidade'] == row['cidade']) &
              (iqe_detalhado2['ETE'] == row['ETE']) &
              (iqe_detalhado2['resultado'] == "Conforme") &
              ~iqe_detalhado2['analise'].isin(['DBO', 'DQO', 'pH', 'Óleos e Graxas', 'OD'])
          ]['trimestre'].notna().sum() - (row['expurgos_conf'] if pd.notna(row['expurgos_conf']) else 0)
      ) if row['parametros'] == "Demais Parâmetros" else (
          iqe_detalhado2[
              (iqe_detalhado2['trimestre'] == row['trimestre']) &
              (iqe_detalhado2['ano_contratual'] == row['ano_contratual']) &
              (iqe_detalhado2['analise'] == row['parametros']) &
              (iqe_detalhado2['cidade'] == row['cidade']) &
              (iqe_detalhado2['ETE'] == row['ETE']) &
              (iqe_detalhado2['resultado'] == "Conforme")
          ]['trimestre'].notna().sum() - (row['expurgos_conf'] if pd.notna(row['expurgos_conf']) else 0)
      ),
      axis=1
    )

    plano_final_iqe['latitude'] = plano_final_iqe.apply(
      lambda row: lat_long[lat_long['cidade'] == row['cidade']]['latitude'].values[0] if any(lat_long['cidade'] == row['cidade']) else None,
      axis=1
    )

    plano_final_iqe['longitude'] = plano_final_iqe.apply(
        lambda row: lat_long[lat_long['cidade'] == row['cidade']]['longitude'].values[0] if any(lat_long['cidade'] == row['cidade']) else None,
        axis=1
    )

    plano_final_iqe['coordenada'] = plano_final_iqe['latitude'].astype(str) + ',' + plano_final_iqe['longitude'].astype(str)


      
    # TABELA IQE FINAL

    prest_sigla_lower = prestadora_sigla.lower()

    coluna_iqe = f'iqe_{prest_sigla_lower}'

    iqe_final = pd.DataFrame({
        'trimestre': [trimestre],
        'ano_contratual': [ano_contratual],
        coluna_iqe: [iqe_prest],
        'iqe_vi': [iqe_vi],
        'meta': [iqe_meta]
    })

    iqe_final['nam_conf'] = iqe_final.apply(
    lambda row: (
        mun_iqe[
            (mun_iqe['trimestre'] == row['trimestre']) &
            (mun_iqe['ano_contratual'] == row['ano_contratual'])
        ]['nam_conf'].sum()
    ),
    axis=1
    )

    iqe_final['nam_realiz'] = iqe_final.apply(
    lambda row: (
        mun_iqe[
            (mun_iqe['trimestre'] == row['trimestre']) &
            (mun_iqe['ano_contratual'] == row['ano_contratual'])
        ]['nam_realiz'].sum()
    ),
    axis=1
    )

    iqe_final['iqe'] = iqe_final.apply(
      lambda row: row['nam_conf'] / row['nam_realiz'] if row['nam_realiz'] != 0 else None,
      axis=1
    )

    iqe_final['expurgos_conformes'] = iqe_final.apply(
    lambda row: (
        mun_iqe[
            (mun_iqe['trimestre'] == row['trimestre']) &
            (mun_iqe['ano_contratual'] == row['ano_contratual'])
        ]['expurgos_conformes'].sum()
    ),
    axis=1
    )

    iqe_final['expurgos_totais'] = iqe_final.apply(
    lambda row: (
        mun_iqe[
            (mun_iqe['trimestre'] == row['trimestre']) &
            (mun_iqe['ano_contratual'] == row['ano_contratual'])
        ]['expurgos'].sum()
    ),
    axis=1
    )

    return realizados, parametros, mun, plano_final, iqa_final, realizados_iqe, parametros_iqe, mun_iqe, plano_final_iqe, iqe_final

  else:

    return realizados, parametros, mun, plano_final, iqa_final



#############################################################################################################################################################################################################################

# ATUALIZAÇÃO DO BANCO DE DADOS

def alimentar_bd (prestadora_sigla, ano, mes, iqa_detalhado, realizados, parametros, mun, plano_final, iqa_final, iqe_detalhado, realizados_iqe, parametros_iqe, mun_iqe, plano_final_iqe, iqe_final, trimestre, ano_contratual):

  prestadora_iqa = prestadora_sigla + '_IQA'

  resposta = supabase.table(prestadora_iqa).select("*").execute()

  df = pd.DataFrame(resposta.data)

  prestadora_detalhado = prestadora_sigla + '_IQA_DETALHADO'
  prestadora_realizados = prestadora_sigla + '_IQA_REALIZADOS'
  prestadora_parametros = prestadora_sigla + '_IQA_PARAMETROS'
  prestadora_mun = prestadora_sigla + '_IQA_MUN'
  prestadora_plano = prestadora_sigla + '_IQA_PLANO'

  resposta = (
      supabase
      .table(prestadora_detalhado)
      .select("id")
      .eq("ano", ano)
      .eq("mes", mes)
      .execute()
  )

  if resposta.data:

      (
          supabase
          .table(prestadora_detalhado)
          .delete()
          .eq("ano", ano)
          .eq("mes", mes)
          .execute()
      )
      (
          supabase
          .table(prestadora_realizados)
          .delete()
          .eq("ano", ano)
          .eq("ano", mes)
          .execute()
      )
      (
          supabase
          .table(prestadora_parametros)
          .delete()
          .eq("ano", ano)
          .eq("mes", mes)
          .execute()
      )
      (
          supabase
          .table(prestadora_mun)
          .delete()
          .eq("ano", ano)
          .eq("mes", mes)
          .execute()
      )
      (
          supabase
          .table(prestadora_plano)
          .delete()
          .eq("ano", ano)
          .eq("mes", mes)
          .execute()
      )
      (
          supabase
          .table(prestadora_iqa)
          .delete()
          .eq("ano", ano)
          .eq("mes", mes)
          .execute()
      )

      iqa_detalhado = iqa_detalhado.astype(object).where(pd.notnull(iqa_detalhado), None)
      realizados = realizados.astype(object).where(pd.notnull(realizados), None)
      parametros = parametros.astype(object).where(pd.notnull(parametros), None)
      mun = mun.astype(object).where(pd.notnull(mun), None)
      plano_final = plano_final.astype(object).where(pd.notnull(plano_final), None)
      iqa_final = iqa_final.astype(object).where(pd.notnull(iqa_final), None)

      supabase.table(prestadora_detalhado).insert(iqa_detalhado.to_dict(orient="records")).execute()
      supabase.table(prestadora_realizados).insert(realizados.to_dict(orient="records")).execute()
      supabase.table(prestadora_parametros).insert(parametros.to_dict(orient="records")).execute()
      supabase.table(prestadora_mun).insert(mun.to_dict(orient="records")).execute()
      supabase.table(prestadora_plano).insert(plano_final.to_dict(orient="records")).execute()
      supabase.table(prestadora_iqa).insert(iqa_final.to_dict(orient="records")).execute()

  else:

      iqa_detalhado = iqa_detalhado.astype(object).where(pd.notnull(iqa_detalhado), None)
      realizados = realizados.astype(object).where(pd.notnull(realizados), None)
      parametros = parametros.astype(object).where(pd.notnull(parametros), None)
      mun = mun.astype(object).where(pd.notnull(mun), None)
      plano_final = plano_final.astype(object).where(pd.notnull(plano_final), None)
      iqa_final = iqa_final.astype(object).where(pd.notnull(iqa_final), None)

      supabase.table(prestadora_detalhado).insert(iqa_detalhado.to_dict(orient="records")).execute()
      supabase.table(prestadora_realizados).insert(realizados.to_dict(orient="records")).execute()
      supabase.table(prestadora_parametros).insert(parametros.to_dict(orient="records")).execute()
      supabase.table(prestadora_mun).insert(mun.to_dict(orient="records")).execute()
      supabase.table(prestadora_plano).insert(plano_final.to_dict(orient="records")).execute()
      supabase.table(prestadora_iqa).insert(iqa_final.to_dict(orient="records")).execute()


  if mes == "November" or mes == "February" or mes == "May" or mes == "August":


    prestadora_iqe = prestadora_sigla + '_IQE'

    resposta_iqe = supabase.table(prestadora_iqe).select("*").execute()

    df_iqe = pd.DataFrame(resposta.data)

    prestadora_detalhado_iqe = prestadora_sigla + '_IQE_DETALHADO'
    prestadora_realizados_iqe = prestadora_sigla + '_IQE_REALIZADOS'
    prestadora_parametros_iqe = prestadora_sigla + '_IQE_PARAMETROS'
    prestadora_mun_iqe = prestadora_sigla + '_IQE_MUN'
    prestadora_plano_iqe = prestadora_sigla + '_IQE_PLANO'

    resposta_iqe = (
        supabase
        .table(prestadora_detalhado_iqe)
        .select("id")
        .eq("trimestre", trimestre)
        .eq("ano_contratual", ano_contratual)
        .execute()
    )

    if resposta_iqe.data:

        (
            supabase
            .table(prestadora_detalhado_iqe)
            .delete()
            .eq("trimestre", trimestre)
            .eq("ano_contratual", ano_contratual)
            .execute()
        )
        (
            supabase
            .table(prestadora_realizados_iqe)
            .delete()
            .eq("trimestre", trimestre)
            .eq("ano_contratual", ano_contratual)
            .execute()
        )
        (
            supabase
            .table(prestadora_parametros_iqe)
            .delete()
            .eq("trimestre", trimestre)
            .eq("ano_contratual", ano_contratual)
            .execute()
        )
        (
            supabase
            .table(prestadora_mun_iqe)
            .delete()
            .eq("trimestre", trimestre)
            .eq("ano_contratual", ano_contratual)
            .execute()
        )
        (
            supabase
            .table(prestadora_plano_iqe)
            .delete()
            .eq("trimestre", trimestre)
            .eq("ano_contratual", ano_contratual)
            .execute()
        )
        (
            supabase
            .table(prestadora_iqe)
            .delete()
            .eq("trimestre", trimestre)
            .eq("ano_contratual", ano_contratual)
            .execute()
        )


        iqe_detalhado = iqe_detalhado.astype(object).where(pd.notnull(iqe_detalhado), None)
        realizados_iqe = realizados_iqe.astype(object).where(pd.notnull(realizados_iqe), None)
        parametros_iqe = parametros_iqe.astype(object).where(pd.notnull(parametros_iqe), None)
        mun_iqe = mun_iqe.astype(object).where(pd.notnull(mun_iqe), None)
        plano_final_iqe = plano_final_iqe.astype(object).where(pd.notnull(plano_final_iqe), None)
        iqe_final = iqe_final.astype(object).where(pd.notnull(iqe_final), None)

        supabase.table(prestadora_detalhado_iqe).insert(iqe_detalhado.to_dict(orient="records")).execute()
        supabase.table(prestadora_realizados_iqe).insert(realizados_iqe.to_dict(orient="records")).execute()
        supabase.table(prestadora_parametros_iqe).insert(parametros_iqe.to_dict(orient="records")).execute()
        supabase.table(prestadora_mun_iqe).insert(mun_iqe.to_dict(orient="records")).execute()
        supabase.table(prestadora_plano_iqe).insert(plano_final_iqe.to_dict(orient="records")).execute()
        supabase.table(prestadora_iqe).insert(iqe_final.to_dict(orient="records")).execute()


    else:

        iqe_detalhado = iqe_detalhado.astype(object).where(pd.notnull(iqe_detalhado), None)
        realizados_iqe = realizados_iqe.astype(object).where(pd.notnull(realizados_iqe), None)
        parametros_iqe = parametros_iqe.astype(object).where(pd.notnull(parametros_iqe), None)
        mun_iqe = mun_iqe.astype(object).where(pd.notnull(mun_iqe), None)
        plano_final_iqe = plano_final_iqe.astype(object).where(pd.notnull(plano_final_iqe), None)
        iqe_final = iqe_final.astype(object).where(pd.notnull(iqe_final), None)

        supabase.table(prestadora_detalhado_iqe).insert(iqe_detalhado.to_dict(orient="records")).execute()
        supabase.table(prestadora_realizados_iqe).insert(realizados_iqe.to_dict(orient="records")).execute()
        supabase.table(prestadora_parametros_iqe).insert(parametros_iqe.to_dict(orient="records")).execute()
        supabase.table(prestadora_mun_iqe).insert(mun_iqe.to_dict(orient="records")).execute()
        supabase.table(prestadora_plano_iqe).insert(plano_final_iqe.to_dict(orient="records")).execute()
        supabase.table(prestadora_iqe).insert(iqe_final.to_dict(orient="records")).execute()



##############################################################################################################################################################################################################################

# BACKUP DO BANCO DE DADOS

load_dotenv()

def backup_bd ():

  PASTA_BACKUP = os.getenv("PASTA_BACKUP", "./backups")

  os.makedirs(PASTA_BACKUP, exist_ok=True)

  TABELAS = [

      "BRK_IQA_DETALHADO",
      "ADS_IQA_DETALHADO",
      "VAA_IQA_DETALHADO",

      "BRK_IQA_REALIZADOS",
      "ADS_IQA_REALIZADOS",
      "VAA_IQA_REALIZADOS",

      "BRK_IQA_PARAMETROS",
      "ADS_IQA_PARAMETROS",
      "VAA_IQA_PARAMETROS",

      "BRK_IQA_MUN",
      "ADS_IQA_MUN",
      "VAA_IQA_MUN",

      "BRK_IQA_PLANO",
      "ADS_IQA_PLANO",
      "VAA_IQA_PLANO",

      "BRK_IQA",
      "ADS_IQA",
      "VAA_IQA",

      "BRK_IQE_DETALHADO",
      "ADS_IQE_DETALHADO",
      "VAA_IQE_DETALHADO",

      "BRK_IQE_REALIZADOS",
      "ADS_IQE_REALIZADOS",
      "VAA_IQE_REALIZADOS",

      "BRK_IQE_PARAMETROS",
      "ADS_IQE_PARAMETROS",
      "VAA_IQE_PARAMETROS",

      "BRK_IQE_MUN",
      "ADS_IQE_MUN",
      "VAA_IQE_MUN",

      "BRK_IQE_PLANO",
      "ADS_IQE_PLANO",
      "VAA_IQE_PLANO",

      "BRK_IQE",
      "ADS_IQE",
      "VAA_IQE",

  ]

    
  for tabela in TABELAS:

      print(f"\nExportando {tabela}...")

      try:

          todos_dados = []
          lote = 1000
          inicio = 0

          while True:

              resposta = (
                  supabase
                  .table(tabela)
                  .select("*")
                  .range(inicio, inicio + lote - 1)
                  .execute()
              )

              dados = resposta.data

              if len(dados) == 0:
                  break

              todos_dados.extend(dados)

              print(f"   {len(todos_dados)} registros baixados...")

              inicio += lote

          df = pd.DataFrame(todos_dados)

          arquivo = os.path.join(
              PASTA_BACKUP,
              f"{tabela}.csv"
          )

          df.to_csv(
              arquivo,
              index=False,
              encoding="utf-8-sig"
          )

          data_backup = datetime.now().strftime("%Y-%m-%d")

          caminho_storage = (
                f"{data_backup}/{tabela}.csv"
          )


          enviar_backup_storage(
                arquivo,
                caminho_storage
          )


          st.success(
                f"Backup enviado: {tabela}.csv"
          )

          print(f"✅ {tabela}: {len(df)} registros exportados.")

      except Exception as erro:

          print(f"❌ ERRO em {tabela}")

          print(erro)

  return "✅ FINALIZADO!"

def enviar_backup_storage(caminho_arquivo, nome_arquivo):

    with open(caminho_arquivo, "rb") as arquivo:

        resposta = (
            supabase
            .storage
            .from_("BACKUPS_IQA")
            .upload(
                nome_arquivo,
                arquivo,
                {
                    "content-type": "text/csv",
                    "upsert": "true"
                }
            )
        )

    return resposta



##############################################################################################################################################################################################################################

# ELABORAÇÃO DOS GRÁFICOS TEMPORAIS

    # 4.1 IQAs / IQEs


def graf_iqa_iqe (prestadora_sigla, mes):

  nome_tabela_iqa = prestadora_sigla + "_IQA"
  nome_tabela_iqe = prestadora_sigla + "_IQE"

  # IQA

  try:

    response = supabase.table(nome_tabela_iqa).select(
        "mes, ano, iqa_ads, iqa_vi, iqa"
    ).execute()

    df = pd.DataFrame(response.data)

    if df.empty:
        print("❌ Nenhum dado encontrado!")
        exit()

    df['ano'] = pd.to_numeric(
        df['ano'],
        errors='coerce'
    )

    df = df.dropna(subset=['ano'])

    df['ano'] = df['ano'].astype(int)

    mapa_meses = {
        'january': 1,
        'february': 2,
        'march': 3,
        'april': 4,
        'may': 5,
        'june': 6,
        'july': 7,
        'august': 8,
        'september': 9,
        'october': 10,
        'november': 11,
        'december': 12,

        'janeiro': 1,
        'fevereiro': 2,
        'março': 3,
        'abril': 4,
        'maio': 5,
        'junho': 6,
        'julho': 7,
        'agosto': 8,
        'setembro': 9,
        'outubro': 10,
        'novembro': 11,
        'dezembro': 12
    }

    df['mes_numero'] = (
        df['mes']
        .astype(str)
        .str.strip()
        .str.lower()
        .map(mapa_meses)
    )

    if df['mes_numero'].isna().any():

        meses_invalidos = df.loc[
            df['mes_numero'].isna(),
            'mes'
        ].unique()

        print(
            f"⚠️ Meses não reconhecidos: "
            f"{meses_invalidos}"
        )

    df = df.dropna(subset=['mes_numero'])

    df['mes_numero'] = df['mes_numero'].astype(int)

    ultimo_ano = df['ano'].max()

    print(f"Último ano encontrado: {ultimo_ano}")

    df_ano = df[
        df['ano'] == ultimo_ano
    ].copy()

    if df_ano.empty:
        print("❌ Nenhum dado encontrado para o último ano!")
        exit()

    ultimo_mes = df_ano['mes_numero'].max()

    print(f"Último mês disponível: {ultimo_mes}")

    df_ano['data'] = pd.to_datetime(
        dict(
            year=df_ano['ano'],
            month=df_ano['mes_numero'],
            day=1
        )
    )

    ultima_data = df_ano['data'].max()

    data_inicial = (
        ultima_data - pd.DateOffset(months=5)
    )

    df_filtrado = df_ano[
        (df_ano['data'] >= data_inicial) &
        (df_ano['data'] <= ultima_data)
    ].copy()

    df_filtrado = df_filtrado.sort_values('data')

    quantidade = len(df_filtrado)

    print(
        f"Meses encontrados: {quantidade}/6"
    )

    print(
        "Detalhes:",
        df_filtrado['mes'].tolist()
    )

    if quantidade == 0:
        print(
            "❌ Nenhum mês encontrado "
            "para o período!"
        )
        exit()

    nomes_meses = {
        1: 'Jan',
        2: 'Fev',
        3: 'Mar',
        4: 'Abr',
        5: 'Mai',
        6: 'Jun',
        7: 'Jul',
        8: 'Ago',
        9: 'Set',
        10: 'Out',
        11: 'Nov',
        12: 'Dez'
    }

    meses_labels = [
        nomes_meses[m]
        for m in df_filtrado['mes_numero']
    ]

    x = np.arange(len(meses_labels))

    width = 0.25

    iqa_ads = pd.to_numeric(
        df_filtrado['iqa_ads'],
        errors='coerce'
    ).fillna(0)

    iqa_vi = pd.to_numeric(
        df_filtrado['iqa_vi'],
        errors='coerce'
    ).fillna(0)

    iqa = pd.to_numeric(
        df_filtrado['iqa'],
        errors='coerce'
    ).fillna(0)

    fig, ax = plt.subplots(figsize=(12, 4.5))

    x = np.arange(len(df_filtrado))
    width = 0.25

    cor_arsal = '#0878E8'
    cor_ads = '#9A9D17'
    cor_vi = '#EC008C'

    barras1 = ax.bar(
        x - width,
        iqa,
        width,
        label='IQA - Arsal',
        color=cor_arsal,
        edgecolor='white',
        linewidth=0.8
    )

    barras2 = ax.bar(
        x,
        iqa_ads,
        width,
        label='IQA - ADS',
        color=cor_ads,
        edgecolor='white',
        linewidth=0.8
    )

    barras3 = ax.bar(
        x + width,
        iqa_vi,
        width,
        label='IQA - VI',
        color=cor_vi,
        edgecolor='white',
        linewidth=0.8
    )

    ax.set_ylim(0, 1.05)

    ax.set_yticks([
        0,
        0.25,
        0.50,
        0.75,
        1.00
    ])

    ax.set_yticklabels([
        '0,00%',
        '25,00%',
        '50,00%',
        '75,00%',
        '100,00%'
    ])

    ax.set_xticks(x)

    ax.set_xticklabels(
        meses_labels,
        fontsize=10
    )

    ax.legend(
        loc='upper left',
        bbox_to_anchor=(0, 1.12),
        ncol=3,
        frameon=False,
        fontsize=10
    )

    ax.grid(
        axis='y',
        linestyle='-',
        linewidth=0.7,
        alpha=0.35
    )

    ax.set_axisbelow(True)

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    ax.set_xlabel('')
    ax.set_ylabel('')

    def colocar_valores(barras):

        for barra in barras:

            altura = barra.get_height()

            if altura > 0:

                ax.text(
                    barra.get_x() + barra.get_width() / 2,
                    altura + 0.008,
                    f'{altura * 100:.1f}%'.replace('.', ','),
                    ha='center',
                    va='bottom',
                    fontsize=10.0
                )

    colocar_valores(barras1)
    colocar_valores(barras2)
    colocar_valores(barras3)

    ax.tick_params(
        axis='both',
        length=0
    )

    plt.tight_layout()

    pasta = "graficos"

    if not os.path.exists(pasta):
        os.makedirs(pasta)

    timestamp = datetime.now().strftime(
        "%Y%m%d_%H%M%S"
    )

    nome_arquivo = (
        "IQAS.png"
    )

    plt.savefig(
        nome_arquivo,
        dpi=300,
        bbox_inches='tight',
        format='png'
    )

    print(
        f"✅ Gráfico salvo: {nome_arquivo}"
    )

    plt.show()

  except Exception as e:

    print(f"❌ Erro: {e}")

    import traceback
    traceback.print_exc()


    
    # IQE

  if mes == "November" or mes == "February" or mes == "May" or mes == "August":

    try:
      response = supabase.table(nome_tabela_iqe).select(
          "trimestre, ano_contratual, iqe_ads, iqe_vi, iqe"
      ).execute()

      df = pd.DataFrame(response.data)

      if df.empty:
          print("❌ Nenhum dado encontrado!")
          exit()

      ultimo_ano = df['ano_contratual'].max()

      print(f"Consultando último ano contratual: {ultimo_ano}")

      df_filtrado = df[
          df['ano_contratual'] == ultimo_ano
      ].copy()

      ordem_trimestres = [
          '1° trimestre',
          '2° trimestre',
          '3° trimestre',
          '4° trimestre'
      ]

      df_filtrado['trimestre'] = pd.Categorical(
          df_filtrado['trimestre'],
          categories=ordem_trimestres,
          ordered=True
      )

      df_filtrado = df_filtrado.sort_values('trimestre')

      df_filtrado = df_filtrado.dropna(
          subset=['trimestre']
      )

      quantidade = len(df_filtrado)

      print(
          f"Trimestres encontrados: {quantidade}/4"
      )

      print(
          f"Detalhes: "
          f"{df_filtrado['trimestre'].astype(str).tolist()}\n"
      )

      if quantidade == 0:
          print(
              "❌ Nenhum trimestre encontrado "
              "para este ano contratual!"
          )
          exit()

      trimestres_labels = [
          t.replace('° trimestre', '° trim')
          for t in df_filtrado['trimestre'].astype(str)
      ]

      x = np.arange(
          len(trimestres_labels)
      )

      width = 0.25

      iqe_ads = pd.to_numeric(
          df_filtrado['iqe_ads'],
          errors='coerce'
      ).fillna(0)

      iqe_vi = pd.to_numeric(
          df_filtrado['iqe_vi'],
          errors='coerce'
      ).fillna(0)

      iqe = pd.to_numeric(
          df_filtrado['iqe'],
          errors='coerce'
      ).fillna(0)

      fig, ax = plt.subplots(
          figsize=(12, 4.5)
      )

      barras1 = ax.bar(
          x - width,
          iqe_ads,
          width,
          label='IQE - ADS',
          color='#9A9D17',
          edgecolor='white',
          linewidth=0.8
      )

      barras2 = ax.bar(
          x,
          iqe_vi,
          width,
          label='IQE - VI',
          color='#EC008C',
          edgecolor='white',
          linewidth=0.8
      )

      barras3 = ax.bar(
          x + width,
          iqe,
          width,
          label='IQE',
          color='#0878E8',
          edgecolor='white',
          linewidth=0.8
      )

      ax.set_ylim(
          0,
          1.05
      )

      ax.set_yticks([
          0,
          0.25,
          0.50,
          0.75,
          1.00
      ])

      ax.set_yticklabels([
          '0,00%',
          '25,00%',
          '50,00%',
          '75,00%',
          '100,00%'
      ])

      ax.set_xticks(x)

      ax.set_xticklabels(
          trimestres_labels,
          fontsize=10
      )

      ax.legend(
          loc='upper left',
          bbox_to_anchor=(0, 1.12),
          ncol=3,
          frameon=False,
          fontsize=10
      )

      ax.grid(
          axis='y',
          linestyle='-',
          linewidth=0.7,
          alpha=0.35
      )

      ax.set_axisbelow(True)

      ax.spines['top'].set_visible(False)
      ax.spines['right'].set_visible(False)

      ax.set_xlabel('')
      ax.set_ylabel('')

      def colocar_valores(barras):

          for barra in barras:

              altura = barra.get_height()

              if altura > 0:

                  ax.text(
                      barra.get_x()
                      + barra.get_width() / 2,
                      altura + 0.008,
                      f'{altura * 100:.1f}%'.replace(
                          '.',
                          ','
                      ),
                      ha='center',
                      va='bottom',
                      fontsize=10.0
                  )

      colocar_valores(barras1)
      colocar_valores(barras2)
      colocar_valores(barras3)

      ax.tick_params(
          axis='both',
          length=0
      )

      plt.tight_layout()

      pasta = "graficos"

      if not os.path.exists(pasta):
          os.makedirs(pasta)

      timestamp = datetime.now().strftime(
          "%Y%m%d_%H%M%S"
      )

      nome_arquivo = (
          "IQES.png"
      )

      plt.savefig(
          nome_arquivo,
          dpi=300,
          bbox_inches='tight',
          format='png'
      )

      print(
          f"✅ Gráfico salvo: {nome_arquivo}"
      )

      if quantidade < 4:

          print(
              f"⚠️ O último ano contratual possui apenas "
              f"{quantidade}/4 trimestres. "
              f"O gráfico mostra somente os "
              f"trimestres existentes."
          )

      plt.show()

    except Exception as e:

      print(f"❌ Erro: {e}")

      import traceback
      traceback.print_exc()



    # 4.2 Nam conf x Nam realiz

def nc_nr(prestadora_sigla, mes):

  nome_tabela_iqa = prestadora_sigla + "_IQA"
  nome_tabela_iqe = prestadora_sigla + "_IQE"

    
  # IQA

  try:

      response = supabase.table(nome_tabela_iqa).select(
          "mes, ano, nam_conf, nam_realiz"
      ).execute()

      df = pd.DataFrame(response.data)

      if df.empty:
          print("❌ Nenhum dado encontrado!")
          exit()

      df['ano'] = pd.to_numeric(
          df['ano'],
          errors='coerce'
      )

      df = df.dropna(subset=['ano'])

      df['ano'] = df['ano'].astype(int)

      mapa_meses = {
          'january': 1,
          'february': 2,
          'march': 3,
          'april': 4,
          'may': 5,
          'june': 6,
          'july': 7,
          'august': 8,
          'september': 9,
          'october': 10,
          'november': 11,
          'december': 12,

          'janeiro': 1,
          'fevereiro': 2,
          'março': 3,
          'abril': 4,
          'maio': 5,
          'junho': 6,
          'julho': 7,
          'agosto': 8,
          'setembro': 9,
          'outubro': 10,
          'novembro': 11,
          'dezembro': 12
      }

      df['mes_numero'] = (
          df['mes']
          .astype(str)
          .str.strip()
          .str.lower()
          .map(mapa_meses)
      )

      if df['mes_numero'].isna().any():

          meses_invalidos = df.loc[
              df['mes_numero'].isna(),
              'mes'
          ].unique()

          print(
              f"⚠️ Meses não reconhecidos: "
              f"{meses_invalidos}"
          )

      df = df.dropna(
          subset=['mes_numero']
      )

      df['mes_numero'] = df['mes_numero'].astype(int)

      ultimo_ano = df['ano'].max()

      print(
          f"Último ano encontrado: {ultimo_ano}"
      )

      df_ano = df[
          df['ano'] == ultimo_ano
      ].copy()

      if df_ano.empty:

          print(
              "❌ Nenhum dado encontrado "
              "para o último ano!"
          )

          exit()

      ultimo_mes = df_ano[
          'mes_numero'
      ].max()

      print(
          f"Último mês disponível: {ultimo_mes}"
      )

      df_ano['data'] = pd.to_datetime(
          dict(
              year=df_ano['ano'],
              month=df_ano['mes_numero'],
              day=1
          )
      )

      ultima_data = df_ano['data'].max()

      data_inicial = (
          ultima_data -
          pd.DateOffset(months=5)
      )

      df_filtrado = df_ano[
          (df_ano['data'] >= data_inicial) &
          (df_ano['data'] <= ultima_data)
      ].copy()

      df_filtrado = df_filtrado.sort_values(
          'data'
      )

      quantidade = len(df_filtrado)

      print(
          f"Meses encontrados: {quantidade}/6"
      )

      print(
          "Detalhes:",
          df_filtrado['mes'].tolist()
      )

      if quantidade == 0:

          print(
              "❌ Nenhum mês encontrado "
              "para o período!"
          )

          exit()

      nomes_meses = {
          1: 'Jan',
          2: 'Fev',
          3: 'Mar',
          4: 'Abr',
          5: 'Mai',
          6: 'Jun',
          7: 'Jul',
          8: 'Ago',
          9: 'Set',
          10: 'Out',
          11: 'Nov',
          12: 'Dez'
      }

      meses_labels = [
          nomes_meses[m]
          for m in df_filtrado['mes_numero']
      ]

      x = np.arange(
          len(meses_labels)
      )

      nam_conf = pd.to_numeric(
          df_filtrado['nam_conf'],
          errors='coerce'
      ).fillna(0)

      nam_realiz = pd.to_numeric(
          df_filtrado['nam_realiz'],
          errors='coerce'
      ).fillna(0)

      fig, ax = plt.subplots(
          figsize=(12, 4.5)
      )

      cor_conf = '#0878E8'
      cor_realiz = '#9A9D17'

      linha1 = ax.plot(
          x,
          nam_conf,
          marker='o',
          markersize=6,
          linewidth=2,
          label='Nam conf',
          color=cor_conf
      )

      linha2 = ax.plot(
          x,
          nam_realiz,
          marker='o',
          markersize=6,
          linewidth=2,
          label='Nam realiz',
          color=cor_realiz
      )

      maior_valor = max(
          nam_conf.max(),
          nam_realiz.max()
      )

      ax.set_ylim(
          0,
          maior_valor * 1.15
          if maior_valor > 0
          else 1
      )

      ax.set_xticks(x)

      ax.set_xticklabels(
          meses_labels,
          fontsize=10
      )

      ax.legend(
          loc='upper left',
          bbox_to_anchor=(0, 1.12),
          ncol=2,
          frameon=False,
          fontsize=10
      )

      ax.grid(
          axis='y',
          linestyle='-',
          linewidth=0.7,
          alpha=0.35
      )

      ax.set_axisbelow(True)

      ax.spines['top'].set_visible(False)
      ax.spines['right'].set_visible(False)

      ax.set_xlabel('')
      ax.set_ylabel('')

      for i, valor in enumerate(nam_conf):

          ax.text(
              x[i],
              valor - (maior_valor * 0.035),
              f'{valor:.0f}'.replace('.', ','),
              ha='center',
              va='top',
              fontsize=10
          )

      for i, valor in enumerate(nam_realiz):

          ax.text(
              x[i],
              valor + (maior_valor * 0.035),
              f'{valor:.0f}'.replace('.', ','),
              ha='center',
              va='bottom',
              fontsize=10
          )

      ax.tick_params(
          axis='both',
          length=0
      )

      plt.tight_layout()

      pasta = "graficos"

      if not os.path.exists(pasta):
          os.makedirs(pasta)

      nome_arquivo = "nc_nr_IQAS.png"

      plt.savefig(
          nome_arquivo,
          dpi=300,
          bbox_inches='tight',
          format='png'
      )

      print(
          f"✅ Gráfico salvo: {nome_arquivo}"
      )

      plt.show()

  except Exception as e:

      print(f"❌ Erro: {e}")

      import traceback
      traceback.print_exc()


    
    # IQE

  if mes == "November" or mes == "February" or mes == "May" or mes == "August":

    try:

        response = supabase.table(nome_tabela_iqe).select(
            "trimestre, ano_contratual, nam_conf, nam_realiz"
        ).execute()

        df = pd.DataFrame(response.data)

        if df.empty:
            print("❌ Nenhum dado encontrado!")
            exit()

        ultimo_ano = df['ano_contratual'].max()

        print(
            f"Consultando último ano contratual: "
            f"{ultimo_ano}"
        )

        df_filtrado = df[
            df['ano_contratual'] == ultimo_ano
        ].copy()

        ordem_trimestres = [
            '1° trimestre',
            '2° trimestre',
            '3° trimestre',
            '4° trimestre'
        ]

        df_filtrado['trimestre'] = pd.Categorical(
            df_filtrado['trimestre'],
            categories=ordem_trimestres,
            ordered=True
        )

        df_filtrado = df_filtrado.sort_values(
            'trimestre'
        )

        df_filtrado = df_filtrado.dropna(
            subset=['trimestre']
        )

        quantidade = len(df_filtrado)

        print(
            f"Trimestres encontrados: {quantidade}/4"
        )

        print(
            f"Detalhes: "
            f"{df_filtrado['trimestre'].astype(str).tolist()}\n"
        )

        if quantidade == 0:

            print(
                "❌ Nenhum trimestre encontrado "
                "para este ano contratual!"
            )

            exit()

        trimestres_labels = [
            t.replace(
                '° trimestre',
                '° trim'
            )
            for t in df_filtrado[
                'trimestre'
            ].astype(str)
        ]

        x = np.arange(
            len(trimestres_labels)
        )

        nam_conf = pd.to_numeric(
            df_filtrado['nam_conf'],
            errors='coerce'
        ).fillna(0)

        nam_realiz = pd.to_numeric(
            df_filtrado['nam_realiz'],
            errors='coerce'
        ).fillna(0)

        fig, ax = plt.subplots(
            figsize=(12, 4.5)
        )

        cor_conf = '#0878E8'
        cor_realiz = '#9A9D17'

        linha1 = ax.plot(
            x,
            nam_conf,
            marker='o',
            markersize=6,
            linewidth=2,
            label='Nam conf',
            color=cor_conf
        )

        linha2 = ax.plot(
            x,
            nam_realiz,
            marker='o',
            markersize=6,
            linewidth=2,
            label='Nam realiz',
            color=cor_realiz
        )

        maior_valor = max(
            nam_conf.max(),
            nam_realiz.max()
        )

        ax.set_ylim(
            0,
            maior_valor * 1.15
            if maior_valor > 0
            else 1
        )

        ax.set_xticks(x)

        ax.set_xticklabels(
            trimestres_labels,
            fontsize=10
        )

        ax.legend(
            loc='upper left',
            bbox_to_anchor=(0, 1.12),
            ncol=2,
            frameon=False,
            fontsize=10
        )

        ax.grid(
            axis='y',
            linestyle='-',
            linewidth=0.7,
            alpha=0.35
        )

        ax.set_axisbelow(True)

        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

        ax.set_xlabel('')
        ax.set_ylabel('')

        for i, valor in enumerate(nam_conf):

            ax.text(
                x[i],
                valor - (maior_valor * 0.035),
                f'{valor:.0f}'.replace('.', ','),
                ha='center',
                va='top',
                fontsize=10
            )

        for i, valor in enumerate(nam_realiz):

            ax.text(
                x[i],
                valor + (maior_valor * 0.035),
                f'{valor:.0f}'.replace('.', ','),
                ha='center',
                va='bottom',
                fontsize=10
            )

        ax.tick_params(
            axis='both',
            length=0
        )

        plt.tight_layout()

        pasta = "graficos"

        if not os.path.exists(pasta):
            os.makedirs(pasta)

        nome_arquivo = "nc_nr_IQES.png"

        plt.savefig(
            nome_arquivo,
            dpi=300,
            bbox_inches='tight',
            format='png'
        )

        print(
            f"✅ Gráfico salvo: {nome_arquivo}"
        )

        if quantidade < 4:

            print(
                f"⚠️ O último ano contratual possui apenas "
                f"{quantidade}/4 trimestres. "
                f"O gráfico mostra somente os "
                f"trimestres existentes."
            )

        plt.show()

    except Exception as e:

        print(f"❌ Erro: {e}")

        import traceback
        traceback.print_exc()



##############################################################################################################################################################################################################################

# ELABORAÇÃO DE RELATÓRIO AUTOMATIZADO - IQA

def rel_iqa():

  # DEFININDO E AJUSTANDO AS TABELAS BASES
    
  def base(PLAN_MUN_TIPO, PLAN_REALIZADAS, PLAN_POND, MES_ANO):

    arquivo_saida   = "PLANO_TRATADO.xlsx"
    coluna_valor    = "PLANO"
    coluna_desc     = "DESCONSIDERAÇÕES - ARSAL"
    coluna_id_pond  = "ID_POND"

    BASE_MUN_TIPO = PLAN_MUN_TIPO

    BASE_MUN_TIPO = BASE_MUN_TIPO.dropna(subset=['PLANO'])

    mes, ano = MES_ANO.split("-")
    ano = ano[:4]

    meses_en = {
      "JANEIRO": "January", "FEVEREIRO": "February", "MARÇO": "March",
      "ABRIL": "April", "MAIO": "May", "JUNHO": "June",
      "JULHO": "July", "AGOSTO": "August", "SETEMBRO": "September",
      "OUTUBRO": "October", "NOVEMBRO": "November", "DEZEMBRO": "December"
    }

    mes = meses_en[mes]

    BASE_MUN_TIPO["Parâmetros"] = BASE_MUN_TIPO["Parâmetros"].replace({
      "Cor": "Cor Aparente", "Residual de Cloro Livre": "Cloro residual livre",
      "Residual de cloro livre": "Cloro residual livre", "Cor aparente": "Cor Aparente",
      "Cloro": "Cloro residual livre", "Coliformes Totais": "Coliformes totais",
      "Escherichia coli": "Escherichia Coli", "Cloro Residual Livre": "Cloro residual livre", "Cloro Residual livre": "Cloro residual livre", "Cloro residual Livre": "Cloro residual livre",
      "Cloro Livre": "Cloro residual livre", "Cloro livre": "Cloro residual livre", "PH": "pH", "ph": "pH", "Ph": "pH"
    })

    parametros_principais = [
      "Turbidez", "Cor Aparente", "pH",
      "Cloro residual livre", "Coliformes totais", "Escherichia Coli"
    ]

    BASE_MUN_TIPO["Parâmetros"] = BASE_MUN_TIPO["Parâmetros"].apply(
      lambda x: x if x in parametros_principais else "Demais Parâmetros"
    )

    BASE_MUN_TIPO[coluna_id_pond] = BASE_MUN_TIPO[coluna_id_pond].fillna("").astype(str).str.strip()

    colunas_chave = ["Cidade", "Parâmetros", coluna_id_pond]

    BASE_MUN_TIPO_agrupado = BASE_MUN_TIPO.groupby(
      colunas_chave, as_index=False
    ).agg(
      **{
        coluna_valor: (coluna_valor, "sum"),
        coluna_desc:  (coluna_desc,  "sum"),
      }
    )

    cols = [c for c in BASE_MUN_TIPO_agrupado.columns if c != coluna_id_pond] + [coluna_id_pond]
    BASE_MUN_TIPO_agrupado = BASE_MUN_TIPO_agrupado[cols]

    BASE_MUN_TIPO_agrupado.insert(0, "MÊS", mes)
    BASE_MUN_TIPO_agrupado.insert(1, "ANO", ano)

    municipios = BASE_MUN_TIPO_agrupado["Cidade"].unique()
    combinacoes_principais = pd.DataFrame(
      list(itertools.product([mes], [ano], municipios, parametros_principais, [""])),
      columns=["MÊS", "ANO", "Cidade", "Parâmetros", coluna_id_pond]
    )

    BASE_MUN_TIPO_principais = pd.merge(
      combinacoes_principais, BASE_MUN_TIPO_agrupado,
      on=["MÊS", "ANO", "Cidade", "Parâmetros", coluna_id_pond], how="left"
    )
    BASE_MUN_TIPO_principais[coluna_valor] = BASE_MUN_TIPO_principais[coluna_valor].fillna(0)
    BASE_MUN_TIPO_principais[coluna_desc]  = BASE_MUN_TIPO_principais[coluna_desc].fillna(0)

    BASE_COM_ID = BASE_MUN_TIPO_agrupado[
      BASE_MUN_TIPO_agrupado[coluna_id_pond] != ""
    ]

    BASE_MUN_TIPO_demais = BASE_MUN_TIPO_agrupado[
      (BASE_MUN_TIPO_agrupado["Parâmetros"] == "Demais Parâmetros") &
      (BASE_MUN_TIPO_agrupado[coluna_id_pond] == "")
    ]

    BASE_MUN_TIPO_completo = pd.concat(
      [BASE_MUN_TIPO_principais, BASE_COM_ID, BASE_MUN_TIPO_demais],
      ignore_index=True
    )

    cols = [c for c in BASE_MUN_TIPO_completo.columns if c != coluna_id_pond] + [coluna_id_pond]
    BASE_MUN_TIPO_completo = BASE_MUN_TIPO_completo[cols]

    BASE_MUN_TIPO_completo.to_excel(arquivo_saida, index=False)

    BASE_MUN = BASE_MUN_TIPO_agrupado.groupby("Cidade", as_index=False)[coluna_valor].sum()
    BASE_MUN = BASE_MUN.rename(columns={"Cidade": "MUNICÍPIO", "PLANO": "PLANO DE AMOSTRAGEM"})
    BASE_MUN["MUNICÍPIO"] = BASE_MUN["MUNICÍPIO"].str.upper().apply(unidecode)

    BASE_TIPO = BASE_MUN_TIPO_agrupado.groupby("Parâmetros", as_index=False)[coluna_valor].sum()
    BASE_TIPO = BASE_TIPO.rename(columns={"Parâmetros": "ANÁLISE", "PLANO": "PLANO DE AMOSTRAGEM"})
    BASE_TIPO["ANÁLISE"] = BASE_TIPO["ANÁLISE"].str.upper().apply(unidecode)

    BASE_REALIZADAS = PLAN_REALIZADAS

    BASE_REALIZADAS["ANÁLISE"] = BASE_REALIZADAS["ANÁLISE"].replace({
      "Cor": "Cor Aparente", "Residual de Cloro Livre": "Cloro residual livre",
      "Residual de cloro livre": "Cloro residual livre", "Cor aparente": "Cor Aparente",
      "Cloro": "Cloro residual livre", "Coliformes Totais": "Coliformes totais",
      "Escherichia coli": "Escherichia Coli", "Cloro Residual Livre": "Cloro residual livre", "Cloro Residual livre": "Cloro residual livre", "Cloro residual Livre": "Cloro residual livre",
      "Cloro Livre": "Cloro residual livre", "Cloro livre": "Cloro residual livre", "PH": "pH", "ph": "pH", "Ph": "pH"})

    BASE_REALIZADAS["CIDADE"] = BASE_REALIZADAS["CIDADE"].str.upper().apply(unidecode)
    BASE_REALIZADAS["ANÁLISE"] = BASE_REALIZADAS["ANÁLISE"].str.upper().apply(unidecode)

    BASE_POND = PLAN_POND

    return (BASE_MUN, BASE_TIPO, BASE_REALIZADAS, BASE_POND, BASE_MUN_TIPO_completo)



  # COLETANDO AS INFORMAÇÕES GERAIS
    


  def infos_gerais (TAB_INFO):

    mes_ing = TAB_INFO["INFORMAÇÃO"][3]
    ano = TAB_INFO["INFORMAÇÃO"][1]

    meses = {
    'January': 'JANEIRO',
    'February': 'FEVEREIRO',
    'March': 'MARÇO',
    'April': 'ABRIL',
    'May': 'MAIO',
    'June': 'JUNHO',
    'July': 'JULHO',
    'August': 'AGOSTO',
    'September': 'SETEMBRO',
    'October': 'OUTUBRO',
    'November': 'NOVEMBRO',
    'December': 'DEZEMBRO'
    }

    mes = meses[mes_ing]
    MES_ANO = mes + '-' + str(ano)
    mes_min = mes.lower()

    trimestre = TAB_INFO["INFORMAÇÃO"][4]
    ANO_CONTRAT = TAB_INFO["INFORMAÇÃO"][2]

    TRIM_ANO = (trimestre + ' - ANO ' + str(ANO_CONTRAT)).upper()
    TRIMMIN_ANO = TRIM_ANO.lower()

    MESMIN_ANO = mes_min + ' de ' + str(ano)

    PRESTADORA_SIGLA = TAB_INFO["INFORMAÇÃO"][0]

    if PRESTADORA_SIGLA == "BRK":
      PRESTADORA = "BRK Ambiental"
    elif PRESTADORA_SIGLA == "ADS":
      PRESTADORA = "Conasa Águas do Sertão"
    else:
      PRESTADORA = "Verde Ambiental Alagoas"

    META = TAB_INFO ["INFORMAÇÃO"][11]

    META = float(META)

    META = "100%" if META >= 1 else f"{META*100:.1f}%".replace('.', ',').rstrip('0').rstrip(',')

    NCOF_P = TAB_INFO ["INFORMAÇÃO"][5]
    NREA_P = TAB_INFO ["INFORMAÇÃO"][6]
    IQA_P = TAB_INFO ["INFORMAÇÃO"][7]

    IQA_P = float(IQA_P)

    IQA_P = "100%" if IQA_P >= 1 else f"{IQA_P*100:.1f}%".replace('.', ',').rstrip('0').rstrip(',')

    NCOF_VI = TAB_INFO ["INFORMAÇÃO"][8]
    NREA_VI = TAB_INFO ["INFORMAÇÃO"][9]
    IQA_VI = TAB_INFO ["INFORMAÇÃO"][10]

    IQA_VI = float(IQA_VI)

    IQA_VI = "100%" if IQA_VI >= 1 else f"{IQA_VI*100:.1f}%".replace('.', ',').rstrip('0').rstrip(',')

    FONTE_NREALIZ = TAB_INFO ["INFORMAÇÃO"][12]
    FONTE_PLANO = TAB_INFO ["INFORMAÇÃO"][13]
    FONT_NCONF = TAB_INFO ["INFORMAÇÃO"][14]
    ACREDITAÇÃO = TAB_INFO ["INFORMAÇÃO"][15]

    return (MES_ANO, ANO_CONTRAT, PRESTADORA, META, NCOF_P, NREA_P, IQA_P, NCOF_VI, NREA_VI, IQA_VI, FONTE_NREALIZ, FONTE_PLANO, FONT_NCONF, ACREDITAÇÃO, MESMIN_ANO)



  # DEFINIÇÃO DO CONJUNTO DE VARIÁVEIS A SEREM INSERIDAS NO RELATÓRIO


  def var_doc (MES_ANO, ANO_CONTRAT, PRESTADORA, META, NCOF_P, NREA_P, IQA_P, NCOF_VI, NREA_VI, IQA_VI, FONTE_NREALIZ, FONTE_PLANO, FONT_NCONF, ACREDITAÇÃO, TAB_MUN_FINAL, QUADRO_POND_NCONF, QUADRO_POND_REALIZ, TAB_MUN_ARSAL, TAB_MUN_VI, TAB_TIPO_FINAL, TAB_TIPO_RESUMIDA_VI, MESMIN_ANO, total_conf_vi_naoconf_arsal, total_naoconf_vi_conf_arsal, desc_plano, TAB_MUN_NCONF, TAB_MUN_NCONF_VI):

    NREA_ARSAL = (TAB_MUN_FINAL["NAM REALIZ"].iloc[-1]).astype(int)
    NREA_EXPURGOS = (TAB_MUN_FINAL["EXPURGOS TOTAIS"].iloc[-1]).astype(int)
    NCONF_ARSAL = (TAB_MUN_FINAL["NAM CONF"].iloc[-1]).astype(int)
    NCONF_EXPURGOS = (TAB_MUN_FINAL["EXPURGOS CONFORMES"].iloc[-1]).astype(int)
    IQA_ARSAL = (NCONF_ARSAL/NREA_ARSAL)*100
    IQA_ARSAL_TXT = "100%" if IQA_ARSAL == 100 else f"{IQA_ARSAL:.1f}%".replace('.', ',')

    if (IQA_VI).strip() == (IQA_ARSAL_TXT).strip():
      IS_SATISF = "satisfatório"

    if (IQA_VI).strip() != (IQA_ARSAL_TXT).strip():
      IS_SATISF = "divergente"

    if PRESTADORA == "BRK Ambiental":
      legis = '5'

    if PRESTADORA != "BRK Ambiental":
      legis = '7'

    def validar(valor):
      return valor if valor is not None else ""


    def formata_milhar(x):
      try:
          return f"{int(float(x)):,}".replace(",", ".")
      except:
          return "-"


    var_dic = {
              '$MES_ANO': str(MES_ANO),
              '$ANO_CONTRAT': str(ANO_CONTRAT),
              '$MESMIN_ANO': str(MESMIN_ANO),
              '$PRESTADORA': str(PRESTADORA),
              '$META':str(META),
              '$NCOF_P': formata_milhar(NCOF_P),
              '$NREA_P': formata_milhar(NREA_P),
              '$IQA_P': str(IQA_P),
              '$NCOF_VI': formata_milhar(NCOF_VI),
              '$NREA_VI': formata_milhar(NREA_VI),
              '$IQA_VI': str(IQA_VI),
              '$FONTE_NREALIZ': str(FONTE_NREALIZ),
              '$FONTE_PLANO': str(FONTE_PLANO),
              '$FONTE_NCONF': str(FONT_NCONF),
              '$ACREDITAÇÃO': str (ACREDITAÇÃO),

              '$QUADRO_NREA_OBS': QUADRO_POND_REALIZ,
              '$QUADRO_NCONF_OBS': QUADRO_POND_NCONF,

              '$QUADRO_NREA_VI': TAB_MUN_VI,
              '$QUADRO_NREA_ARSAL': TAB_MUN_ARSAL,

              '$QUADRO_NCONF_VI': TAB_TIPO_RESUMIDA_VI,
              '$QUADRO_NCONF_ARSAL': TAB_TIPO_FINAL,

              '$QUADRO_NCONFMUN_VI': TAB_MUN_NCONF_VI,
              '$QUADRO_NCONFMUN_ARSAL': TAB_MUN_NCONF,


              '$NREA_ARSAL': formata_milhar(NREA_ARSAL),
              '$NREA_EXPURGOS': str (NREA_EXPURGOS),
              '$NCOF_ARSAL': formata_milhar(NCONF_ARSAL),
              '$NCONF_EXPURGOS': str (NCONF_EXPURGOS),
              '$IS_SATISFATORIO': IS_SATISF,
              '$IQA_ARSAL': IQA_ARSAL_TXT,
              '$LEGIS' : legis,
              '$CONF_NCONF': str(total_conf_vi_naoconf_arsal),
              '$NCONF_CONF' : str(total_naoconf_vi_conf_arsal),
              '$NREA__PLAN' : str(desc_plano)
              }

    return var_dic



  # CONVERSÃO DE DATAFRAME PARA TABELA EDITÁVEL NO RELATÓRIO

  def df_para_tabela_word(doc, df, cor_cabecalho="BFBFBF", cor_par="FFFFFF", cor_impar="FFFFFF"):
      
      from docx.oxml.ns import qn
      from docx.oxml import OxmlElement
      from docx.shared import Pt, RGBColor, Cm
      from docx.enum.text import WD_ALIGN_PARAGRAPH
      from docx.enum.table import WD_TABLE_ALIGNMENT

      def set_bg(cell, cor):
          tcPr = cell._tc.get_or_add_tcPr()
          shd = OxmlElement("w:shd")
          shd.set(qn("w:val"), "clear")
          shd.set(qn("w:color"), "auto")
          shd.set(qn("w:fill"), cor)
          tcPr.append(shd)

      def set_borders(cell, cor="000000", tamanho="4"):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for lado in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{lado}")
            border.set(qn("w:val"),   "single")
            border.set(qn("w:sz"),    tamanho)
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), cor)
            tcBorders.append(border)
        tcPr.append(tcBorders)

      def set_col_widths(tabela, larguras_cm):

          tbl = tabela._tbl
          tblGrid = tbl.find(qn("w:tblGrid"))
          if tblGrid is None:
              tblGrid = OxmlElement("w:tblGrid")
              tbl.insert(0, tblGrid)
          else:
              for col in tblGrid.findall(qn("w:gridCol")):
                  tblGrid.remove(col)

          for largura in larguras_cm:
              gridCol = OxmlElement("w:gridCol")
              gridCol.set(qn("w:w"), str(int(largura / 635)))
              tblGrid.append(gridCol)

          for row in tabela.rows:
              for i, cell in enumerate(row.cells):
                  tc = cell._tc
                  tcPr = tc.get_or_add_tcPr()
                  tcW = tcPr.find(qn("w:tcW"))
                  if tcW is None:
                      tcW = OxmlElement("w:tcW")
                      tcPr.append(tcW)
                  tcW.set(qn("w:w"),    str(int(larguras_cm[i] / 635)))
                  tcW.set(qn("w:type"), "dxa")

      tabela = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
      tabela.style = "Normal Table"
      tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

      num_cols = len(df.columns)
      larguras = []

      MUNICIPIO_ANALISE = Cm(5.5)
      EX_LARGA          = Cm(7.0)
      LARGA             = Cm(3.2)
      MEDIA             = Cm(2.5)
      ESTREITA          = Cm(1.6)
      CONVERSAO         = Cm(3.0)

      for c in range(num_cols):
          col_name = df.columns[c]

          if c == 0:

              if col_name in ("MUNICÍPIO", "ANÁLISE"):
                  larguras.append(MUNICIPIO_ANALISE)
              else:
                  larguras.append(ESTREITA)
          else:

              if col_name in ("CONF VI → NÃO CONF ARSAL", "NÃO CONF VI → CONF ARSAL",
                              "CONVERSÃO: CONFORME → NÃO CONFORME", "CONVERSÃO: NÃO CONFORME → CONFORME"):
                  larguras.append(CONVERSAO)

              elif col_name in ("PONDERAÇÕES (PRESTADORA)", "PONDERAÇÕES (VI)", "PONDERAÇÕES (ARSAL)"):
                  larguras.append(EX_LARGA)

              elif col_name in ("PLANO DE AMOSTRAGEM", "PLANO DE AMOSTRAGEM AJUSTADO",
                                "PLANO DE AMOSTRAGEM AJUSTADO - ARSAL", "DESCONSIDERAÇÕES DO PLANO",
                                "ANALISES REALIZADAS", "ANALISES CONFORMES", "ANALISES NÃO CONFORMES",
                                "EXPURGOS TOTAIS", "EXPURGOS CONFORMES", "EXPURGOS NAO CONFORMES",
                                "NAM REALIZ", "NAM REALIZ (VI)", "NAM CONF", "NAM CONF (VI)"):
                  larguras.append(MEDIA)

              else:
                  larguras.append(ESTREITA)

      set_col_widths(tabela, larguras)

      for c, col in enumerate(df.columns):
          cell = tabela.rows[0].cells[c]

          set_bg(cell, cor_cabecalho)
          set_borders(cell)

          cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

          p = cell.paragraphs[0]
          p.clear()

          p.paragraph_format.space_after = Pt(0)   # ← AQUI
          p.paragraph_format.space_before = Pt(0)  # ← AQUI

          run = p.add_run(str(col))
          run.bold = True
          run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
          run.font.name = "Times New Roman"
          run.font.size = Pt(9.0)
          p.alignment = WD_ALIGN_PARAGRAPH.CENTER

      for r, (_, linha) in enumerate(df.iterrows()):
          cor = cor_par if r % 2 == 0 else cor_impar
          is_ultima_linha = (df.columns[-1] in ("NAM REALIZ", "NAM REALIZ (VI)", "NAM CONF", "NAM CONF (VI)")) and (r == len(df) - 1)
          for c, valor in enumerate(linha):
              cell = tabela.rows[r + 1].cells[c]
              set_bg(cell, cor)
              set_borders(cell)

              cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

              p = cell.paragraphs[0]
              p.clear()

              p.paragraph_format.space_after = Pt(0)   # ← AQUI
              p.paragraph_format.space_before = Pt(0)  # ← AQUI

              run = p.add_run(str(valor) if valor is not None else "")

              run.bold = is_ultima_linha

              run.font.name = "Times New Roman"
              run.font.size = Pt(9.0)
              p.alignment = WD_ALIGN_PARAGRAPH.CENTER

      doc.element.body.remove(tabela._tbl)
      return tabela._tbl



  # ALIMENTANDO INFORMAÇÕES AO MODELO DE RELATÓRIO

  def substituir_var(documento, var_dic):
      doc = Document(documento)

      def process_paragraph(paragraph):
          full_text = "".join(run.text for run in paragraph.runs)

          if any(key in full_text for key in var_dic.keys()):
              for key, value in var_dic.items():
                  if key in full_text:
                      if isinstance(value, pd.DataFrame):
                        
                          tbl = df_para_tabela_word(doc, value)
                          paragraph._element.addprevious(tbl)

                          for run in paragraph.runs:
                              run.text = ""

                          return
                      else:
                          full_text = full_text.replace(key, value)

              for run in paragraph.runs:
                  run.text = ""
              paragraph.runs[0].text = full_text

      for p in doc.paragraphs:
          process_paragraph(p)

      for table in doc.tables:
          for row in table.rows:
              for cell in row.cells:
                  for p in cell.paragraphs:
                      process_paragraph(p)

      def replace_placeholder_with_image(doc, placeholder, image_path, width_inches):
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:

                paragraph.clear()
    
                paragraph.add_run().add_picture(image_path, width=Inches(width_inches))

      placeholders = [

        ('$GRAF_COMPARATIVO_IQAS', 'IQAS.png', 6.5),

        ('$GRAF_CONF_REALIZ_TEMPO', 'nc_nr_IQAS.png', 6.5),
        ('$GRAF_COMPARATIVO_IQAS', 'IQAS.png', 5.5),

        ('$GRAF_PLANO_REALIZ_ANALISE', 'grafico_plan_realiz_tipo.png', 6.5),
        ('$GRAF_PLANO_REALIZ_MUN', 'grafico_plan_realiz_mun.png', 9.0),#####
        ('$GRAF_CONF_REALIZ_ANALISE', 'grafico_conf_realiz_tipo.png', 6.5),
        ('$GRAF_CONF_REALIZ_MUN', 'grafico_conf_realiz_mun.png', 9.5),
        ('$GRAF_IQAS_MUN', 'grafico_iqa_mun.png', 6.5)
      ]

      for ph, img, width in placeholders:
        replace_placeholder_with_image(doc, ph, img, width)

      return doc



# ANÁLISE POR MUNICÍPIO

  def analise_mun(TAB_MUN, TAB_REALIZADAS, BASE_MUN_TIPO_completo):

    BASE_MUN_TIPO_completo["Cidade"] = BASE_MUN_TIPO_completo["Cidade"].str.upper().apply(unidecode)

    DESC_MUN = (
      BASE_MUN_TIPO_completo
      .groupby("Cidade", as_index=False)["DESCONSIDERAÇÕES - ARSAL"]
      .sum()
      .rename(columns={
        "Cidade": "MUNICÍPIO",
        "DESCONSIDERAÇÕES - ARSAL": "DESCONSIDERAÇÕES DO PLANO"
      })
    )

    TAB_MUN = TAB_MUN.merge(DESC_MUN, on="MUNICÍPIO", how="left")
    TAB_MUN["DESCONSIDERAÇÕES DO PLANO"] = TAB_MUN["DESCONSIDERAÇÕES DO PLANO"].fillna(0).astype(int)
    TAB_MUN["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"] = (TAB_MUN["PLANO DE AMOSTRAGEM"] - TAB_MUN["DESCONSIDERAÇÕES DO PLANO"]).astype(int)

    TAB_REALIZADAS_ACEITO          = TAB_REALIZADAS[TAB_REALIZADAS["ACEITO?"] == "Validado"]
    TAB_REALIZADAS_ACEITO_CONF     = TAB_REALIZADAS_ACEITO[TAB_REALIZADAS["CONFORMIDADE - ARSAL"] == "Conforme"]
    TAB_REALIZADAS_ACEITO_CONF_VI     = TAB_REALIZADAS_ACEITO[TAB_REALIZADAS["CONFORMIDADE - VI"] == "Conforme"]
    TAB_REALIZADAS_ACEITO_NAOCONF  = TAB_REALIZADAS_ACEITO[(TAB_REALIZADAS["CONFORMIDADE - ARSAL"] != "Conforme") & (TAB_REALIZADAS["CONFORMIDADE - ARSAL"].notna())]
    TAB_REALIZADAS_ACEITO_NAOCONF_VI  = TAB_REALIZADAS_ACEITO[(TAB_REALIZADAS["CONFORMIDADE - VI"] != "Conforme") & (TAB_REALIZADAS["CONFORMIDADE - VI"].notna())]
    TAB_REALIZADAS_ACEITO_CONF_EXP    = TAB_REALIZADAS_ACEITO_CONF[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_CONF_EXP_VI    = TAB_REALIZADAS_ACEITO_CONF_VI[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_NAOCONF_EXP = TAB_REALIZADAS_ACEITO_NAOCONF[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI = TAB_REALIZADAS_ACEITO_NAOCONF_VI[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]

    CONF_VI_NAOCONF_ARSAL = TAB_REALIZADAS_ACEITO[
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - VI"] == "Conforme") &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - ARSAL"] != "Conforme") &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - ARSAL"].notna()) &
        (TAB_REALIZADAS_ACEITO["EXPURGOS - ARSAL"] != "EXPURGAR")
    ]

    NAOCONF_VI_CONF_ARSAL = TAB_REALIZADAS_ACEITO[
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - VI"] != "Conforme") &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - VI"].notna()) &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - ARSAL"] == "Conforme") &
        (TAB_REALIZADAS_ACEITO["EXPURGOS - ARSAL"] != "EXPURGAR")
    ]

    contagem_conf_vi_naoconf_arsal = (
        CONF_VI_NAOCONF_ARSAL.groupby("CIDADE").size()
        .reset_index(name="CONF VI → NÃO CONF ARSAL")
    )

    contagem_naoconf_vi_conf_arsal = (
        NAOCONF_VI_CONF_ARSAL.groupby("CIDADE").size()
        .reset_index(name="NÃO CONF VI → CONF ARSAL")
    )

    TAB_MUN = TAB_MUN.merge(contagem_conf_vi_naoconf_arsal, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN = TAB_MUN.drop(columns='CIDADE')
    TAB_MUN["CONF VI → NÃO CONF ARSAL"] = TAB_MUN["CONF VI → NÃO CONF ARSAL"].fillna(0).astype(int)

    TAB_MUN = TAB_MUN.merge(contagem_naoconf_vi_conf_arsal, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN = TAB_MUN.drop(columns='CIDADE')
    TAB_MUN["NÃO CONF VI → CONF ARSAL"] = TAB_MUN["NÃO CONF VI → CONF ARSAL"].fillna(0).astype(int)

    contagem_tot = TAB_REALIZADAS_ACEITO.groupby('CIDADE').size().reset_index(name='ANALISES REALIZADAS')
    TAB_MUN = TAB_MUN.merge(contagem_tot, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN = TAB_MUN.drop(columns='CIDADE')

    contagem_conf = TAB_REALIZADAS_ACEITO_CONF_VI.groupby('CIDADE').size().reset_index(name='ANALISES CONFORMES')
    TAB_MUN = TAB_MUN.merge(contagem_conf, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN = TAB_MUN.drop(columns='CIDADE')

    contagem_naoconf = TAB_REALIZADAS_ACEITO_NAOCONF_VI.groupby('CIDADE').size().reset_index(name='ANALISES NÃO CONFORMES')
    TAB_MUN = TAB_MUN.merge(contagem_naoconf, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN = TAB_MUN.drop(columns='CIDADE')
    TAB_MUN['ANALISES NÃO CONFORMES'] = TAB_MUN['ANALISES NÃO CONFORMES'].fillna(0).astype(int)

    contagem_conf_exp = TAB_REALIZADAS_ACEITO_CONF_EXP_VI.groupby('CIDADE').size().reset_index(name='EXPURGOS CONFORMES')
    TAB_MUN = TAB_MUN.merge(contagem_conf_exp, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN = TAB_MUN.drop(columns='CIDADE')
    TAB_MUN['EXPURGOS CONFORMES'] = TAB_MUN['EXPURGOS CONFORMES'].fillna(0).astype(int)

    contagem_naoconf_exp = TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI.groupby('CIDADE').size().reset_index(name='EXPURGOS NAO CONFORMES')
    TAB_MUN = TAB_MUN.merge(contagem_naoconf_exp, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN = TAB_MUN.drop(columns='CIDADE')
    TAB_MUN['EXPURGOS NAO CONFORMES'] = TAB_MUN['EXPURGOS NAO CONFORMES'].fillna(0).astype(int)

    TAB_MUN['EXPURGOS TOTAIS'] = TAB_MUN['EXPURGOS NAO CONFORMES'] + TAB_MUN['EXPURGOS CONFORMES']

    TAB_MUN["NAM REALIZ"] = (np.maximum(TAB_MUN['PLANO DE AMOSTRAGEM AJUSTADO - ARSAL'], TAB_MUN['ANALISES REALIZADAS']) - TAB_MUN['EXPURGOS TOTAIS']).astype(int)

    TAB_MUN["NAM CONF"] = (TAB_MUN['ANALISES CONFORMES'] - TAB_MUN['EXPURGOS CONFORMES'] + TAB_MUN['NÃO CONF VI → CONF ARSAL'] - TAB_MUN['CONF VI → NÃO CONF ARSAL']).astype(int)


    TAB_MUN["IQA (%)"] = (((TAB_MUN["NAM CONF"]) / (TAB_MUN["NAM REALIZ"])) * 100).round(1)

    TAB_MUN_FINAL = TAB_MUN[[
      "MUNICÍPIO", "PLANO DE AMOSTRAGEM", "DESCONSIDERAÇÕES DO PLANO",
      "PLANO DE AMOSTRAGEM AJUSTADO - ARSAL",
      "ANALISES REALIZADAS", "EXPURGOS TOTAIS", "NAM REALIZ",
      "ANALISES CONFORMES", "EXPURGOS CONFORMES", "CONF VI → NÃO CONF ARSAL", "NÃO CONF VI → CONF ARSAL", "NAM CONF", "IQA (%)"
    ]]
    TAB_MUN_FINAL.loc['Total'] = TAB_MUN_FINAL.sum()
    TAB_MUN_FINAL["MUNICÍPIO"].iat[-1] = ' '
    TAB_MUN_FINAL["IQA (%)"].iat[-1] = (((TAB_MUN_FINAL["NAM CONF"].iat[-1]) / (TAB_MUN_FINAL["NAM REALIZ"].iat[-1])) * 100).round(1)

    desc_plano = int(TAB_MUN_FINAL["DESCONSIDERAÇÕES DO PLANO"].iloc[-1])

    TAB_MUN_NCONF = TAB_MUN[["MUNICÍPIO", "ANALISES REALIZADAS", "ANALISES CONFORMES", "EXPURGOS CONFORMES", "CONF VI → NÃO CONF ARSAL", "NÃO CONF VI → CONF ARSAL", "NAM CONF"]]
    TAB_MUN_NCONF_VI = TAB_MUN[["MUNICÍPIO", "ANALISES REALIZADAS", "ANALISES CONFORMES"]]

    def add_total(df):
      total = df.select_dtypes(include='number').sum()
      total['MUNICÍPIO'] = 'TOTAL'
      return pd.concat([df, total.to_frame().T], ignore_index=True)

    TAB_MUN_NCONF = add_total(TAB_MUN_NCONF)
    TAB_MUN_NCONF_VI = add_total(TAB_MUN_NCONF_VI)

    TAB_MUN_RESUMIDA = TAB_MUN_FINAL[[
      "MUNICÍPIO", "PLANO DE AMOSTRAGEM", "DESCONSIDERAÇÕES DO PLANO",
      "PLANO DE AMOSTRAGEM AJUSTADO - ARSAL",
      "ANALISES REALIZADAS", "EXPURGOS TOTAIS", "NAM REALIZ"
    ]]
    TAB_MUN_RESUMIDA = TAB_MUN_RESUMIDA.rename(columns={"PLANO DE AMOSTRAGEM": "PLANO DE AMOSTRAGEM AJUSTADO"})
    TAB_MUN_RESUMIDA["PLANO DE AMOSTRAGEM AJUSTADO"] = TAB_MUN_RESUMIDA["PLANO DE AMOSTRAGEM AJUSTADO"].astype(int)
    TAB_MUN_RESUMIDA.loc[TAB_MUN_RESUMIDA.index[-1], "MUNICÍPIO"] = "TOTAL"

    TAB_MUN_RESUMIDA_VI = TAB_MUN_RESUMIDA[[
      "MUNICÍPIO", "PLANO DE AMOSTRAGEM AJUSTADO", "ANALISES REALIZADAS"
    ]]
    TAB_MUN_RESUMIDA_VI["NAM REALIZ (VI)"] = (np.maximum(
      TAB_MUN_RESUMIDA_VI["PLANO DE AMOSTRAGEM AJUSTADO"],
      TAB_MUN_RESUMIDA_VI['ANALISES REALIZADAS']
    )).astype(int)
    TAB_MUN_RESUMIDA_VI = TAB_MUN_RESUMIDA_VI[:-1]
    TAB_MUN_RESUMIDA_VI.loc['Total'] = TAB_MUN_RESUMIDA_VI.sum()
    TAB_MUN_RESUMIDA_VI.iloc[-1, TAB_MUN_RESUMIDA_VI.columns.get_loc("MUNICÍPIO")] = "TOTAL"

    return TAB_MUN_FINAL, TAB_MUN_RESUMIDA, TAB_MUN_RESUMIDA_VI, desc_plano, TAB_MUN_NCONF, TAB_MUN_NCONF_VI




# ANÁLISE POR PARÂMETRO

  def analise_tipo(TAB_TIPO, TAB_REALIZADAS, BASE_MUN_TIPO_completo):

    mapa_parametros = {
      "Turbidez":            "TURBIDEZ",
      "Cor Aparente":        "COR APARENTE",
      "pH":                  "PH",
      "Cloro residual livre": "CLORO RESIDUAL LIVRE",
      "Coliformes totais":   "COLIFORMES TOTAIS",
      "Escherichia Coli":    "ESCHERICHIA COLI",
      "Demais Parâmetros":   "DEMAIS PARAMETROS",
    }

    DESC_TIPO = (
      BASE_MUN_TIPO_completo
      .groupby("Parâmetros", as_index=False)["DESCONSIDERAÇÕES - ARSAL"]
      .sum()
    )
    DESC_TIPO["Parâmetros"] = DESC_TIPO["Parâmetros"].map(mapa_parametros)
    DESC_TIPO = DESC_TIPO.rename(columns={
      "Parâmetros":              "ANÁLISE",
      "DESCONSIDERAÇÕES - ARSAL": "DESCONSIDERAÇÕES DO PLANO"
    })

    TAB_TIPO = TAB_TIPO.merge(DESC_TIPO, on="ANÁLISE", how="left")
    TAB_TIPO["DESCONSIDERAÇÕES DO PLANO"] = TAB_TIPO["DESCONSIDERAÇÕES DO PLANO"].fillna(0).astype(int)
    TAB_TIPO["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"] = (TAB_TIPO["PLANO DE AMOSTRAGEM"] - TAB_TIPO["DESCONSIDERAÇÕES DO PLANO"]).astype(int)

    TAB_REALIZADAS = TAB_REALIZADAS.replace("CLORO", "CLORO RESIDUAL LIVRE")

    TAB_REALIZADAS['ANÁLISE'] = np.where(
      TAB_REALIZADAS['ANÁLISE'].isin(['TURBIDEZ', 'COR APARENTE', 'PH', 'COLIFORMES TOTAIS', 'CLORO RESIDUAL LIVRE', 'ESCHERICHIA COLI']),
      TAB_REALIZADAS['ANÁLISE'], 'DEMAIS PARAMETROS'
    )

    TAB_TIPO_VI = TAB_TIPO

    TAB_REALIZADAS_ACEITO = TAB_REALIZADAS[TAB_REALIZADAS["ACEITO?"] == "Validado"]

    CONF_VI_NAOCONF_ARSAL = TAB_REALIZADAS_ACEITO[
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - VI"] == "Conforme") &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - ARSAL"] != "Conforme") &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - ARSAL"].notna()) &
        (TAB_REALIZADAS_ACEITO["EXPURGOS - ARSAL"] != "EXPURGAR")
    ]

    NAOCONF_VI_CONF_ARSAL = TAB_REALIZADAS_ACEITO[
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - VI"] != "Conforme") &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - VI"].notna()) &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - ARSAL"] == "Conforme") &
        (TAB_REALIZADAS_ACEITO["EXPURGOS - ARSAL"] != "EXPURGAR")
    ]

    contagem_conf_vi_naoconf_arsal = (
        CONF_VI_NAOCONF_ARSAL.groupby("ANÁLISE").size()
        .reset_index(name="CONVERSÃO: CONFORME → NÃO CONFORME")
    )

    contagem_naoconf_vi_conf_arsal = (
        NAOCONF_VI_CONF_ARSAL.groupby("ANÁLISE").size()
        .reset_index(name="CONVERSÃO: NÃO CONFORME → CONFORME")
    )

    TAB_TIPO = TAB_TIPO.merge(contagem_conf_vi_naoconf_arsal, on="ANÁLISE", how="left")
    TAB_TIPO["CONVERSÃO: CONFORME → NÃO CONFORME"] = TAB_TIPO["CONVERSÃO: CONFORME → NÃO CONFORME"].fillna(0).astype(int)

    TAB_TIPO = TAB_TIPO.merge(contagem_naoconf_vi_conf_arsal, on="ANÁLISE", how="left")
    TAB_TIPO["CONVERSÃO: NÃO CONFORME → CONFORME"] = TAB_TIPO["CONVERSÃO: NÃO CONFORME → CONFORME"].fillna(0).astype(int)

    TAB_REALIZADAS_ACEITO_CONF = TAB_REALIZADAS_ACEITO[TAB_REALIZADAS["CONFORMIDADE - ARSAL"] == "Conforme"]
    TAB_REALIZADAS_ACEITO_CONF_VI = TAB_REALIZADAS_ACEITO[TAB_REALIZADAS["CONFORMIDADE - VI"] == "Conforme"]
    TAB_REALIZADAS_ACEITO_NAOCONF = TAB_REALIZADAS_ACEITO[(TAB_REALIZADAS["CONFORMIDADE - ARSAL"] != "Conforme") & (TAB_REALIZADAS["CONFORMIDADE - ARSAL"].notna())]
    TAB_REALIZADAS_ACEITO_NAOCONF_VI = TAB_REALIZADAS_ACEITO[(TAB_REALIZADAS["CONFORMIDADE - VI"] != "Conforme") & (TAB_REALIZADAS["CONFORMIDADE - VI"].notna())]
    TAB_REALIZADAS_ACEITO_CONF_EXP = TAB_REALIZADAS_ACEITO_CONF[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_CONF_EXP_VI = TAB_REALIZADAS_ACEITO_CONF_VI[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_NAOCONF_EXP = TAB_REALIZADAS_ACEITO_NAOCONF[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI = TAB_REALIZADAS_ACEITO_NAOCONF_VI[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]

    contagem_tot = TAB_REALIZADAS_ACEITO.groupby('ANÁLISE').size().reset_index(name='ANALISES REALIZADAS')
    TAB_TIPO = TAB_TIPO.merge(contagem_tot, on='ANÁLISE', how='left')
    TAB_TIPO['ANALISES REALIZADAS'] = TAB_TIPO['ANALISES REALIZADAS'].astype(int)

    contagem_conf = TAB_REALIZADAS_ACEITO_CONF_VI.groupby('ANÁLISE').size().reset_index(name='ANALISES CONFORMES')
    TAB_TIPO = TAB_TIPO.merge(contagem_conf, on='ANÁLISE', how='left')

    contagem_naoconf = TAB_REALIZADAS_ACEITO_NAOCONF_VI.groupby('ANÁLISE').size().reset_index(name='ANALISES NÃO CONFORMES')
    TAB_TIPO = TAB_TIPO.merge(contagem_naoconf, on='ANÁLISE', how='left')
    TAB_TIPO['ANALISES NÃO CONFORMES'] = TAB_TIPO['ANALISES NÃO CONFORMES'].fillna(0).astype(int)

    contagem_conf_exp = TAB_REALIZADAS_ACEITO_CONF_EXP_VI.groupby('ANÁLISE').size().reset_index(name='EXPURGOS CONFORMES')
    TAB_TIPO = TAB_TIPO.merge(contagem_conf_exp, on='ANÁLISE', how='left')
    TAB_TIPO['EXPURGOS CONFORMES'] = TAB_TIPO['EXPURGOS CONFORMES'].fillna(0).astype(int)

    contagem_naoconf_exp = TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI.groupby('ANÁLISE').size().reset_index(name='EXPURGOS NAO CONFORMES')
    TAB_TIPO = TAB_TIPO.merge(contagem_naoconf_exp, on='ANÁLISE', how='left')
    TAB_TIPO['EXPURGOS NAO CONFORMES'] = TAB_TIPO['EXPURGOS NAO CONFORMES'].fillna(0).astype(int)

    TAB_TIPO['EXPURGOS TOTAIS'] = TAB_TIPO['EXPURGOS NAO CONFORMES'] + TAB_TIPO['EXPURGOS CONFORMES']

    TAB_TIPO["NAM REALIZ"] = (np.maximum(TAB_TIPO['PLANO DE AMOSTRAGEM AJUSTADO - ARSAL'], TAB_TIPO['ANALISES REALIZADAS']) - TAB_TIPO['EXPURGOS TOTAIS']).astype(int)

    TAB_TIPO["NAM CONF"] = (TAB_TIPO['ANALISES CONFORMES'] - TAB_TIPO['EXPURGOS CONFORMES'] + TAB_TIPO['CONVERSÃO: NÃO CONFORME → CONFORME'] - TAB_TIPO['CONVERSÃO: CONFORME → NÃO CONFORME']).astype(int)

    TAB_TIPO["IQA (%)"] = (((TAB_TIPO["NAM CONF"]) / (TAB_TIPO["NAM REALIZ"])) * 100).round(1)

    TAB_TIPO_FINAL = TAB_TIPO[[
      "ANÁLISE", "PLANO DE AMOSTRAGEM", "DESCONSIDERAÇÕES DO PLANO",
      "PLANO DE AMOSTRAGEM AJUSTADO - ARSAL",
      "ANALISES REALIZADAS", "ANALISES CONFORMES", "EXPURGOS CONFORMES",
      "CONVERSÃO: CONFORME → NÃO CONFORME", "CONVERSÃO: NÃO CONFORME → CONFORME", "NAM CONF"
    ]]
    TAB_TIPO_FINAL.loc['Total'] = TAB_TIPO_FINAL.sum()
    TAB_TIPO_FINAL["ANÁLISE"].iat[-1] = ' '
    TAB_TIPO_FINAL.loc[TAB_TIPO_FINAL.index[-1], "ANÁLISE"] = "TOTAL"

    TAB_TIPO_FINAL.rename(columns={"PLANO DE AMOSTRAGEM": "PLANO DE AMOSTRAGEM AJUSTADO"})

    total_conf_vi_naoconf_arsal = int(TAB_TIPO_FINAL["CONVERSÃO: CONFORME → NÃO CONFORME"].iloc[-1])
    total_naoconf_vi_conf_arsal = int(TAB_TIPO_FINAL["CONVERSÃO: NÃO CONFORME → CONFORME"].iloc[-1])

    TAB_TIPO_FINAL_RESUMIDA = TAB_TIPO_FINAL[["ANÁLISE", "ANALISES REALIZADAS", "ANALISES CONFORMES", "EXPURGOS CONFORMES", "CONVERSÃO: CONFORME → NÃO CONFORME", "CONVERSÃO: NÃO CONFORME → CONFORME", "NAM CONF"]]

    TAB_REALIZADAS_ACEITO_CONF_VI = TAB_REALIZADAS_ACEITO[TAB_REALIZADAS["CONFORMIDADE - VI"] == "Conforme"]
    TAB_REALIZADAS_ACEITO_NAOCONF_VI = TAB_REALIZADAS_ACEITO[(TAB_REALIZADAS["CONFORMIDADE - VI"] != "Conforme") & (TAB_REALIZADAS["CONFORMIDADE - VI"].notna())]
    TAB_REALIZADAS_ACEITO_CONF_EXP_VI = TAB_REALIZADAS_ACEITO_CONF_VI[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI = TAB_REALIZADAS_ACEITO_NAOCONF_VI[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]

    contagem_tot_VI = TAB_REALIZADAS_ACEITO.groupby('ANÁLISE').size().reset_index(name='ANALISES REALIZADAS')
    TAB_TIPO_VI = TAB_TIPO_VI.merge(contagem_tot_VI, on='ANÁLISE', how='left')

    TAB_TIPO_VI = TAB_TIPO_VI.fillna(0)
    TAB_TIPO_VI['ANALISES REALIZADAS'] = TAB_TIPO_VI['ANALISES REALIZADAS'].astype(int)

    contagem_conf_VI = TAB_REALIZADAS_ACEITO_CONF_VI.groupby('ANÁLISE').size().reset_index(name='ANALISES CONFORMES')
    TAB_TIPO_VI = TAB_TIPO_VI.merge(contagem_conf_VI, on='ANÁLISE', how='left')

    contagem_naoconf_VI = TAB_REALIZADAS_ACEITO_NAOCONF_VI.groupby('ANÁLISE').size().reset_index(name='ANALISES NÃO CONFORMES')
    TAB_TIPO_VI = TAB_TIPO_VI.merge(contagem_naoconf_VI, on='ANÁLISE', how='left')
    TAB_TIPO_VI['ANALISES NÃO CONFORMES'] = TAB_TIPO_VI['ANALISES NÃO CONFORMES'].fillna(0).astype(int)

    contagem_conf_exp_VI = TAB_REALIZADAS_ACEITO_CONF_EXP_VI.groupby('ANÁLISE').size().reset_index(name='EXPURGOS CONFORMES')
    TAB_TIPO_VI = TAB_TIPO_VI.merge(contagem_conf_exp_VI, on='ANÁLISE', how='left')
    TAB_TIPO_VI['EXPURGOS CONFORMES'] = TAB_TIPO_VI['EXPURGOS CONFORMES'].fillna(0).astype(int)

    contagem_naoconf_exp_VI = TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI.groupby('ANÁLISE').size().reset_index(name='EXPURGOS NAO CONFORMES')
    TAB_TIPO_VI = TAB_TIPO_VI.merge(contagem_naoconf_exp_VI, on='ANÁLISE', how='left')
    TAB_TIPO_VI['EXPURGOS NAO CONFORMES'] = TAB_TIPO_VI['EXPURGOS NAO CONFORMES'].fillna(0).astype(int)

    TAB_TIPO_VI['EXPURGOS TOTAIS'] = TAB_TIPO_VI['EXPURGOS NAO CONFORMES'] + TAB_TIPO_VI['EXPURGOS CONFORMES']

    TAB_TIPO_VI["NAM REALIZ"] = (np.maximum(TAB_TIPO_VI['PLANO DE AMOSTRAGEM AJUSTADO - ARSAL'], TAB_TIPO_VI['ANALISES REALIZADAS']) - TAB_TIPO_VI['EXPURGOS TOTAIS']).astype(int)

    TAB_TIPO_VI = TAB_TIPO_VI.fillna(0)
    TAB_TIPO_VI["NAM CONF"] = (TAB_TIPO_VI['ANALISES CONFORMES'] - TAB_TIPO_VI['EXPURGOS CONFORMES']).astype(int)

    TAB_TIPO_VI["IQA (%)"] = (((TAB_TIPO_VI["NAM CONF"]) / (TAB_TIPO_VI["NAM REALIZ"])) * 100).round(1)

    TAB_TIPO_FINAL_VI = TAB_TIPO_VI[[
      "ANÁLISE", "PLANO DE AMOSTRAGEM", "DESCONSIDERAÇÕES DO PLANO",
      "PLANO DE AMOSTRAGEM AJUSTADO - ARSAL",
      "ANALISES REALIZADAS", "ANALISES CONFORMES", "EXPURGOS CONFORMES", "NAM CONF"
    ]]
    TAB_TIPO_FINAL_VI.loc['Total'] = TAB_TIPO_FINAL_VI.sum()
    TAB_TIPO_FINAL_VI["ANÁLISE"].iat[-1] = ' '
    TAB_TIPO_FINAL_VI.loc[TAB_TIPO_FINAL_VI.index[-1], "ANÁLISE"] = "TOTAL"

    TAB_TIPO_RESUMIDA_VI = TAB_TIPO_FINAL_VI[["ANÁLISE", "ANALISES REALIZADAS", "ANALISES CONFORMES"]]
    TAB_TIPO_RESUMIDA_VI["NAM CONF (VI)"] = TAB_TIPO_RESUMIDA_VI["ANALISES CONFORMES"]

    return TAB_TIPO, TAB_TIPO_FINAL_RESUMIDA, TAB_TIPO_RESUMIDA_VI, total_conf_vi_naoconf_arsal, total_naoconf_vi_conf_arsal



  # ELABORAÇÃO DA TABELA DE PONDERAÇÕES

  def pond(TAB_POND, TAB_REALIZADAS, BASE_MUN_TIPO_completo):

    TAB_REALIZADAS["ID_POND"] = TAB_REALIZADAS["ID_POND"].fillna(0)

    BASE_MUN_TIPO_completo["ID_POND"] = BASE_MUN_TIPO_completo["ID_POND"].fillna("").astype(str).str.strip()

    DESC_POR_ID = (
      BASE_MUN_TIPO_completo[BASE_MUN_TIPO_completo["ID_POND"] != ""]
      .groupby("ID_POND", as_index=False)["DESCONSIDERAÇÕES - ARSAL"]
      .sum()
    )
    DESC_POR_ID["ID_POND"] = DESC_POR_ID["ID_POND"].astype(float).astype("Int64")

    TAB_REALIZADAS['ANÁLISE'] = np.where(
        TAB_REALIZADAS['ANÁLISE'].isin(['TURBIDEZ', 'COR APARENTE', 'PH', 'COLIFORMES TOTAIS', 'ESCHERICHIA COLI']),
        TAB_REALIZADAS['ANÁLISE'],
        'DEMAIS PARAMETROS'
    )

    TAB_REALIZADAS_ACEITO = TAB_REALIZADAS[TAB_REALIZADAS["ACEITO?"] == "Validado"]

    TAB_EXPURGOS = TAB_REALIZADAS_ACEITO[
        TAB_REALIZADAS_ACEITO["EXPURGOS - ARSAL"] == "EXPURGAR"
    ]

    TAB_EXPURGOS_CONF = TAB_EXPURGOS[
        TAB_EXPURGOS["CONFORMIDADE - VI"] == "Conforme"
    ]

    TAB_SEM_EXPURGO = TAB_REALIZADAS_ACEITO[
        TAB_REALIZADAS_ACEITO["EXPURGOS - ARSAL"] != "EXPURGAR"
    ]

    TAB_CONF_VI_NAOCONF_ARSAL = TAB_SEM_EXPURGO[
        (TAB_SEM_EXPURGO["CONFORMIDADE - VI"] == "Conforme") &
        (TAB_SEM_EXPURGO["CONFORMIDADE - ARSAL"] != "Conforme") &
        (TAB_SEM_EXPURGO["CONFORMIDADE - ARSAL"].notna())
    ]

    TAB_NAOCONF_VI_CONF_ARSAL = TAB_SEM_EXPURGO[
        (TAB_SEM_EXPURGO["CONFORMIDADE - VI"] != "Conforme") &
        (TAB_SEM_EXPURGO["CONFORMIDADE - VI"].notna()) &
        (TAB_SEM_EXPURGO["CONFORMIDADE - ARSAL"] == "Conforme")
    ]

    QUADRO_POND_NCONF = TAB_POND.copy()
    QUADRO_POND_NCONF = QUADRO_POND_NCONF.rename(columns={
      QUADRO_POND_NCONF.columns[0]: "ID",
      QUADRO_POND_NCONF.columns[1]: "PONDERAÇÕES (PRESTADORA)",
      QUADRO_POND_NCONF.columns[2]: "PONDERAÇÕES (VI)",
      QUADRO_POND_NCONF.columns[3]: "PONDERAÇÕES (ARSAL)",
    })

    contagem_exp = TAB_EXPURGOS_CONF.groupby('ID_POND').size().reset_index(name='EXPURGOS')
    QUADRO_POND_NCONF = QUADRO_POND_NCONF.merge(contagem_exp, left_on='ID', right_on='ID_POND', how='left')
    QUADRO_POND_NCONF = QUADRO_POND_NCONF.drop(columns=["ID_POND"])
    QUADRO_POND_NCONF["EXPURGOS"] = QUADRO_POND_NCONF["EXPURGOS"].fillna(0).astype("Int64")
    QUADRO_POND_NCONF["EXPURGOS"] = QUADRO_POND_NCONF["EXPURGOS"].astype("object").replace(0, "")

    contagem_cv_na = TAB_CONF_VI_NAOCONF_ARSAL.groupby('ID_POND').size().reset_index(name='CONVERSÃO: CONFORME → NÃO CONFORME')
    QUADRO_POND_NCONF = QUADRO_POND_NCONF.merge(contagem_cv_na, left_on='ID', right_on='ID_POND', how='left')
    QUADRO_POND_NCONF = QUADRO_POND_NCONF.drop(columns=["ID_POND"], errors="ignore")
    QUADRO_POND_NCONF["CONVERSÃO: CONFORME → NÃO CONFORME"] = QUADRO_POND_NCONF["CONVERSÃO: CONFORME → NÃO CONFORME"].fillna(0).astype("Int64")
    QUADRO_POND_NCONF["CONVERSÃO: CONFORME → NÃO CONFORME"] = QUADRO_POND_NCONF["CONVERSÃO: CONFORME → NÃO CONFORME"].astype("object").replace(0, "")

    contagem_na_cv = TAB_NAOCONF_VI_CONF_ARSAL.groupby('ID_POND').size().reset_index(name='CONVERSÃO: NÃO CONFORME → CONFORME')
    QUADRO_POND_NCONF = QUADRO_POND_NCONF.merge(contagem_na_cv, left_on='ID', right_on='ID_POND', how='left')
    QUADRO_POND_NCONF = QUADRO_POND_NCONF.drop(columns=["ID_POND"], errors="ignore")
    QUADRO_POND_NCONF["CONVERSÃO: NÃO CONFORME → CONFORME"] = QUADRO_POND_NCONF["CONVERSÃO: NÃO CONFORME → CONFORME"].fillna(0).astype("Int64")
    QUADRO_POND_NCONF["CONVERSÃO: NÃO CONFORME → CONFORME"] = QUADRO_POND_NCONF["CONVERSÃO: NÃO CONFORME → CONFORME"].astype("object").replace(0, "")

    colunas_impacto = ['EXPURGOS', 'CONVERSÃO: CONFORME → NÃO CONFORME', 'CONVERSÃO: NÃO CONFORME → CONFORME']
    QUADRO_POND_NCONF = QUADRO_POND_NCONF[
        QUADRO_POND_NCONF[colunas_impacto].apply(lambda row: any(v != "" for v in row), axis=1)
    ]

    for col in colunas_impacto:
        if col in QUADRO_POND_NCONF.columns:
            if QUADRO_POND_NCONF[col].replace("", pd.NA).isna().all():
                QUADRO_POND_NCONF = QUADRO_POND_NCONF.drop(columns=[col])

    if not TAB_EXPURGOS.empty:

      QUADRO_POND_REALIZ = TAB_POND.copy()
      QUADRO_POND_REALIZ = QUADRO_POND_REALIZ.rename(columns={
        QUADRO_POND_REALIZ.columns[0]: "ID",
        QUADRO_POND_REALIZ.columns[1]: "PONDERAÇÕES (PRESTADORA)",
        QUADRO_POND_REALIZ.columns[2]: "PONDERAÇÕES (VI)",
        QUADRO_POND_REALIZ.columns[3]: "PONDERAÇÕES (ARSAL)",
      })

      contagem = TAB_EXPURGOS.groupby('ID_POND').size().reset_index(name='EXPURGOS')
      QUADRO_POND_REALIZ = QUADRO_POND_REALIZ.merge(contagem, left_on='ID', right_on='ID_POND', how='left')
      QUADRO_POND_REALIZ = QUADRO_POND_REALIZ.drop(columns=["ID_POND"])
      QUADRO_POND_REALIZ["EXPURGOS"] = QUADRO_POND_REALIZ["EXPURGOS"].fillna(0).astype("Int64")
      QUADRO_POND_REALIZ["EXPURGOS"] = QUADRO_POND_REALIZ["EXPURGOS"].astype("object").replace(0, "")

      QUADRO_POND_REALIZ = QUADRO_POND_REALIZ.merge(DESC_POR_ID, left_on='ID', right_on='ID_POND', how='left')
      QUADRO_POND_REALIZ = QUADRO_POND_REALIZ.drop(columns=["ID_POND"])
      QUADRO_POND_REALIZ = QUADRO_POND_REALIZ.rename(columns={"DESCONSIDERAÇÕES - ARSAL": "DESCONSIDERAÇÕES (PLANO)"})
      QUADRO_POND_REALIZ["DESCONSIDERAÇÕES (PLANO)"] = QUADRO_POND_REALIZ["DESCONSIDERAÇÕES (PLANO)"].fillna(0).astype("Int64")
      QUADRO_POND_REALIZ["DESCONSIDERAÇÕES (PLANO)"] = QUADRO_POND_REALIZ["DESCONSIDERAÇÕES (PLANO)"].astype("object").replace(0, "")

    else:
      QUADRO_POND_REALIZ = "SEM EXPURGOS"

    return QUADRO_POND_NCONF, QUADRO_POND_REALIZ



  # CONSTRUÇÃO DOS GRÁFICOS

  def graf(TAB_MUN_FINAL, TAB_TIPO):

    TAB_MUN_FINAL = TAB_MUN_FINAL[:-1]

    def remove_acentos(txt):
      return ''.join(c for c in unicodedata.normalize('NFD', txt)
                    if unicodedata.category(c) != 'Mn')

    TAB_MUN_FINAL['MUNICÍPIO_SORT'] = TAB_MUN_FINAL['MUNICÍPIO'].apply(remove_acentos)

    TAB_MUN_FINAL = TAB_MUN_FINAL.sort_values(by='MUNICÍPIO_SORT', ascending=False)
    TAB_MUN_FINAL = TAB_MUN_FINAL.drop(columns=['MUNICÍPIO_SORT'])

    TAB_TIPO= TAB_TIPO.sort_values(by="ANÁLISE", ascending=False)


    # Plano de amostragem vs análises realizadas por município

    espacamento = 1.5
    y = np.arange(len(TAB_MUN_FINAL["MUNICÍPIO"])) * espacamento
    largura = 0.4  # menor para caber 3 barras separadas

    fig, ax3 = plt.subplots(figsize=(12, len(TAB_MUN_FINAL["MUNICÍPIO"]) * 0.4 * espacamento))

    bars1 = ax3.barh(y + largura,   TAB_MUN_FINAL["ANALISES REALIZADAS"],                  height=largura, label="ANALISES REALIZADAS",                  color="green")
    bars2 = ax3.barh(y,             TAB_MUN_FINAL["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"], height=largura, label="PLANO DE AMOSTRAGEM AJUSTADO - ARSAL", color="skyblue")
    bars3 = ax3.barh(y - largura,   TAB_MUN_FINAL["EXPURGOS TOTAIS"],                      height=largura, label="EXPURGOS",                              color="orange")

    ax3.set_yticks(y)
    ax3.set_yticklabels(TAB_MUN_FINAL["MUNICÍPIO"])
    ax3.set_xlabel("Amostras")

    for bar in bars1:
      val = int(bar.get_width())
      if val > 0:
        ax3.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                f'{val}', va='center', fontsize=10)

    for bar in bars2:
      val = int(bar.get_width())
      if val > 0:
        ax3.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                f'{val}', va='center', fontsize=10)

    for bar in bars3:
      val = int(bar.get_width())
      if val > 0:
        ax3.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                f'{val}', va='center', fontsize=10)

    ax3.tick_params(axis='y', labelsize=11)
    ax3.grid(axis='x', linestyle='--', alpha=0.1)
    ax3.legend()

    plt.tight_layout()
    plt.savefig("grafico_plan_realiz_mun.png", dpi=300, bbox_inches="tight")
    plt.show()


    # Análises realizadas vs Análises conformes por município

    espacamento = 1.5
    y = np.arange(len(TAB_MUN_FINAL["MUNICÍPIO"])) * espacamento
    largura = 0.6

    fig, ax4 = plt.subplots(figsize=(12, len((TAB_MUN_FINAL["MUNICÍPIO"]))*0.4*espacamento))

    bars1 = ax4.barh(y + largura/2, TAB_MUN_FINAL["ANALISES CONFORMES"]-TAB_MUN_FINAL["EXPURGOS CONFORMES"], height=largura, label="ANALISES CONFORMES (SUBTRAINDO OS EXPURGOS CONFORMES)", color="green")
    bars2 = ax4.barh(y - largura/2, TAB_MUN_FINAL["ANALISES REALIZADAS"]-TAB_MUN_FINAL["EXPURGOS TOTAIS"], height=largura, label="ANALISES REALIZADAS (SUBTRAINDO OS EXPURGOS TOTAIS)", color="orange")

    ax4.set_yticks(y)
    ax4.set_yticklabels(TAB_MUN_FINAL["MUNICÍPIO"])
    ax4.set_xlabel("Amostras")

    for bar in bars1:
      ax4.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{int(bar.get_width())}', va='center', fontsize=10)

    for bar in bars2:
      ax4.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{int(bar.get_width())}', va='center', fontsize=10)

    ax4.tick_params(axis='y', labelsize=11)

    ax4.grid(axis='x', linestyle='--', alpha=0.1)
    ax4.legend()

    plt.savefig("grafico_conf_realiz_mun.png", dpi=300, bbox_inches="tight")

    plt.show()


    # IQA por município

    espacamento = 1.5
    y = np.arange(len(TAB_MUN_FINAL["MUNICÍPIO"])) * espacamento
    largura = 0.6

    fig, ax5 = plt.subplots(figsize=(12, len((TAB_MUN_FINAL["MUNICÍPIO"]))*0.4*espacamento))

    bars1 = ax5.barh(y + largura/2, TAB_MUN_FINAL["IQA (%)"], height=largura, label="IQA(%)", color="orange")

    ax5.set_yticks(y)
    ax5.set_yticklabels(TAB_MUN_FINAL["MUNICÍPIO"])
    ax5.set_xlabel("IQA (%)")

    for bar in bars1:
      ax5.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{bar.get_width():.1f}', va='center', fontsize=10)

    ax5.tick_params(axis='y', labelsize=11)

    ax5.grid(axis='x', linestyle='--', alpha=0.1)
    ax5.legend()

    plt.savefig("grafico_iqa_mun.png", dpi=300, bbox_inches="tight")

    plt.show()


    # Plano de amostragem vs análises realizadas por tipo de análise

    TAB_TIPO["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"] = TAB_TIPO["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"].astype(int)

    x = np.arange(len(TAB_TIPO["ANÁLISE"]))
    largura = 0.6

    fig, ax6 = plt.subplots(figsize=(10, 8))

    barras = ax6.bar(x, TAB_TIPO["ANALISES REALIZADAS"], width=largura, label="ANALISES REALIZADAS", color='skyblue')
    ax6.plot(x, TAB_TIPO["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"], label="PLANO DE AMOSTRAGEM AJUSTADO - ARSAL", color='green', marker='o', linewidth=2, linestyle='--')

    for barra in barras:
      altura = barra.get_height()
      ax6.text(barra.get_x() + barra.get_width()/2, altura/2, f'{altura}', ha='center', va='center', color='black', fontsize=10)

    for p, valor in enumerate(TAB_TIPO["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"]):
      ax6.text(x[p] + 0.1, valor + 1.5, f'{valor}', ha='left', va='bottom', color='green', fontsize=10)

    ax6.grid(True, linestyle='--', alpha=0.1)

    ax6.set_xticks(x)
    ax6.set_xticklabels(TAB_TIPO["ANÁLISE"], rotation=30)

    ax6.legend()

    ax6.set_ylabel('Amostras')

    plt.savefig("grafico_plan_realiz_tipo.png", dpi=300, bbox_inches="tight")

    plt.show()


    # Análises realizadas vs Análises conformes por tipo de análise

    espacamento = 1.5
    y = np.arange(len(TAB_TIPO["ANÁLISE"])) * espacamento
    largura = 0.6

    fig, ax7 = plt.subplots(figsize=(12, len((TAB_TIPO["ANÁLISE"]))*0.5*espacamento))

    bars1 = ax7.barh(y + largura/2, TAB_TIPO["ANALISES CONFORMES"] - TAB_TIPO["EXPURGOS CONFORMES"], height=largura, label="ANALISES CONFORMES (SUBTRAINDO OS EXPURGOS CONFORMES)", color="green")
    bars2 = ax7.barh(y - largura/2, TAB_TIPO["ANALISES REALIZADAS"] - TAB_TIPO["EXPURGOS TOTAIS"], height=largura, label="ANALISES REALIZADAS (SUBTRAINDO OS EXPURGOS TOTAIS)", color="orange")

    ax7.set_yticks(y)
    ax7.set_yticklabels(TAB_TIPO["ANÁLISE"])
    ax7.set_xlabel("Amostras")

    for bar in bars1:
      ax7.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{int(bar.get_width())}', va='center', fontsize=10)

    for bar in bars2:
      ax7.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{int(bar.get_width())}', va='center', fontsize=10)

    ax7.grid(axis='x', linestyle='--', alpha=0.1)
    ax7.legend()

    plt.savefig("grafico_conf_realiz_tipo.png", dpi=300, bbox_inches="tight")

    plt.show()


  # FUNÇÃO PRINCIPAL E EXECUÇÃO - IQA

  def MAIN_IQA():

    arquivo_gerais = pd.read_excel('ATT_SQL.xlsx', sheet_name='GERAIS')

    MES_ANO, ANO_CONTRAT, PRESTADORA, META, NCOF_P, NREA_P, IQA_P, NCOF_VI, NREA_VI, IQA_VI, FONTE_NREALIZ, FONTE_PLANO, FONT_NCONF, ACREDITAÇÃO, MESMIN_ANO = infos_gerais (arquivo_gerais)


    PLAN_MUN_TIPO = pd.read_excel ('ATT_SQL.xlsx', sheet_name='IQA_PLANO')

    PLAN_MUN_TIPO = PLAN_MUN_TIPO.rename(columns={'DESCONSIDERACOES - ARSAL' : 'DESCONSIDERAÇÕES - ARSAL'})

    AMOSTRAS_REALIZADAS = pd.read_excel ('ATT_SQL.xlsx', sheet_name='IQA_DETALHADO')

    AMOSTRAS_REALIZADAS = AMOSTRAS_REALIZADAS.rename(columns={
    'id_amostra': 'ID_AMOSTRA',
    'saa': 'CONTA',
    'analise': 'ANÁLISE',
    'municipio': 'CIDADE',
    'parecer': 'ACEITO?',
    'conformidade_vi': 'CONFORMIDADE - VI',
    'resultado': 'CONFORMIDADE - ARSAL',
    'expurgos': 'EXPURGOS - ARSAL',
    'id_pond': 'ID_POND',
    })

    PONDERACOES = pd.read_excel ('ATT_SQL.xlsx', sheet_name='IQA_PONDERACOES')


    PLAN_MUN_TIPO = PLAN_MUN_TIPO.fillna(0)
    AMOSTRAS_REALIZADAS = AMOSTRAS_REALIZADAS.fillna(0)
    PONDERACOES = PONDERACOES.fillna(0)


    TAB_MUN, TAB_TIPO, TAB_REALIZADAS, TAB_POND, BASE_MUN_TIPO_completo = base(PLAN_MUN_TIPO, AMOSTRAS_REALIZADAS, PONDERACOES, MES_ANO)

    TAB_MUN_FINAL, TAB_MUN_ARSAL, TAB_MUN_VI, desc_plano, TAB_MUN_NCONF, TAB_MUN_NCONF_VI = analise_mun(TAB_MUN, TAB_REALIZADAS, BASE_MUN_TIPO_completo)

    AN_TIPO, TAB_TIPO_FINAL, TAB_TIPO_RESUMIDA_VI, total_conf_vi_naoconf_arsal, total_naoconf_vi_conf_arsal  = analise_tipo(TAB_TIPO, TAB_REALIZADAS, BASE_MUN_TIPO_completo)

    graf(TAB_MUN_FINAL, AN_TIPO)

    QUADRO_POND_NCONF, QUADRO_POND_REALIZ = pond(TAB_POND, TAB_REALIZADAS, BASE_MUN_TIPO_completo)

    dic_info = var_doc (MES_ANO, ANO_CONTRAT, PRESTADORA, META, NCOF_P, NREA_P, IQA_P, NCOF_VI, NREA_VI, IQA_VI, FONTE_NREALIZ, FONTE_PLANO, FONT_NCONF, ACREDITAÇÃO, TAB_MUN_FINAL, QUADRO_POND_NCONF, QUADRO_POND_REALIZ, TAB_MUN_ARSAL, TAB_MUN_VI, TAB_TIPO_FINAL, TAB_TIPO_RESUMIDA_VI, MESMIN_ANO, total_conf_vi_naoconf_arsal, total_naoconf_vi_conf_arsal, desc_plano, TAB_MUN_NCONF, TAB_MUN_NCONF_VI)

    arquivo_entrada = "modelo_iqa_automatico.docx"

    doc_modificado = substituir_var(arquivo_entrada, dic_info)

    arquivo_saida = f'INDICADORES- ' + str(PRESTADORA) + ' - ' + str(MES_ANO) + '.docx'

    doc_modificado.save(arquivo_saida)

    return arquivo_saida


  lat_long = pd.read_excel('lat_long.xlsx')

  url = "https://github.com/carlosfarsal/DASH_ATT/raw/main/modelo_iqa_automatico.docx"

  output = "modelo_iqa_automatico.docx"

  print(f"⏳ Baixando {output} de {url}...")
    
  try:
    gdown.download(url, output, quiet=False)
  except Exception as e:
    print(f"❌ Erro ao baixar: {e}")

  # Verifica se o arquivo foi baixado
  if not os.path.exists(output):
    raise FileNotFoundError(
        f"❌ ARQUIVO NÃO ENCONTRADO: {output}\n"
        f"URL tentada: {url}\n"
        f"Verifique se o arquivo existe no repositório GitHub."
    )
  
  nome_arq = MAIN_IQA()

  return nome_arq


##############################################################################################################################################################################################################################

# ELABORAÇÃO DE RELATÓRIO AUTOMATIZADO - IQA E IQE

def rel_iqa_iqe():

  # DEFININDO E AJUSTANDO AS TABELAS BASES - IQA
  
  def base(PLAN_MUN_TIPO, PLAN_REALIZADAS, PLAN_POND, MES_ANO):

    arquivo_saida   = "PLANO_TRATADO.xlsx"
    coluna_valor    = "PLANO"
    coluna_desc     = "DESCONSIDERAÇÕES - ARSAL"
    coluna_id_pond  = "ID_POND"

    BASE_MUN_TIPO = PLAN_MUN_TIPO

    BASE_MUN_TIPO = BASE_MUN_TIPO.dropna(subset=['PLANO'])

    mes, ano = MES_ANO.split("-")
    ano = ano[:4]

    meses_en = {
      "JANEIRO": "January", "FEVEREIRO": "February", "MARÇO": "March",
      "ABRIL": "April", "MAIO": "May", "JUNHO": "June",
      "JULHO": "July", "AGOSTO": "August", "SETEMBRO": "September",
      "OUTUBRO": "October", "NOVEMBRO": "November", "DEZEMBRO": "December"
    }

    mes = meses_en[mes]

    BASE_MUN_TIPO["Parâmetros"] = BASE_MUN_TIPO["Parâmetros"].replace({
      "Cor": "Cor Aparente", "Residual de Cloro Livre": "Cloro residual livre",
      "Residual de cloro livre": "Cloro residual livre", "Cor aparente": "Cor Aparente",
      "Cloro": "Cloro residual livre", "Coliformes Totais": "Coliformes totais",
      "Escherichia coli": "Escherichia Coli", "Cloro Residual Livre": "Cloro residual livre", "Cloro Residual livre": "Cloro residual livre", "Cloro residual Livre": "Cloro residual livre",
      "Cloro Livre": "Cloro residual livre", "Cloro livre": "Cloro residual livre", "PH": "pH", "ph": "pH", "Ph": "pH"
    })

    parametros_principais = [
      "Turbidez", "Cor Aparente", "pH",
      "Cloro residual livre", "Coliformes totais", "Escherichia Coli"
    ]

    BASE_MUN_TIPO["Parâmetros"] = BASE_MUN_TIPO["Parâmetros"].apply(
      lambda x: x if x in parametros_principais else "Demais Parâmetros"
    )

    BASE_MUN_TIPO[coluna_id_pond] = BASE_MUN_TIPO[coluna_id_pond].fillna("").astype(str).str.strip()

    colunas_chave = ["Cidade", "Parâmetros", coluna_id_pond]

    BASE_MUN_TIPO_agrupado = BASE_MUN_TIPO.groupby(
      colunas_chave, as_index=False
    ).agg(
      **{
        coluna_valor: (coluna_valor, "sum"),
        coluna_desc:  (coluna_desc,  "sum"),
      }
    )

    cols = [c for c in BASE_MUN_TIPO_agrupado.columns if c != coluna_id_pond] + [coluna_id_pond]
    BASE_MUN_TIPO_agrupado = BASE_MUN_TIPO_agrupado[cols]

    BASE_MUN_TIPO_agrupado.insert(0, "MÊS", mes)
    BASE_MUN_TIPO_agrupado.insert(1, "ANO", ano)

    municipios = BASE_MUN_TIPO_agrupado["Cidade"].unique()
    combinacoes_principais = pd.DataFrame(
      list(itertools.product([mes], [ano], municipios, parametros_principais, [""])),
      columns=["MÊS", "ANO", "Cidade", "Parâmetros", coluna_id_pond]
    )

    BASE_MUN_TIPO_principais = pd.merge(
      combinacoes_principais, BASE_MUN_TIPO_agrupado,
      on=["MÊS", "ANO", "Cidade", "Parâmetros", coluna_id_pond], how="left"
    )
    BASE_MUN_TIPO_principais[coluna_valor] = BASE_MUN_TIPO_principais[coluna_valor].fillna(0)
    BASE_MUN_TIPO_principais[coluna_desc]  = BASE_MUN_TIPO_principais[coluna_desc].fillna(0)

    BASE_COM_ID = BASE_MUN_TIPO_agrupado[
      BASE_MUN_TIPO_agrupado[coluna_id_pond] != ""
    ]

    BASE_MUN_TIPO_demais = BASE_MUN_TIPO_agrupado[
      (BASE_MUN_TIPO_agrupado["Parâmetros"] == "Demais Parâmetros") &
      (BASE_MUN_TIPO_agrupado[coluna_id_pond] == "")
    ]

    BASE_MUN_TIPO_completo = pd.concat(
      [BASE_MUN_TIPO_principais, BASE_COM_ID, BASE_MUN_TIPO_demais],
      ignore_index=True
    )

    cols = [c for c in BASE_MUN_TIPO_completo.columns if c != coluna_id_pond] + [coluna_id_pond]
    BASE_MUN_TIPO_completo = BASE_MUN_TIPO_completo[cols]

    BASE_MUN_TIPO_completo.to_excel(arquivo_saida, index=False)

    BASE_MUN = BASE_MUN_TIPO_agrupado.groupby("Cidade", as_index=False)[coluna_valor].sum()
    BASE_MUN = BASE_MUN.rename(columns={"Cidade": "MUNICÍPIO", "PLANO": "PLANO DE AMOSTRAGEM"})
    BASE_MUN["MUNICÍPIO"] = BASE_MUN["MUNICÍPIO"].str.upper().apply(unidecode)

    BASE_TIPO = BASE_MUN_TIPO_agrupado.groupby("Parâmetros", as_index=False)[coluna_valor].sum()
    BASE_TIPO = BASE_TIPO.rename(columns={"Parâmetros": "ANÁLISE", "PLANO": "PLANO DE AMOSTRAGEM"})
    BASE_TIPO["ANÁLISE"] = BASE_TIPO["ANÁLISE"].str.upper().apply(unidecode)

    BASE_REALIZADAS = PLAN_REALIZADAS

    BASE_REALIZADAS["ANÁLISE"] = BASE_REALIZADAS["ANÁLISE"].replace({
      "Cor": "Cor Aparente", "Residual de Cloro Livre": "Cloro residual livre",
      "Residual de cloro livre": "Cloro residual livre", "Cor aparente": "Cor Aparente",
      "Cloro": "Cloro residual livre", "Coliformes Totais": "Coliformes totais",
      "Escherichia coli": "Escherichia Coli", "Cloro Residual Livre": "Cloro residual livre", "Cloro Residual livre": "Cloro residual livre", "Cloro residual Livre": "Cloro residual livre",
      "Cloro Livre": "Cloro residual livre", "Cloro livre": "Cloro residual livre", "PH": "pH", "ph": "pH", "Ph": "pH"})

    BASE_REALIZADAS["CIDADE"] = BASE_REALIZADAS["CIDADE"].str.upper().apply(unidecode)
    BASE_REALIZADAS["ANÁLISE"] = BASE_REALIZADAS["ANÁLISE"].str.upper().apply(unidecode)

    BASE_POND = PLAN_POND

    return (BASE_MUN, BASE_TIPO, BASE_REALIZADAS, BASE_POND, BASE_MUN_TIPO_completo)



  # DEFININDO E AJUSTANDO AS TABELAS BASES - IQE

  def base_IQE(PLAN_MUN_TIPO_IQE, PLAN_REALIZADAS_IQE, PLAN_POND_IQE, TRIM_ANO):

    arquivo_saida  = "PLANO_TRATADO_IQE.xlsx"
    coluna_valor   = "PLANO"
    coluna_desc    = "DESCONSIDERAÇÕES - ARSAL"
    coluna_id_pond = "ID_POND"

    BASE_MUN_TIPO_IQE = PLAN_MUN_TIPO_IQE

    BASE_MUN_TIPO_IQE = BASE_MUN_TIPO_IQE.dropna(subset=[coluna_valor])

    trimestre, ano = TRIM_ANO.split(" - ANO ")
    ano_contratual = ano[:1]

    BASE_MUN_TIPO_IQE["Parâmetros"] = BASE_MUN_TIPO_IQE["Parâmetros"].replace({
      "DBO - 5 dias"          : "DBO",
      "Óleos e Graxas Totais" : "Óleos e Graxas",
      "Partículas Flutuantes" : "Materiais Flutuantes",
      "pH (a 25°C)"           : "pH"
    })

    parametros_principais = [
      "DBO",
      "DQO",
      "Temperatura",
      "Óleos e Graxas",
      "Materiais Flutuantes",
      "pH"
    ]

    BASE_MUN_TIPO_IQE["Parâmetros"] = BASE_MUN_TIPO_IQE["Parâmetros"].apply(
      lambda x: x if x in parametros_principais else "Demais Parâmetros"
    )

    BASE_MUN_TIPO_IQE[coluna_id_pond] = (
      BASE_MUN_TIPO_IQE[coluna_id_pond].fillna("").astype(str).str.strip()
    )

    colunas_chave = ["Cidade", "Parâmetros", coluna_id_pond]

    BASE_MUN_TIPO_IQE_agrupado = BASE_MUN_TIPO_IQE.groupby(
      colunas_chave, as_index=False
    ).agg(
      **{
        coluna_valor : (coluna_valor, "sum"),
        coluna_desc  : (coluna_desc,  "sum"),
      }
    )

    cols = [c for c in BASE_MUN_TIPO_IQE_agrupado.columns if c != coluna_id_pond] + [coluna_id_pond]
    BASE_MUN_TIPO_IQE_agrupado = BASE_MUN_TIPO_IQE_agrupado[cols]

    BASE_MUN_TIPO_IQE_agrupado.insert(0, "TRIMESTRE", trimestre)
    BASE_MUN_TIPO_IQE_agrupado.insert(1, "ANO CONTRATUAL", ano_contratual)

    municipios = BASE_MUN_TIPO_IQE_agrupado["Cidade"].unique()
    combinacoes_principais = pd.DataFrame(
      list(itertools.product([trimestre], [ano_contratual], municipios, parametros_principais, [""])),
      columns=["TRIMESTRE", "ANO CONTRATUAL", "Cidade", "Parâmetros", coluna_id_pond]
    )

    BASE_MUN_TIPO_IQE_principais = pd.merge(
      combinacoes_principais, BASE_MUN_TIPO_IQE_agrupado,
      on=["TRIMESTRE", "ANO CONTRATUAL", "Cidade", "Parâmetros", coluna_id_pond], how="left"
    )
    BASE_MUN_TIPO_IQE_principais[coluna_valor] = BASE_MUN_TIPO_IQE_principais[coluna_valor].fillna(0)
    BASE_MUN_TIPO_IQE_principais[coluna_desc]  = BASE_MUN_TIPO_IQE_principais[coluna_desc].fillna(0)

    BASE_COM_ID_IQE = BASE_MUN_TIPO_IQE_agrupado[
      BASE_MUN_TIPO_IQE_agrupado[coluna_id_pond] != ""
    ]

    BASE_MUN_TIPO_IQE_demais = BASE_MUN_TIPO_IQE_agrupado[
      (BASE_MUN_TIPO_IQE_agrupado["Parâmetros"] == "Demais Parâmetros") &
      (BASE_MUN_TIPO_IQE_agrupado[coluna_id_pond] == "")
    ]

    BASE_MUN_TIPO_IQE_completo = pd.concat(
      [BASE_MUN_TIPO_IQE_principais, BASE_COM_ID_IQE, BASE_MUN_TIPO_IQE_demais],
      ignore_index=True
    )

    cols = [c for c in BASE_MUN_TIPO_IQE_completo.columns if c != coluna_id_pond] + [coluna_id_pond]
    BASE_MUN_TIPO_IQE_completo = BASE_MUN_TIPO_IQE_completo[cols]

    BASE_MUN_TIPO_IQE_completo.to_excel(arquivo_saida, index=False)

    BASE_MUN_IQE = BASE_MUN_TIPO_IQE_agrupado.groupby("Cidade", as_index=False)[coluna_valor].sum()
    BASE_MUN_IQE = BASE_MUN_IQE.rename(columns={"Cidade": "MUNICÍPIO", "PLANO": "PLANO DE AMOSTRAGEM"})
    BASE_MUN_IQE["MUNICÍPIO"] = BASE_MUN_IQE["MUNICÍPIO"].str.upper().apply(unidecode)

    BASE_TIPO_IQE = BASE_MUN_TIPO_IQE_agrupado.groupby("Parâmetros", as_index=False)[coluna_valor].sum()
    BASE_TIPO_IQE = BASE_TIPO_IQE.rename(columns={"Parâmetros": "ANÁLISE", "PLANO": "PLANO DE AMOSTRAGEM"})
    BASE_TIPO_IQE["ANÁLISE"] = BASE_TIPO_IQE["ANÁLISE"].str.upper().apply(unidecode)

    BASE_REALIZADAS_IQE = PLAN_REALIZADAS_IQE

    BASE_REALIZADAS_IQE["ANÁLISE"] = BASE_REALIZADAS_IQE["ANÁLISE"].replace({
      "DBO - 5 dias"          : "DBO",
      "Óleos e Graxas Totais" : "Óleos e Graxas",
      "Partículas Flutuantes" : "Materiais Flutuantes",
      "pH (a 25°C)"           : "pH"
    })
    BASE_REALIZADAS_IQE["CIDADE"]  = BASE_REALIZADAS_IQE["CIDADE"].str.upper().apply(unidecode)
    BASE_REALIZADAS_IQE["ANÁLISE"] = BASE_REALIZADAS_IQE["ANÁLISE"].str.upper().apply(unidecode)

    BASE_POND_IQE = PLAN_POND_IQE

    return (BASE_MUN_IQE, BASE_TIPO_IQE, BASE_REALIZADAS_IQE, BASE_POND_IQE, BASE_MUN_TIPO_IQE_completo)



  # COLETANDO AS INFORMAÇÕES GERAIS
  
  def infos_gerais (TAB_INFO):

    mes_ing = TAB_INFO["INFORMAÇÃO"][3]
    ano = TAB_INFO["INFORMAÇÃO"][1]

    meses = {
    'January': 'JANEIRO',
    'February': 'FEVEREIRO',
    'March': 'MARÇO',
    'April': 'ABRIL',
    'May': 'MAIO',
    'June': 'JUNHO',
    'July': 'JULHO',
    'August': 'AGOSTO',
    'September': 'SETEMBRO',
    'October': 'OUTUBRO',
    'November': 'NOVEMBRO',
    'December': 'DEZEMBRO'
    }

    mes = meses[mes_ing]
    MES_ANO = mes + '-' + str(ano)
    mes_min = mes.lower()

    trimestre = TAB_INFO["INFORMAÇÃO"][4]
    ANO_CONTRAT = TAB_INFO["INFORMAÇÃO"][2]

    TRIM_ANO = (trimestre + ' - ANO ' + str(ANO_CONTRAT)).upper()
    TRIMMIN_ANO = TRIM_ANO.lower()

    MESMIN_ANO = mes_min + ' de ' + str(ano)

    PRESTADORA_SIGLA = TAB_INFO["INFORMAÇÃO"][0]

    if PRESTADORA_SIGLA == "BRK":
      PRESTADORA = "BRK Ambiental"
    elif PRESTADORA_SIGLA == "ADS":
      PRESTADORA = "Conasa Águas do Sertão"
    else:
      PRESTADORA = "Verde Ambiental Alagoas"

    META = TAB_INFO ["INFORMAÇÃO"][11]

    META = float(META)

    META = "100%" if META >= 1 else f"{META*100:.1f}%".replace('.', ',').rstrip('0').rstrip(',')

    NCOF_P = TAB_INFO ["INFORMAÇÃO"][5]
    NREA_P = TAB_INFO ["INFORMAÇÃO"][6]
    IQA_P = TAB_INFO ["INFORMAÇÃO"][7]

    IQA_P = float(IQA_P)

    IQA_P = "100%" if IQA_P >= 1 else f"{IQA_P*100:.1f}%".replace('.', ',').rstrip('0').rstrip(',')

    NCOF_VI = TAB_INFO ["INFORMAÇÃO"][8]
    NREA_VI = TAB_INFO ["INFORMAÇÃO"][9]
    IQA_VI = TAB_INFO ["INFORMAÇÃO"][10]

    IQA_VI = float(IQA_VI)

    IQA_VI = "100%" if IQA_VI >= 1 else f"{IQA_VI*100:.1f}%".replace('.', ',').rstrip('0').rstrip(',')

    FONTE_NREALIZ = TAB_INFO ["INFORMAÇÃO"][12]
    FONTE_PLANO = TAB_INFO ["INFORMAÇÃO"][13]
    FONT_NCONF = TAB_INFO ["INFORMAÇÃO"][14]
    ACREDITAÇÃO = TAB_INFO ["INFORMAÇÃO"][15]

    META_IQE = TAB_INFO ["INFORMAÇÃO"][22]

    META_IQE = float(META_IQE)

    META_IQE = "100%" if META_IQE >= 1 else f"{META_IQE*100:.1f}%".replace('.', ',').rstrip('0').rstrip(',')

    NCOF_P_IQE = TAB_INFO ["INFORMAÇÃO"][16]
    NREA_P_IQE = TAB_INFO ["INFORMAÇÃO"][17]
    IQE_P = TAB_INFO ["INFORMAÇÃO"][18]

    IQE_P = float(IQE_P)

    IQE_P = "100%" if IQE_P >= 1 else f"{IQE_P*100:.1f}%".replace('.', ',').rstrip('0').rstrip(',')

    NCOF_VI_IQE = TAB_INFO ["INFORMAÇÃO"][19]
    NREA_VI_IQE = TAB_INFO ["INFORMAÇÃO"][20]
    IQE_VI = TAB_INFO ["INFORMAÇÃO"][21]

    IQE_VI = float(IQE_VI)

    IQE_VI = "100%" if IQE_VI >= 1 else f"{IQE_VI*100:.1f}%".replace('.', ',').rstrip('0').rstrip(',')

    FONTE_NREALIZ_IQE = TAB_INFO ["INFORMAÇÃO"][23]
    FONTE_PLANO_IQE = TAB_INFO ["INFORMAÇÃO"][24]
    FONT_NCONF_IQE = TAB_INFO ["INFORMAÇÃO"][25]
    ACREDITAÇÃO_IQE = TAB_INFO ["INFORMAÇÃO"][26]

    return (MES_ANO, TRIM_ANO, TRIMMIN_ANO, ANO_CONTRAT, PRESTADORA, META, NCOF_P, NREA_P, IQA_P, NCOF_VI, NREA_VI, IQA_VI, FONTE_NREALIZ, FONTE_PLANO, FONT_NCONF, ACREDITAÇÃO, META_IQE, NCOF_P_IQE, NREA_P_IQE, IQE_P, NCOF_VI_IQE, NREA_VI_IQE, IQE_VI, FONTE_NREALIZ_IQE, FONTE_PLANO_IQE, FONT_NCONF_IQE, ACREDITAÇÃO_IQE, MESMIN_ANO)



# DEFINIÇÃO DO CONJUNTO DE VARIÁVEIS A SEREM INSERIDAS NO RELATÓRIO

  def var_doc (MES_ANO, TRIM_ANO, TRIMMIN_ANO, ANO_CONTRAT, PRESTADORA, META, NCOF_P, NREA_P, IQA_P, NCOF_VI, NREA_VI, IQA_VI, FONTE_NREALIZ, FONTE_PLANO, FONT_NCONF, ACREDITAÇÃO, TAB_MUN_FINAL, TAB_MUN_FINAL_IQE, QUADRO_POND_NCONF, QUADRO_POND_REALIZ, TAB_MUN_ARSAL, TAB_MUN_VI, TAB_TIPO_FINAL, TAB_TIPO_RESUMIDA_VI, MESMIN_ANO, total_conf_vi_naoconf_arsal, total_naoconf_vi_conf_arsal, desc_plano, TAB_MUN_NCONF, TAB_MUN_NCONF_VI, META_IQE, NCOF_P_IQE, NREA_P_IQE, IQE_P, NCOF_VI_IQE, NREA_VI_IQE, IQE_VI, FONTE_NREALIZ_IQE, FONTE_PLANO_IQE, FONT_NCONF_IQE, ACREDITA_IQE, QUADRO_POND_REALIZ_IQE, QUADRO_POND_NCONF_IQE, TAB_MUN_VI_IQE, TAB_MUN_ARSAL_IQE, TAB_TIPO_RESUMIDA_VI_IQE, TAB_TIPO_FINAL_IQE, total_conf_vi_naoconf_arsal_IQE, total_naoconf_vi_conf_arsal_IQE, desc_plano_IQE, TAB_MUN_NCONF_IQE, TAB_MUN_NCONF_VI_IQE):

    NREA_ARSAL = (TAB_MUN_FINAL["NAM REALIZ"].iloc[-1]).astype(int)
    NREA_EXPURGOS = (TAB_MUN_FINAL["EXPURGOS TOTAIS"].iloc[-1]).astype(int)
    NCONF_ARSAL = (TAB_MUN_FINAL["NAM CONF"].iloc[-1]).astype(int)
    NCONF_EXPURGOS = (TAB_MUN_FINAL["EXPURGOS CONFORMES"].iloc[-1]).astype(int)
    IQA_ARSAL = (NCONF_ARSAL/NREA_ARSAL)*100
    IQA_ARSAL_TXT = "100%" if IQA_ARSAL == 100 else f"{IQA_ARSAL:.1f}%".replace('.', ',')

    if (IQA_VI).strip() == (IQA_ARSAL_TXT).strip():
      IS_SATISF = "satisfatório"

    if (IQA_VI).strip() != (IQA_ARSAL_TXT).strip():
      IS_SATISF = "divergente"

    if PRESTADORA == "BRK Ambiental":
      legis = '5'

    if PRESTADORA != "BRK Ambiental":
      legis = '7'

    def validar(valor):
      return valor if valor is not None else ""

    def formata_milhar(x):
      try:
          return f"{int(float(x)):,}".replace(",", ".")
      except:
          return "-"

    NREA_ARSAL_IQE = (TAB_MUN_FINAL_IQE["NAM REALIZ"].iloc[-1]).astype(int)
    NREA_EXPURGOS_IQE = (TAB_MUN_FINAL_IQE["EXPURGOS TOTAIS"].iloc[-1]).astype(int)
    NCONF_ARSAL_IQE = (TAB_MUN_FINAL_IQE["NAM CONF"].iloc[-1]).astype(int)
    NCONF_EXPURGOS_IQE = (TAB_MUN_FINAL_IQE["EXPURGOS CONFORMES"].iloc[-1]).astype(int)
    IQE_ARSAL = (NCONF_ARSAL_IQE/NREA_ARSAL_IQE)*100
    IQE_ARSAL_TXT = "100%" if IQE_ARSAL == 100 else f"{IQE_ARSAL:.1f}%".replace('.', ',')

    print ("####################################################")

    print(type(IQE_ARSAL_TXT))
    print(type(IQE_VI))

    print(IQE_ARSAL_TXT)
    print(IQE_VI)

    if (IQE_VI).strip() == (IQE_ARSAL_TXT).strip():
      IS_SATISF_IQE = "satisfatório"

    if (IQE_VI).strip() != (IQE_ARSAL_TXT).strip():
      IS_SATISF_IQE = "divergente"

    if PRESTADORA == "BRK Ambiental":
      legis = '5'

    if PRESTADORA != "BRK Ambiental":
      legis = '7'

    def validar(valor):
      return valor if valor is not None else ""

    var_dic = {
              '$MES_ANO': str(MES_ANO),
              '$ANO_CONTRAT': str(ANO_CONTRAT),
              '$MESMIN_ANO': str(MESMIN_ANO),
              '$PRESTADORA': str(PRESTADORA),
              '$TRIMMIN_ANO': str(TRIMMIN_ANO),
              '$TRIM_ANO': str(TRIM_ANO),


              '$META':str(META),
              '$NCOF_P': formata_milhar(NCOF_P),
              '$NREA_P': formata_milhar(NREA_P),
              '$IQA_P': str(IQA_P),
              '$NCOF_VI': formata_milhar(NCOF_VI),
              '$NREA_VI': formata_milhar(NREA_VI),
              '$IQA_VI': str(IQA_VI),
              '$FONTE_NREALIZ': str(FONTE_NREALIZ),
              '$FONTE_PLANO': str(FONTE_PLANO),
              '$FONTE_NCONF': str(FONT_NCONF),
              '$ACREDITAÇÃO': str (ACREDITAÇÃO),
              '$QUADRO_NREA_OBS': QUADRO_POND_REALIZ,
              '$QUADRO_NCONF_OBS': QUADRO_POND_NCONF,
              '$QUADRO_NREA_VI': TAB_MUN_VI,
              '$QUADRO_NREA_ARSAL': TAB_MUN_ARSAL,
              '$QUADRO_NCONF_VI': TAB_TIPO_RESUMIDA_VI,
              '$QUADRO_NCONF_ARSAL': TAB_TIPO_FINAL,
              #'$QUADRO_NCONFMUN_VI': TAB_MUN_NCONF_VI,
              #'$QUADRO_NCONFMUN_ARSAL': TAB_MUN_NCONF,
              '$NREA_ARSAL': formata_milhar(NREA_ARSAL),
              '$NREA_EXPURGOS': str (NREA_EXPURGOS),
              '$NCOF_ARSAL': formata_milhar(NCONF_ARSAL),
              '$NCONF_EXPURGOS': str (NCONF_EXPURGOS),
              '$IS_SATISFATORIO': IS_SATISF,
              '$IQA_ARSAL': IQA_ARSAL_TXT,
              '$LEGIS' : legis,
              '$CONF_NCONF': str(total_conf_vi_naoconf_arsal),
              '$NCONF_CONF' : str(total_naoconf_vi_conf_arsal),
              '$NREA__PLAN' : str(desc_plano),


              '$MEETA_IQE':str(META_IQE),
              '$NCOF__P_IQE': formata_milhar(NCOF_P_IQE),
              '$NREA__P_IQE': formata_milhar(NREA_P_IQE),
              '$IQE_P': str(IQE_P),
              '$NCOF__VI_IQE': formata_milhar(NCOF_VI_IQE),
              '$NREA__VI_IQE': formata_milhar(NREA_VI_IQE),
              '$IQE_VI': str(IQE_VI),
              '$FONTE__NREALIZ_IQE': str(FONTE_NREALIZ_IQE),
              '$FONTE__PLANO_IQE': str(FONTE_PLANO_IQE),
              '$FONTE__NCONF_IQE': str(FONT_NCONF_IQE),
              '$ACREDITA_IQE': str (ACREDITA_IQE),
              '$QUADRO__NREA_OBS_IQE': QUADRO_POND_REALIZ_IQE,
              '$QUADRO__NCONF_OBS_IQE': QUADRO_POND_NCONF_IQE,
              '$QUADRO__NREA_VI_IQE': TAB_MUN_VI_IQE,
              '$QUADRO__NREA_ARSAL_IQE': TAB_MUN_ARSAL_IQE,
              '$QUADRO__NCONF_VI_IQE': TAB_TIPO_RESUMIDA_VI_IQE,
              '$QUADRO__NCONF_ARSAL_IQE': TAB_TIPO_FINAL_IQE,
              # '$QUADRO_NCONFMUN_VI': TAB_MUN_NCONF_VI,
              # '$QUADRO_NCONFMUN_ARSAL': TAB_MUN_NCONF,
              '$NREA__ARSAL_IQE': formata_milhar(NREA_ARSAL_IQE),
              '$NREA__EXPURGOS_IQE': str (NREA_EXPURGOS_IQE),
              '$NCOF__ARSAL_IQE': formata_milhar(NCONF_ARSAL_IQE),
              '$NCONF__EXPURGOS_IQE': str (NCONF_EXPURGOS_IQE),
              '$IS__SATISFATORIO_IQE': IS_SATISF_IQE,
              '$IQE_ARSAL': IQE_ARSAL_TXT,
              '$CONF__NCONF_IQE': str(total_conf_vi_naoconf_arsal_IQE),
              '$NCONF__CONF_IQE' : str(total_naoconf_vi_conf_arsal_IQE),
              '$NREAL__PLAN_IQE' : str(desc_plano_IQE)

              }


    print (IS_SATISF_IQE)

    return var_dic



  # CONVERSÃO DE DATAFRAME PARA TABELA EDITÁVEL NO RELATÓRIO

  def df_para_tabela_word(doc, df, cor_cabecalho="BFBFBF", cor_par="FFFFFF", cor_impar="FFFFFF"):
    
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    def set_bg(cell, cor):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), cor)
        tcPr.append(shd)

    def set_borders(cell, cor="000000", tamanho="4"):
      tcPr = cell._tc.get_or_add_tcPr()
      tcBorders = OxmlElement("w:tcBorders")
      for lado in ("top", "left", "bottom", "right"):
          border = OxmlElement(f"w:{lado}")
          border.set(qn("w:val"),   "single")
          border.set(qn("w:sz"),    tamanho)
          border.set(qn("w:space"), "0")
          border.set(qn("w:color"), cor)
          tcBorders.append(border)
      tcPr.append(tcBorders)

    def set_col_widths(tabela, larguras_cm):
        
        tbl = tabela._tbl
        tblGrid = tbl.find(qn("w:tblGrid"))
        if tblGrid is None:
            tblGrid = OxmlElement("w:tblGrid")
            tbl.insert(0, tblGrid)
        else:
            for col in tblGrid.findall(qn("w:gridCol")):
                tblGrid.remove(col)

        for largura in larguras_cm:
            gridCol = OxmlElement("w:gridCol")
            gridCol.set(qn("w:w"), str(int(largura / 635)))
            tblGrid.append(gridCol)

        for row in tabela.rows:
            for i, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.append(tcW)
                tcW.set(qn("w:w"),    str(int(larguras_cm[i] / 635)))
                tcW.set(qn("w:type"), "dxa")

    tabela = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    tabela.style = "Normal Table"
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

    num_cols = len(df.columns)
    larguras = []

    MUNICIPIO_ANALISE = Cm(5.5)
    EX_LARGA          = Cm(7.0)
    LARGA             = Cm(3.2)
    MEDIA             = Cm(2.5)
    ESTREITA          = Cm(1.6)
    CONVERSAO         = Cm(3.0)

    for c in range(num_cols):
        col_name = df.columns[c]

        if c == 0:

            if col_name in ("MUNICÍPIO", "ANÁLISE"):
                larguras.append(MUNICIPIO_ANALISE)
            else:
                larguras.append(ESTREITA)
        else:

            if col_name in ("CONF VI → NÃO CONF ARSAL", "NÃO CONF VI → CONF ARSAL",
                            "CONVERSÃO: CONFORME → NÃO CONFORME", "CONVERSÃO: NÃO CONFORME → CONFORME"):
                larguras.append(CONVERSAO)
  
            elif col_name in ("PONDERAÇÕES (PRESTADORA)", "PONDERAÇÕES (VI)", "PONDERAÇÕES (ARSAL)"):
                larguras.append(EX_LARGA)
    
            elif col_name in ("PLANO DE AMOSTRAGEM", "PLANO DE AMOSTRAGEM AJUSTADO",
                              "PLANO DE AMOSTRAGEM AJUSTADO - ARSAL", "DESCONSIDERAÇÕES DO PLANO",
                              "ANALISES REALIZADAS", "ANALISES CONFORMES", "ANALISES NÃO CONFORMES",
                              "EXPURGOS TOTAIS", "EXPURGOS CONFORMES", "EXPURGOS NAO CONFORMES",
                              "NAM REALIZ", "NAM REALIZ (VI)", "NAM CONF", "NAM CONF (VI)"):
                larguras.append(MEDIA)
    
            else:
                larguras.append(ESTREITA)


    set_col_widths(tabela, larguras)

    for c, col in enumerate(df.columns):
        cell = tabela.rows[0].cells[c]

        set_bg(cell, cor_cabecalho)
        set_borders(cell)

        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p = cell.paragraphs[0]
        p.clear()

        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

        run = p.add_run(str(col))
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.name = "Times New Roman"
        run.font.size = Pt(9.0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r, (_, linha) in enumerate(df.iterrows()):
        cor = cor_par if r % 2 == 0 else cor_impar
        is_ultima_linha = (df.columns[-1] in ("NAM REALIZ", "NAM REALIZ (VI)", "NAM CONF", "NAM CONF (VI)")) and (r == len(df) - 1)
        for c, valor in enumerate(linha):
            cell = tabela.rows[r + 1].cells[c]
            set_bg(cell, cor)
            set_borders(cell)

            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.clear()

            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

            run = p.add_run(str(valor) if valor is not None else "")

            run.bold = is_ultima_linha

            run.font.name = "Times New Roman"
            run.font.size = Pt(9.0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.element.body.remove(tabela._tbl)
    return tabela._tbl



# ALIMENTANDO INFORMAÇÕES AO MODELO DE RELATÓRIO

  def substituir_var(documento, var_dic):
    doc = Document(documento)

    def process_paragraph(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)

        if any(key in full_text for key in var_dic.keys()):
            for key, value in var_dic.items():
                if key in full_text:
                    if isinstance(value, pd.DataFrame):
                        
                        tbl = df_para_tabela_word(doc, value)
                        paragraph._element.addprevious(tbl)

                        for run in paragraph.runs:
                            run.text = ""

                        return
                    else:
                        full_text = full_text.replace(key, value)

            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = full_text

    for p in doc.paragraphs:
        process_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)

    def replace_placeholder_with_image(doc, placeholder, image_path, width_inches):
      for paragraph in doc.paragraphs:
          if placeholder in paragraph.text:
              
              paragraph.clear()
              
              paragraph.add_run().add_picture(image_path, width=Inches(width_inches))

    placeholders = [

      ('$GRAF_COMPARATIVO_IQAS', 'IQAS.png', 6.5),
      ('$GRAF_CONF_REALIZ_TEMPO', 'nc_nr_IQAS.png', 6.5),
      ('$GRAF_PLANO_REALIZ_ANALISE', 'grafico_plan_realiz_tipo.png', 6.5),
      ('$GRAF_PLANO_REALIZ_MUN', 'grafico_plan_realiz_mun.png', 9.0),
      ('$GRAF_CONF_REALIZ_ANALISE', 'grafico_conf_realiz_tipo.png', 6.5),
      ('$GRAF_CONF_REALIZ_MUN', 'grafico_conf_realiz_mun.png', 9.5),
      ('$GRAF_IQAS_MUN', 'grafico_iqa_mun.png', 6.5),

      ('$GRAF_COMPARATIVO_IQES', 'IQES.png', 6.5),
      ('$GRAF__CONF_REALIZ_TEMPO_IQE', 'nc_nr_IQES.png', 6.5),
      ('$GRAF__PLANO_REALIZ_ANALISE_IQE', 'grafico_plan_realiz_tipo_IQE.png', 6.5),
      ('$GRAF__PLANO_REALIZ_MUN_IQE', 'grafico_plan_realiz_mun_IQE.png', 9.0),
      ('$GRAF__CONF_REALIZ_ANALISE_IQE', 'grafico_conf_realiz_tipo_IQE.png', 6.5),
      ('$GRAF__CONF_REALIZ_MUN_IQE', 'grafico_conf_realiz_mun_IQE.png', 9.5),
      ('$GRAF_IQES_MUN', 'grafico_iqa_mun_IQE.png', 6.5)
    ]

    for ph, img, width in placeholders:
      replace_placeholder_with_image(doc, ph, img, width)

    return doc

  

  # ANÁLISE POR MUNICÍPIO

  def analise_mun(TAB_MUN, TAB_REALIZADAS, BASE_MUN_TIPO_completo, TAB_MUN_IQE, TAB_REALIZADAS_IQE, BASE_MUN_TIPO_IQE_completo):

    # IQA

    BASE_MUN_TIPO_completo["Cidade"] = BASE_MUN_TIPO_completo["Cidade"].str.upper().apply(unidecode)

    DESC_MUN = (
      BASE_MUN_TIPO_completo
      .groupby("Cidade", as_index=False)["DESCONSIDERAÇÕES - ARSAL"]
      .sum()
      .rename(columns={
        "Cidade": "MUNICÍPIO",
        "DESCONSIDERAÇÕES - ARSAL": "DESCONSIDERAÇÕES DO PLANO"
      })
    )

    TAB_MUN = TAB_MUN.merge(DESC_MUN, on="MUNICÍPIO", how="left")
    TAB_MUN["DESCONSIDERAÇÕES DO PLANO"] = TAB_MUN["DESCONSIDERAÇÕES DO PLANO"].fillna(0).astype(int)
    TAB_MUN["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"] = (TAB_MUN["PLANO DE AMOSTRAGEM"] - TAB_MUN["DESCONSIDERAÇÕES DO PLANO"]).astype(int)

    TAB_REALIZADAS_ACEITO          = TAB_REALIZADAS[TAB_REALIZADAS["ACEITO?"] == "Validado"]
    TAB_REALIZADAS_ACEITO_CONF     = TAB_REALIZADAS_ACEITO[TAB_REALIZADAS["CONFORMIDADE - ARSAL"] == "Conforme"]
    TAB_REALIZADAS_ACEITO_CONF_VI     = TAB_REALIZADAS_ACEITO[TAB_REALIZADAS["CONFORMIDADE - VI"] == "Conforme"]
    TAB_REALIZADAS_ACEITO_NAOCONF  = TAB_REALIZADAS_ACEITO[(TAB_REALIZADAS["CONFORMIDADE - ARSAL"] != "Conforme") & (TAB_REALIZADAS["CONFORMIDADE - ARSAL"].notna())]
    TAB_REALIZADAS_ACEITO_NAOCONF_VI  = TAB_REALIZADAS_ACEITO[(TAB_REALIZADAS["CONFORMIDADE - VI"] != "Conforme") & (TAB_REALIZADAS["CONFORMIDADE - VI"].notna())]
    TAB_REALIZADAS_ACEITO_CONF_EXP    = TAB_REALIZADAS_ACEITO_CONF[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_CONF_EXP_VI    = TAB_REALIZADAS_ACEITO_CONF_VI[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_NAOCONF_EXP = TAB_REALIZADAS_ACEITO_NAOCONF[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI = TAB_REALIZADAS_ACEITO_NAOCONF_VI[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]

    CONF_VI_NAOCONF_ARSAL = TAB_REALIZADAS_ACEITO[
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - VI"] == "Conforme") &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - ARSAL"] != "Conforme") &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - ARSAL"].notna()) &
        (TAB_REALIZADAS_ACEITO["EXPURGOS - ARSAL"] != "EXPURGAR")
    ]

    NAOCONF_VI_CONF_ARSAL = TAB_REALIZADAS_ACEITO[
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - VI"] != "Conforme") &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - VI"].notna()) &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - ARSAL"] == "Conforme") &
        (TAB_REALIZADAS_ACEITO["EXPURGOS - ARSAL"] != "EXPURGAR")
    ]

    contagem_conf_vi_naoconf_arsal = (
        CONF_VI_NAOCONF_ARSAL.groupby("CIDADE").size()
        .reset_index(name="CONF VI → NÃO CONF ARSAL")
    )

    contagem_naoconf_vi_conf_arsal = (
        NAOCONF_VI_CONF_ARSAL.groupby("CIDADE").size()
        .reset_index(name="NÃO CONF VI → CONF ARSAL")
    )

    TAB_MUN = TAB_MUN.merge(contagem_conf_vi_naoconf_arsal, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN = TAB_MUN.drop(columns='CIDADE')
    TAB_MUN["CONF VI → NÃO CONF ARSAL"] = TAB_MUN["CONF VI → NÃO CONF ARSAL"].fillna(0).astype(int)

    TAB_MUN = TAB_MUN.merge(contagem_naoconf_vi_conf_arsal, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN = TAB_MUN.drop(columns='CIDADE')
    TAB_MUN["NÃO CONF VI → CONF ARSAL"] = TAB_MUN["NÃO CONF VI → CONF ARSAL"].fillna(0).astype(int)

    contagem_tot = TAB_REALIZADAS_ACEITO.groupby('CIDADE').size().reset_index(name='ANALISES REALIZADAS')
    TAB_MUN = TAB_MUN.merge(contagem_tot, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN = TAB_MUN.drop(columns='CIDADE')

    contagem_conf = TAB_REALIZADAS_ACEITO_CONF_VI.groupby('CIDADE').size().reset_index(name='ANALISES CONFORMES')
    TAB_MUN = TAB_MUN.merge(contagem_conf, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN = TAB_MUN.drop(columns='CIDADE')

    contagem_naoconf = TAB_REALIZADAS_ACEITO_NAOCONF_VI.groupby('CIDADE').size().reset_index(name='ANALISES NÃO CONFORMES')
    TAB_MUN = TAB_MUN.merge(contagem_naoconf, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN = TAB_MUN.drop(columns='CIDADE')
    TAB_MUN['ANALISES NÃO CONFORMES'] = TAB_MUN['ANALISES NÃO CONFORMES'].fillna(0).astype(int)

    contagem_conf_exp = TAB_REALIZADAS_ACEITO_CONF_EXP_VI.groupby('CIDADE').size().reset_index(name='EXPURGOS CONFORMES')
    TAB_MUN = TAB_MUN.merge(contagem_conf_exp, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN = TAB_MUN.drop(columns='CIDADE')
    TAB_MUN['EXPURGOS CONFORMES'] = TAB_MUN['EXPURGOS CONFORMES'].fillna(0).astype(int)

    contagem_naoconf_exp = TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI.groupby('CIDADE').size().reset_index(name='EXPURGOS NAO CONFORMES')
    TAB_MUN = TAB_MUN.merge(contagem_naoconf_exp, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN = TAB_MUN.drop(columns='CIDADE')
    TAB_MUN['EXPURGOS NAO CONFORMES'] = TAB_MUN['EXPURGOS NAO CONFORMES'].fillna(0).astype(int)

    TAB_MUN['EXPURGOS TOTAIS'] = TAB_MUN['EXPURGOS NAO CONFORMES'] + TAB_MUN['EXPURGOS CONFORMES']

    TAB_MUN["NAM REALIZ"] = (np.maximum(TAB_MUN['PLANO DE AMOSTRAGEM AJUSTADO - ARSAL'].fillna(0), TAB_MUN['ANALISES REALIZADAS'].fillna(0)) - TAB_MUN['EXPURGOS TOTAIS'].fillna(0)).astype(int)

    TAB_MUN["NAM CONF"] = (TAB_MUN['ANALISES CONFORMES'].fillna(0) - TAB_MUN['EXPURGOS CONFORMES'].fillna(0) + TAB_MUN['NÃO CONF VI → CONF ARSAL'].fillna(0) - TAB_MUN['CONF VI → NÃO CONF ARSAL'].fillna(0)).astype(int)

    TAB_MUN["IQA (%)"] = (((TAB_MUN["NAM CONF"]) / (TAB_MUN["NAM REALIZ"])) * 100).round(1)

    TAB_MUN_FINAL = TAB_MUN[[
      "MUNICÍPIO", "PLANO DE AMOSTRAGEM", "DESCONSIDERAÇÕES DO PLANO",
      "PLANO DE AMOSTRAGEM AJUSTADO - ARSAL",
      "ANALISES REALIZADAS", "EXPURGOS TOTAIS", "NAM REALIZ",
      "ANALISES CONFORMES", "EXPURGOS CONFORMES", "CONF VI → NÃO CONF ARSAL", "NÃO CONF VI → CONF ARSAL", "NAM CONF", "IQA (%)"
    ]]
    TAB_MUN_FINAL.loc['Total'] = TAB_MUN_FINAL.sum()
    TAB_MUN_FINAL["MUNICÍPIO"].iat[-1] = ' '
    TAB_MUN_FINAL["IQA (%)"].iat[-1] = (((TAB_MUN_FINAL["NAM CONF"].iat[-1]) / (TAB_MUN_FINAL["NAM REALIZ"].iat[-1])) * 100).round(1)

    desc_plano = int(TAB_MUN_FINAL["DESCONSIDERAÇÕES DO PLANO"].iloc[-1])

    TAB_MUN_NCONF = TAB_MUN[["MUNICÍPIO", "ANALISES REALIZADAS", "ANALISES CONFORMES", "EXPURGOS CONFORMES", "CONF VI → NÃO CONF ARSAL", "NÃO CONF VI → CONF ARSAL", "NAM CONF"]]
    TAB_MUN_NCONF_VI = TAB_MUN[["MUNICÍPIO", "ANALISES REALIZADAS", "ANALISES CONFORMES"]]

    def add_total(df):
      total = df.select_dtypes(include='number').sum()
      total['MUNICÍPIO'] = 'TOTAL'
      return pd.concat([df, total.to_frame().T], ignore_index=True)

    TAB_MUN_NCONF = add_total(TAB_MUN_NCONF)
    TAB_MUN_NCONF_VI = add_total(TAB_MUN_NCONF_VI)

    TAB_MUN_RESUMIDA = TAB_MUN_FINAL[[
      "MUNICÍPIO", "PLANO DE AMOSTRAGEM", "DESCONSIDERAÇÕES DO PLANO",
      "PLANO DE AMOSTRAGEM AJUSTADO - ARSAL",
      "ANALISES REALIZADAS", "EXPURGOS TOTAIS", "NAM REALIZ"
    ]]
    TAB_MUN_RESUMIDA = TAB_MUN_RESUMIDA.rename(columns={"PLANO DE AMOSTRAGEM": "PLANO DE AMOSTRAGEM AJUSTADO"})
    TAB_MUN_RESUMIDA["PLANO DE AMOSTRAGEM AJUSTADO"] = TAB_MUN_RESUMIDA["PLANO DE AMOSTRAGEM AJUSTADO"].astype(int)
    TAB_MUN_RESUMIDA.loc[TAB_MUN_RESUMIDA.index[-1], "MUNICÍPIO"] = "TOTAL"

    TAB_MUN_RESUMIDA_VI = TAB_MUN_RESUMIDA[[
      "MUNICÍPIO", "PLANO DE AMOSTRAGEM AJUSTADO", "ANALISES REALIZADAS"
    ]]

    TAB_MUN_RESUMIDA_VI = TAB_MUN_RESUMIDA_VI.fillna(0)

    TAB_MUN_RESUMIDA_VI["NAM REALIZ (VI)"] = (np.maximum(
      TAB_MUN_RESUMIDA_VI["PLANO DE AMOSTRAGEM AJUSTADO"],
      TAB_MUN_RESUMIDA_VI['ANALISES REALIZADAS']
    )).astype(int)
    TAB_MUN_RESUMIDA_VI = TAB_MUN_RESUMIDA_VI[:-1]
    TAB_MUN_RESUMIDA_VI.loc['Total'] = TAB_MUN_RESUMIDA_VI.sum()
    TAB_MUN_RESUMIDA_VI.iloc[-1, TAB_MUN_RESUMIDA_VI.columns.get_loc("MUNICÍPIO")] = "TOTAL"

    

    # IQE

    BASE_MUN_TIPO_IQE_completo["Cidade"] = BASE_MUN_TIPO_IQE_completo["Cidade"].str.upper().apply(unidecode)

    DESC_MUN = (
      BASE_MUN_TIPO_IQE_completo
      .groupby("Cidade", as_index=False)["DESCONSIDERAÇÕES - ARSAL"]
      .sum()
      .rename(columns={
        "Cidade": "MUNICÍPIO",
        "DESCONSIDERAÇÕES - ARSAL": "DESCONSIDERAÇÕES DO PLANO"
      })
    )

    TAB_MUN_IQE = TAB_MUN_IQE.merge(DESC_MUN, on="MUNICÍPIO", how="left")
    TAB_MUN_IQE["DESCONSIDERAÇÕES DO PLANO"] = TAB_MUN_IQE["DESCONSIDERAÇÕES DO PLANO"].fillna(0).astype(int)
    TAB_MUN_IQE["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"] = (TAB_MUN_IQE["PLANO DE AMOSTRAGEM"] - TAB_MUN_IQE["DESCONSIDERAÇÕES DO PLANO"]).astype(int)

    TAB_REALIZADAS_ACEITO_IQE          = TAB_REALIZADAS_IQE[TAB_REALIZADAS_IQE["ACEITO?"] == "Validado"]
    TAB_REALIZADAS_ACEITO_CONF_IQE     = TAB_REALIZADAS_ACEITO_IQE[TAB_REALIZADAS_IQE["CONFORMIDADE - ARSAL"] == "Conforme"]
    TAB_REALIZADAS_ACEITO_CONF_VI_IQE     = TAB_REALIZADAS_ACEITO_IQE[TAB_REALIZADAS_IQE["CONFORMIDADE - VI"] == "Conforme"]
    TAB_REALIZADAS_ACEITO_NAOCONF_IQE  = TAB_REALIZADAS_ACEITO_IQE[(TAB_REALIZADAS_IQE["CONFORMIDADE - ARSAL"] != "Conforme") & (TAB_REALIZADAS_IQE["CONFORMIDADE - ARSAL"].notna())]
    TAB_REALIZADAS_ACEITO_NAOCONF_VI_IQE  = TAB_REALIZADAS_ACEITO_IQE[(TAB_REALIZADAS_IQE["CONFORMIDADE - VI"] != "Conforme") & (TAB_REALIZADAS_IQE["CONFORMIDADE - VI"].notna())]
    TAB_REALIZADAS_ACEITO_CONF_EXP_IQE    = TAB_REALIZADAS_ACEITO_CONF_IQE[TAB_REALIZADAS_IQE["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_CONF_EXP_VI_IQE    = TAB_REALIZADAS_ACEITO_CONF_VI_IQE[TAB_REALIZADAS_IQE["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_NAOCONF_EXP_IQE = TAB_REALIZADAS_ACEITO_NAOCONF_IQE[TAB_REALIZADAS_IQE["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI_IQE = TAB_REALIZADAS_ACEITO_NAOCONF_VI_IQE[TAB_REALIZADAS_IQE["EXPURGOS - ARSAL"] == "EXPURGAR"]

    CONF_VI_NAOCONF_ARSAL_IQE = TAB_REALIZADAS_ACEITO_IQE[
        (TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - VI"] == "Conforme") &
        (TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - ARSAL"] != "Conforme") &
        (TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - ARSAL"].notna()) &
        (TAB_REALIZADAS_ACEITO_IQE["EXPURGOS - ARSAL"] != "EXPURGAR")
    ]

    NAOCONF_VI_CONF_ARSAL_IQE = TAB_REALIZADAS_ACEITO_IQE[
        (TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - VI"] != "Conforme") &
        (TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - VI"].notna()) &
        (TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - ARSAL"] == "Conforme") &
        (TAB_REALIZADAS_ACEITO_IQE["EXPURGOS - ARSAL"] != "EXPURGAR")
    ]

    contagem_conf_vi_naoconf_arsal_IQE = (
        CONF_VI_NAOCONF_ARSAL_IQE.groupby("CIDADE").size()
        .reset_index(name="CONF VI → NÃO CONF ARSAL")
    )

    contagem_naoconf_vi_conf_arsal_IQE = (
        NAOCONF_VI_CONF_ARSAL_IQE.groupby("CIDADE").size()
        .reset_index(name="NÃO CONF VI → CONF ARSAL")
    )

    TAB_MUN_IQE = TAB_MUN_IQE.merge(contagem_conf_vi_naoconf_arsal_IQE, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN_IQE = TAB_MUN_IQE.drop(columns='CIDADE')
    TAB_MUN_IQE["CONF VI → NÃO CONF ARSAL"] = TAB_MUN_IQE["CONF VI → NÃO CONF ARSAL"].fillna(0).astype(int)

    TAB_MUN_IQE = TAB_MUN_IQE.merge(contagem_naoconf_vi_conf_arsal_IQE, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN_IQE = TAB_MUN_IQE.drop(columns='CIDADE')
    TAB_MUN_IQE["NÃO CONF VI → CONF ARSAL"] = TAB_MUN_IQE["NÃO CONF VI → CONF ARSAL"].fillna(0).astype(int)

    contagem_tot_IQE = TAB_REALIZADAS_ACEITO_IQE.groupby('CIDADE').size().reset_index(name='ANALISES REALIZADAS')
    TAB_MUN_IQE = TAB_MUN_IQE.merge(contagem_tot_IQE, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN_IQE = TAB_MUN_IQE.drop(columns='CIDADE')

    contagem_conf_IQE = TAB_REALIZADAS_ACEITO_CONF_VI_IQE.groupby('CIDADE').size().reset_index(name='ANALISES CONFORMES')
    TAB_MUN_IQE = TAB_MUN_IQE.merge(contagem_conf_IQE, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN_IQE = TAB_MUN_IQE.drop(columns='CIDADE')

    contagem_naoconf_IQE = TAB_REALIZADAS_ACEITO_NAOCONF_VI_IQE.groupby('CIDADE').size().reset_index(name='ANALISES NÃO CONFORMES')
    TAB_MUN_IQE = TAB_MUN_IQE.merge(contagem_naoconf_IQE, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN_IQE = TAB_MUN_IQE.drop(columns='CIDADE')
    TAB_MUN_IQE['ANALISES NÃO CONFORMES'] = TAB_MUN_IQE['ANALISES NÃO CONFORMES'].fillna(0).astype(int)

    contagem_conf_exp_IQE = TAB_REALIZADAS_ACEITO_CONF_EXP_VI_IQE.groupby('CIDADE').size().reset_index(name='EXPURGOS CONFORMES')
    TAB_MUN_IQE = TAB_MUN_IQE.merge(contagem_conf_exp_IQE, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN_IQE = TAB_MUN_IQE.drop(columns='CIDADE')
    TAB_MUN_IQE['EXPURGOS CONFORMES'] = TAB_MUN_IQE['EXPURGOS CONFORMES'].fillna(0).astype(int)

    contagem_naoconf_exp_IQE = TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI_IQE.groupby('CIDADE').size().reset_index(name='EXPURGOS NAO CONFORMES')
    TAB_MUN_IQE = TAB_MUN_IQE.merge(contagem_naoconf_exp_IQE, left_on='MUNICÍPIO', right_on='CIDADE', how='left')
    TAB_MUN_IQE = TAB_MUN_IQE.drop(columns='CIDADE')
    TAB_MUN_IQE['EXPURGOS NAO CONFORMES'] = TAB_MUN_IQE['EXPURGOS NAO CONFORMES'].fillna(0).astype(int)

    TAB_MUN_IQE['EXPURGOS TOTAIS'] = TAB_MUN_IQE['EXPURGOS NAO CONFORMES'] + TAB_MUN_IQE['EXPURGOS CONFORMES']

    TAB_MUN_IQE["NAM REALIZ"] = (np.maximum(TAB_MUN_IQE['PLANO DE AMOSTRAGEM AJUSTADO - ARSAL'], TAB_MUN_IQE['ANALISES REALIZADAS']) - TAB_MUN_IQE['EXPURGOS TOTAIS']).astype(int)

    TAB_MUN_IQE = TAB_MUN_IQE.fillna(0)

    TAB_MUN_IQE["NAM CONF"] = (TAB_MUN_IQE['ANALISES CONFORMES'] - TAB_MUN_IQE['EXPURGOS CONFORMES'] + TAB_MUN_IQE['NÃO CONF VI → CONF ARSAL'] - TAB_MUN_IQE['CONF VI → NÃO CONF ARSAL']).astype(int)

    TAB_MUN_IQE["IQE (%)"] = (((TAB_MUN_IQE["NAM CONF"]) / (TAB_MUN_IQE["NAM REALIZ"])) * 100).round(1)

    TAB_MUN_FINAL_IQE = TAB_MUN_IQE[[
      "MUNICÍPIO", "PLANO DE AMOSTRAGEM", "DESCONSIDERAÇÕES DO PLANO",
      "PLANO DE AMOSTRAGEM AJUSTADO - ARSAL",
      "ANALISES REALIZADAS", "EXPURGOS TOTAIS", "NAM REALIZ",
      "ANALISES CONFORMES", "EXPURGOS CONFORMES", "CONF VI → NÃO CONF ARSAL", "NÃO CONF VI → CONF ARSAL", "NAM CONF", "IQE (%)"
    ]]
    TAB_MUN_FINAL_IQE.loc['Total'] = TAB_MUN_FINAL_IQE.sum()
    TAB_MUN_FINAL_IQE["MUNICÍPIO"].iat[-1] = ' '
    TAB_MUN_FINAL_IQE["IQE (%)"].iat[-1] = (((TAB_MUN_FINAL_IQE["NAM CONF"].iat[-1]) / (TAB_MUN_FINAL_IQE["NAM REALIZ"].iat[-1])) * 100).round(1)

    desc_plano_IQE = int(TAB_MUN_FINAL_IQE["DESCONSIDERAÇÕES DO PLANO"].iloc[-1])

    TAB_MUN_NCONF_IQE = TAB_MUN_IQE[["MUNICÍPIO", "ANALISES REALIZADAS", "ANALISES CONFORMES", "EXPURGOS CONFORMES", "CONF VI → NÃO CONF ARSAL", "NÃO CONF VI → CONF ARSAL", "NAM CONF"]]
    TAB_MUN_NCONF_VI_IQE = TAB_MUN_IQE[["MUNICÍPIO", "ANALISES REALIZADAS", "ANALISES CONFORMES"]]

    def add_total(df):
      total = df.select_dtypes(include='number').sum()
      total['MUNICÍPIO'] = 'TOTAL'
      return pd.concat([df, total.to_frame().T], ignore_index=True)

    TAB_MUN_NCONF_IQE = add_total(TAB_MUN_NCONF_IQE)
    TAB_MUN_NCONF_VI_IQE = add_total(TAB_MUN_NCONF_VI_IQE)

    TAB_MUN_RESUMIDA_IQE = TAB_MUN_FINAL_IQE[[
      "MUNICÍPIO", "PLANO DE AMOSTRAGEM", "DESCONSIDERAÇÕES DO PLANO",
      "PLANO DE AMOSTRAGEM AJUSTADO - ARSAL",
      "ANALISES REALIZADAS", "EXPURGOS TOTAIS", "NAM REALIZ"
    ]]
    TAB_MUN_RESUMIDA_IQE = TAB_MUN_RESUMIDA_IQE.rename(columns={"PLANO DE AMOSTRAGEM": "PLANO DE AMOSTRAGEM AJUSTADO"})
    TAB_MUN_RESUMIDA_IQE["PLANO DE AMOSTRAGEM AJUSTADO"] = TAB_MUN_RESUMIDA_IQE["PLANO DE AMOSTRAGEM AJUSTADO"].astype(int)
    TAB_MUN_RESUMIDA_IQE.loc[TAB_MUN_RESUMIDA_IQE.index[-1], "MUNICÍPIO"] = "TOTAL"

    TAB_MUN_RESUMIDA_VI_IQE = TAB_MUN_RESUMIDA_IQE[[
      "MUNICÍPIO", "PLANO DE AMOSTRAGEM AJUSTADO", "ANALISES REALIZADAS"
    ]]

    TAB_MUN_RESUMIDA_VI_IQE = TAB_MUN_RESUMIDA_VI_IQE.fillna(0)

    TAB_MUN_RESUMIDA_VI_IQE["NAM REALIZ (VI)"] = (np.maximum(
      TAB_MUN_RESUMIDA_VI_IQE["PLANO DE AMOSTRAGEM AJUSTADO"],
      TAB_MUN_RESUMIDA_VI_IQE['ANALISES REALIZADAS']
    )).astype(int)
    TAB_MUN_RESUMIDA_VI_IQE = TAB_MUN_RESUMIDA_VI_IQE[:-1]
    TAB_MUN_RESUMIDA_VI_IQE.loc['Total'] = TAB_MUN_RESUMIDA_VI_IQE.sum()
    TAB_MUN_RESUMIDA_VI_IQE.iloc[-1, TAB_MUN_RESUMIDA_VI_IQE.columns.get_loc("MUNICÍPIO")] = "TOTAL"

    return TAB_MUN_FINAL, TAB_MUN_RESUMIDA, TAB_MUN_RESUMIDA_VI, desc_plano, TAB_MUN_NCONF, TAB_MUN_NCONF_VI, TAB_MUN_FINAL_IQE, TAB_MUN_RESUMIDA_IQE, TAB_MUN_RESUMIDA_VI_IQE, desc_plano_IQE, TAB_MUN_NCONF_IQE, TAB_MUN_NCONF_VI_IQE



  # ANÁLISE POR PARÂMETRO

  def analise_tipo(TAB_TIPO, TAB_REALIZADAS, BASE_MUN_TIPO_completo, TAB_TIPO_IQE, TAB_REALIZADAS_IQE, BASE_MUN_TIPO_IQE_completo):

    # IQA

    mapa_parametros = {
      "Turbidez":            "TURBIDEZ",
      "Cor Aparente":        "COR APARENTE",
      "pH":                  "PH",
      "Cloro residual livre": "CLORO RESIDUAL LIVRE",
      "Coliformes totais":   "COLIFORMES TOTAIS",
      "Escherichia Coli":    "ESCHERICHIA COLI",
      "Demais Parâmetros":   "DEMAIS PARAMETROS",
    }

    DESC_TIPO = (
      BASE_MUN_TIPO_completo
      .groupby("Parâmetros", as_index=False)["DESCONSIDERAÇÕES - ARSAL"]
      .sum()
    )
    DESC_TIPO["Parâmetros"] = DESC_TIPO["Parâmetros"].map(mapa_parametros)
    DESC_TIPO = DESC_TIPO.rename(columns={
      "Parâmetros":              "ANÁLISE",
      "DESCONSIDERAÇÕES - ARSAL": "DESCONSIDERAÇÕES DO PLANO"
    })

    TAB_TIPO = TAB_TIPO.merge(DESC_TIPO, on="ANÁLISE", how="left")
    TAB_TIPO["DESCONSIDERAÇÕES DO PLANO"] = TAB_TIPO["DESCONSIDERAÇÕES DO PLANO"].fillna(0).astype(int)
    TAB_TIPO["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"] = (TAB_TIPO["PLANO DE AMOSTRAGEM"] - TAB_TIPO["DESCONSIDERAÇÕES DO PLANO"]).astype(int)

    TAB_REALIZADAS = TAB_REALIZADAS.replace("CLORO", "CLORO RESIDUAL LIVRE")

    TAB_REALIZADAS['ANÁLISE'] = np.where(
      TAB_REALIZADAS['ANÁLISE'].isin(['TURBIDEZ', 'COR APARENTE', 'PH', 'COLIFORMES TOTAIS', 'CLORO RESIDUAL LIVRE', 'ESCHERICHIA COLI']),
      TAB_REALIZADAS['ANÁLISE'], 'DEMAIS PARAMETROS'
    )

    TAB_TIPO_VI = TAB_TIPO

    TAB_REALIZADAS_ACEITO = TAB_REALIZADAS[TAB_REALIZADAS["ACEITO?"] == "Validado"]

    CONF_VI_NAOCONF_ARSAL = TAB_REALIZADAS_ACEITO[
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - VI"] == "Conforme") &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - ARSAL"] != "Conforme") &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - ARSAL"].notna()) &
        (TAB_REALIZADAS_ACEITO["EXPURGOS - ARSAL"] != "EXPURGAR")
    ]

    NAOCONF_VI_CONF_ARSAL = TAB_REALIZADAS_ACEITO[
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - VI"] != "Conforme") &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - VI"].notna()) &
        (TAB_REALIZADAS_ACEITO["CONFORMIDADE - ARSAL"] == "Conforme") &
        (TAB_REALIZADAS_ACEITO["EXPURGOS - ARSAL"] != "EXPURGAR")
    ]

    contagem_conf_vi_naoconf_arsal = (
        CONF_VI_NAOCONF_ARSAL.groupby("ANÁLISE").size()
        .reset_index(name="CONVERSÃO: CONFORME → NÃO CONFORME")
    )

    contagem_naoconf_vi_conf_arsal = (
        NAOCONF_VI_CONF_ARSAL.groupby("ANÁLISE").size()
        .reset_index(name="CONVERSÃO: NÃO CONFORME → CONFORME")
    )

    TAB_TIPO = TAB_TIPO.merge(contagem_conf_vi_naoconf_arsal, on="ANÁLISE", how="left")
    TAB_TIPO["CONVERSÃO: CONFORME → NÃO CONFORME"] = TAB_TIPO["CONVERSÃO: CONFORME → NÃO CONFORME"].fillna(0).astype(int)

    TAB_TIPO = TAB_TIPO.merge(contagem_naoconf_vi_conf_arsal, on="ANÁLISE", how="left")
    TAB_TIPO["CONVERSÃO: NÃO CONFORME → CONFORME"] = TAB_TIPO["CONVERSÃO: NÃO CONFORME → CONFORME"].fillna(0).astype(int)

    TAB_REALIZADAS_ACEITO_CONF = TAB_REALIZADAS_ACEITO[TAB_REALIZADAS["CONFORMIDADE - ARSAL"] == "Conforme"]
    TAB_REALIZADAS_ACEITO_CONF_VI = TAB_REALIZADAS_ACEITO[TAB_REALIZADAS["CONFORMIDADE - VI"] == "Conforme"]
    TAB_REALIZADAS_ACEITO_NAOCONF = TAB_REALIZADAS_ACEITO[(TAB_REALIZADAS["CONFORMIDADE - ARSAL"] != "Conforme") & (TAB_REALIZADAS["CONFORMIDADE - ARSAL"].notna())]
    TAB_REALIZADAS_ACEITO_NAOCONF_VI = TAB_REALIZADAS_ACEITO[(TAB_REALIZADAS["CONFORMIDADE - VI"] != "Conforme") & (TAB_REALIZADAS["CONFORMIDADE - VI"].notna())]
    TAB_REALIZADAS_ACEITO_CONF_EXP = TAB_REALIZADAS_ACEITO_CONF[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_CONF_EXP_VI = TAB_REALIZADAS_ACEITO_CONF_VI[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_NAOCONF_EXP = TAB_REALIZADAS_ACEITO_NAOCONF[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI = TAB_REALIZADAS_ACEITO_NAOCONF_VI[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]


    print("\n===== DEBUG IQA =====")

    print("\nANÁLISES:")
    print(TAB_REALIZADAS["ANÁLISE"].value_counts(dropna=False))

    print("\nACEITOS:")
    print(TAB_REALIZADAS["ACEITO?"].value_counts(dropna=False))

    print("\nCONFORMIDADE VI:")
    print(TAB_REALIZADAS_ACEITO["CONFORMIDADE - VI"].value_counts(dropna=False))

    print("\nCONFORMIDADE ARSAL:")
    print(TAB_REALIZADAS_ACEITO["CONFORMIDADE - ARSAL"].value_counts(dropna=False))

    print("\nTOTAL REALIZADAS ACEITAS:")
    print(len(TAB_REALIZADAS_ACEITO))

    print("\nTOTAL CONFORMES VI:")
    print(len(TAB_REALIZADAS_ACEITO_CONF_VI))

    print("\nTOTAL NÃO CONFORMES VI:")
    print(len(TAB_REALIZADAS_ACEITO_NAOCONF_VI))

    contagem_tot = TAB_REALIZADAS_ACEITO.groupby('ANÁLISE').size().reset_index(name='ANALISES REALIZADAS')
    TAB_TIPO = TAB_TIPO.merge(contagem_tot, on='ANÁLISE', how='left')

    TAB_TIPO = TAB_TIPO.fillna(0)

    TAB_TIPO['ANALISES REALIZADAS'] = TAB_TIPO['ANALISES REALIZADAS'].astype(int)

    contagem_conf = TAB_REALIZADAS_ACEITO_CONF_VI.groupby('ANÁLISE').size().reset_index(name='ANALISES CONFORMES')
    TAB_TIPO = TAB_TIPO.merge(contagem_conf, on='ANÁLISE', how='left')

    contagem_naoconf = TAB_REALIZADAS_ACEITO_NAOCONF_VI.groupby('ANÁLISE').size().reset_index(name='ANALISES NÃO CONFORMES')
    TAB_TIPO = TAB_TIPO.merge(contagem_naoconf, on='ANÁLISE', how='left')
    TAB_TIPO['ANALISES NÃO CONFORMES'] = TAB_TIPO['ANALISES NÃO CONFORMES'].fillna(0).astype(int)

    contagem_conf_exp = TAB_REALIZADAS_ACEITO_CONF_EXP_VI.groupby('ANÁLISE').size().reset_index(name='EXPURGOS CONFORMES')
    TAB_TIPO = TAB_TIPO.merge(contagem_conf_exp, on='ANÁLISE', how='left')
    TAB_TIPO['EXPURGOS CONFORMES'] = TAB_TIPO['EXPURGOS CONFORMES'].fillna(0).astype(int)

    contagem_naoconf_exp = TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI.groupby('ANÁLISE').size().reset_index(name='EXPURGOS NAO CONFORMES')
    TAB_TIPO = TAB_TIPO.merge(contagem_naoconf_exp, on='ANÁLISE', how='left')
    TAB_TIPO['EXPURGOS NAO CONFORMES'] = TAB_TIPO['EXPURGOS NAO CONFORMES'].fillna(0).astype(int)

    TAB_TIPO['EXPURGOS TOTAIS'] = TAB_TIPO['EXPURGOS NAO CONFORMES'] + TAB_TIPO['EXPURGOS CONFORMES']

    TAB_TIPO["NAM REALIZ"] = (np.maximum(TAB_TIPO['PLANO DE AMOSTRAGEM AJUSTADO - ARSAL'], TAB_TIPO['ANALISES REALIZADAS']) - TAB_TIPO['EXPURGOS TOTAIS']).astype(int)

    TAB_TIPO = TAB_TIPO.fillna(0)

    TAB_TIPO["NAM CONF"] = (TAB_TIPO['ANALISES CONFORMES'] - TAB_TIPO['EXPURGOS CONFORMES'] + TAB_TIPO['CONVERSÃO: NÃO CONFORME → CONFORME'] - TAB_TIPO['CONVERSÃO: CONFORME → NÃO CONFORME']).astype(int)

    TAB_TIPO["IQA (%)"] = (((TAB_TIPO["NAM CONF"]) / (TAB_TIPO["NAM REALIZ"])) * 100).round(1)

    TAB_TIPO_FINAL = TAB_TIPO[[
      "ANÁLISE", "PLANO DE AMOSTRAGEM", "DESCONSIDERAÇÕES DO PLANO",
      "PLANO DE AMOSTRAGEM AJUSTADO - ARSAL",
      "ANALISES REALIZADAS", "ANALISES CONFORMES", "EXPURGOS CONFORMES",
      "CONVERSÃO: CONFORME → NÃO CONFORME", "CONVERSÃO: NÃO CONFORME → CONFORME", "NAM CONF"
    ]]
    TAB_TIPO_FINAL.loc['Total'] = TAB_TIPO_FINAL.sum()
    TAB_TIPO_FINAL["ANÁLISE"].iat[-1] = ' '
    TAB_TIPO_FINAL.loc[TAB_TIPO_FINAL.index[-1], "ANÁLISE"] = "TOTAL"

    TAB_TIPO_FINAL.rename(columns={"PLANO DE AMOSTRAGEM": "PLANO DE AMOSTRAGEM AJUSTADO"})

    total_conf_vi_naoconf_arsal = int(TAB_TIPO_FINAL["CONVERSÃO: CONFORME → NÃO CONFORME"].iloc[-1])
    total_naoconf_vi_conf_arsal = int(TAB_TIPO_FINAL["CONVERSÃO: NÃO CONFORME → CONFORME"].iloc[-1])

    TAB_TIPO_FINAL_RESUMIDA = TAB_TIPO_FINAL[["ANÁLISE", "ANALISES REALIZADAS", "ANALISES CONFORMES", "EXPURGOS CONFORMES", "CONVERSÃO: CONFORME → NÃO CONFORME", "CONVERSÃO: NÃO CONFORME → CONFORME", "NAM CONF"]]

    TAB_REALIZADAS_ACEITO_CONF_VI = TAB_REALIZADAS_ACEITO[TAB_REALIZADAS["CONFORMIDADE - VI"] == "Conforme"]
    TAB_REALIZADAS_ACEITO_NAOCONF_VI = TAB_REALIZADAS_ACEITO[(TAB_REALIZADAS["CONFORMIDADE - VI"] != "Conforme") & (TAB_REALIZADAS["CONFORMIDADE - VI"].notna())]
    TAB_REALIZADAS_ACEITO_CONF_EXP_VI = TAB_REALIZADAS_ACEITO_CONF_VI[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI = TAB_REALIZADAS_ACEITO_NAOCONF_VI[TAB_REALIZADAS["EXPURGOS - ARSAL"] == "EXPURGAR"]

    contagem_tot_VI = TAB_REALIZADAS_ACEITO.groupby('ANÁLISE').size().reset_index(name='ANALISES REALIZADAS')
    TAB_TIPO_VI = TAB_TIPO_VI.merge(contagem_tot_VI, on='ANÁLISE', how='left')
    TAB_TIPO_VI = TAB_TIPO_VI.fillna(0)
    TAB_TIPO_VI['ANALISES REALIZADAS'] = TAB_TIPO_VI['ANALISES REALIZADAS'].astype(int)

    contagem_conf_VI = TAB_REALIZADAS_ACEITO_CONF_VI.groupby('ANÁLISE').size().reset_index(name='ANALISES CONFORMES')
    TAB_TIPO_VI = TAB_TIPO_VI.merge(contagem_conf_VI, on='ANÁLISE', how='left')

    contagem_naoconf_VI = TAB_REALIZADAS_ACEITO_NAOCONF_VI.groupby('ANÁLISE').size().reset_index(name='ANALISES NÃO CONFORMES')
    TAB_TIPO_VI = TAB_TIPO_VI.merge(contagem_naoconf_VI, on='ANÁLISE', how='left')
    TAB_TIPO_VI['ANALISES NÃO CONFORMES'] = TAB_TIPO_VI['ANALISES NÃO CONFORMES'].fillna(0).astype(int)

    contagem_conf_exp_VI = TAB_REALIZADAS_ACEITO_CONF_EXP_VI.groupby('ANÁLISE').size().reset_index(name='EXPURGOS CONFORMES')
    TAB_TIPO_VI = TAB_TIPO_VI.merge(contagem_conf_exp_VI, on='ANÁLISE', how='left')
    TAB_TIPO_VI['EXPURGOS CONFORMES'] = TAB_TIPO_VI['EXPURGOS CONFORMES'].fillna(0).astype(int)

    contagem_naoconf_exp_VI = TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI.groupby('ANÁLISE').size().reset_index(name='EXPURGOS NAO CONFORMES')
    TAB_TIPO_VI = TAB_TIPO_VI.merge(contagem_naoconf_exp_VI, on='ANÁLISE', how='left')
    TAB_TIPO_VI['EXPURGOS NAO CONFORMES'] = TAB_TIPO_VI['EXPURGOS NAO CONFORMES'].fillna(0).astype(int)

    TAB_TIPO_VI['EXPURGOS TOTAIS'] = TAB_TIPO_VI['EXPURGOS NAO CONFORMES'] + TAB_TIPO_VI['EXPURGOS CONFORMES']

    TAB_TIPO_VI["NAM REALIZ"] = (np.maximum(TAB_TIPO_VI['PLANO DE AMOSTRAGEM AJUSTADO - ARSAL'], TAB_TIPO_VI['ANALISES REALIZADAS']) - TAB_TIPO_VI['EXPURGOS TOTAIS']).astype(int)

    TAB_TIPO_VI = TAB_TIPO_VI.fillna(0)
    TAB_TIPO_VI["NAM CONF"] = (TAB_TIPO_VI['ANALISES CONFORMES'] - TAB_TIPO_VI['EXPURGOS CONFORMES']).astype(int)

    TAB_TIPO_VI["IQA (%)"] = (((TAB_TIPO_VI["NAM CONF"]) / (TAB_TIPO_VI["NAM REALIZ"])) * 100).round(1)

    TAB_TIPO_FINAL_VI = TAB_TIPO_VI[[
      "ANÁLISE", "PLANO DE AMOSTRAGEM", "DESCONSIDERAÇÕES DO PLANO",
      "PLANO DE AMOSTRAGEM AJUSTADO - ARSAL",
      "ANALISES REALIZADAS", "ANALISES CONFORMES", "EXPURGOS CONFORMES", "NAM CONF"
    ]]
    TAB_TIPO_FINAL_VI.loc['Total'] = TAB_TIPO_FINAL_VI.sum()
    TAB_TIPO_FINAL_VI["ANÁLISE"].iat[-1] = ' '
    TAB_TIPO_FINAL_VI.loc[TAB_TIPO_FINAL_VI.index[-1], "ANÁLISE"] = "TOTAL"

    TAB_TIPO_RESUMIDA_VI = TAB_TIPO_FINAL_VI[["ANÁLISE", "ANALISES REALIZADAS", "ANALISES CONFORMES"]]
    TAB_TIPO_RESUMIDA_VI["NAM CONF (VI)"] = TAB_TIPO_RESUMIDA_VI["ANALISES CONFORMES"]


    # IQE

    mapa_parametros_IQE = {
      "DBO":                 "DBO",
      "DQO":                 "DQO",
      "Temperatura":         "TEMPERATURA",
      "Óleos e Graxas":      "OLEOS E GRAXAS",
      "Materiais Flutuantes": "MATERIAIS FLUTUANTES",
      "pH":                  "PH",
      "Demais Parâmetros":   "DEMAIS PARAMETROS",
    }

    DESC_TIPO_IQE = (
      BASE_MUN_TIPO_IQE_completo
      .groupby("Parâmetros", as_index=False)["DESCONSIDERAÇÕES - ARSAL"]
      .sum()
    )
    DESC_TIPO_IQE["Parâmetros"] = DESC_TIPO_IQE["Parâmetros"].map(mapa_parametros_IQE)
    DESC_TIPO_IQE = DESC_TIPO_IQE.rename(columns={
      "Parâmetros":               "ANÁLISE",
      "DESCONSIDERAÇÕES - ARSAL": "DESCONSIDERAÇÕES DO PLANO"
    })

    TAB_TIPO_IQE = TAB_TIPO_IQE.merge(DESC_TIPO_IQE, on="ANÁLISE", how="left")
    TAB_TIPO_IQE["DESCONSIDERAÇÕES DO PLANO"] = TAB_TIPO_IQE["DESCONSIDERAÇÕES DO PLANO"].fillna(0).astype(int)
    TAB_TIPO_IQE["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"] = (
      TAB_TIPO_IQE["PLANO DE AMOSTRAGEM"] - TAB_TIPO_IQE["DESCONSIDERAÇÕES DO PLANO"]
    ).astype(int)

    TAB_REALIZADAS_IQE['ANÁLISE'] = np.where(
      TAB_REALIZADAS_IQE['ANÁLISE'].isin(['DBO', 'DQO', 'PH', 'MATERIAIS FLUTUANTES', 'OLEOS E GRAXAS', 'TEMPERATURA']),
      TAB_REALIZADAS_IQE['ANÁLISE'], 'DEMAIS PARAMETROS'
    )

    TAB_TIPO_VI_IQE = TAB_TIPO_IQE.copy()

    TAB_REALIZADAS_ACEITO_IQE = TAB_REALIZADAS_IQE[TAB_REALIZADAS_IQE["ACEITO?"] == "Validado"]

    CONF_VI_NAOCONF_ARSAL_IQE = TAB_REALIZADAS_ACEITO_IQE[
      (TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - VI"]    == "Conforme") &
      (TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - ARSAL"] != "Conforme") &
      (TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - ARSAL"].notna()) &
      (TAB_REALIZADAS_ACEITO_IQE["EXPURGOS - ARSAL"]     != "EXPURGAR")
    ]

    NAOCONF_VI_CONF_ARSAL_IQE = TAB_REALIZADAS_ACEITO_IQE[
      (TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - VI"]    != "Conforme") &
      (TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - VI"].notna()) &
      (TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - ARSAL"] == "Conforme") &
      (TAB_REALIZADAS_ACEITO_IQE["EXPURGOS - ARSAL"]     != "EXPURGAR")
    ]

    contagem_conf_vi_naoconf_arsal_IQE = (
      CONF_VI_NAOCONF_ARSAL_IQE.groupby("ANÁLISE").size()
      .reset_index(name="CONVERSÃO: CONFORME → NÃO CONFORME")
    )
    contagem_naoconf_vi_conf_arsal_IQE = (
      NAOCONF_VI_CONF_ARSAL_IQE.groupby("ANÁLISE").size()
      .reset_index(name="CONVERSÃO: NÃO CONFORME → CONFORME")
    )

    TAB_TIPO_IQE = TAB_TIPO_IQE.merge(contagem_conf_vi_naoconf_arsal_IQE, on="ANÁLISE", how="left")
    TAB_TIPO_IQE["CONVERSÃO: CONFORME → NÃO CONFORME"] = TAB_TIPO_IQE["CONVERSÃO: CONFORME → NÃO CONFORME"].fillna(0).astype(int)

    TAB_TIPO_IQE = TAB_TIPO_IQE.merge(contagem_naoconf_vi_conf_arsal_IQE, on="ANÁLISE", how="left")
    TAB_TIPO_IQE["CONVERSÃO: NÃO CONFORME → CONFORME"] = TAB_TIPO_IQE["CONVERSÃO: NÃO CONFORME → CONFORME"].fillna(0).astype(int)

    TAB_REALIZADAS_ACEITO_CONF_IQE       = TAB_REALIZADAS_ACEITO_IQE[TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - ARSAL"] == "Conforme"]
    TAB_REALIZADAS_ACEITO_NAOCONF_IQE    = TAB_REALIZADAS_ACEITO_IQE[(TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - ARSAL"] != "Conforme") & (TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - ARSAL"].notna())]
    TAB_REALIZADAS_ACEITO_CONF_EXP_IQE   = TAB_REALIZADAS_ACEITO_CONF_IQE[TAB_REALIZADAS_ACEITO_IQE["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_NAOCONF_EXP_IQE = TAB_REALIZADAS_ACEITO_NAOCONF_IQE[TAB_REALIZADAS_ACEITO_IQE["EXPURGOS - ARSAL"] == "EXPURGAR"]

    TAB_REALIZADAS_ACEITO_CONF_VI_IQE       = TAB_REALIZADAS_ACEITO_IQE[TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - VI"] == "Conforme"]
    TAB_REALIZADAS_ACEITO_NAOCONF_VI_IQE    = TAB_REALIZADAS_ACEITO_IQE[(TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - VI"] != "Conforme") & (TAB_REALIZADAS_ACEITO_IQE["CONFORMIDADE - VI"].notna())]
    TAB_REALIZADAS_ACEITO_CONF_EXP_VI_IQE   = TAB_REALIZADAS_ACEITO_CONF_VI_IQE[TAB_REALIZADAS_ACEITO_IQE["EXPURGOS - ARSAL"] == "EXPURGAR"]
    TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI_IQE = TAB_REALIZADAS_ACEITO_NAOCONF_VI_IQE[TAB_REALIZADAS_ACEITO_IQE["EXPURGOS - ARSAL"] == "EXPURGAR"]

    contagem_tot_IQE = TAB_REALIZADAS_ACEITO_IQE.groupby('ANÁLISE').size().reset_index(name='ANALISES REALIZADAS')
    TAB_TIPO_IQE = TAB_TIPO_IQE.merge(contagem_tot_IQE, on='ANÁLISE', how='left')
    TAB_TIPO_IQE['ANALISES REALIZADAS'] = TAB_TIPO_IQE['ANALISES REALIZADAS'].astype(int)

    contagem_conf_IQE = TAB_REALIZADAS_ACEITO_CONF_VI_IQE.groupby('ANÁLISE').size().reset_index(name='ANALISES CONFORMES')
    TAB_TIPO_IQE = TAB_TIPO_IQE.merge(contagem_conf_IQE, on='ANÁLISE', how='left')

    contagem_naoconf_IQE = TAB_REALIZADAS_ACEITO_NAOCONF_VI_IQE.groupby('ANÁLISE').size().reset_index(name='ANALISES NÃO CONFORMES')
    TAB_TIPO_IQE = TAB_TIPO_IQE.merge(contagem_naoconf_IQE, on='ANÁLISE', how='left')
    TAB_TIPO_IQE['ANALISES NÃO CONFORMES'] = TAB_TIPO_IQE['ANALISES NÃO CONFORMES'].fillna(0).astype(int)

    contagem_conf_exp_IQE = TAB_REALIZADAS_ACEITO_CONF_EXP_VI_IQE.groupby('ANÁLISE').size().reset_index(name='EXPURGOS CONFORMES')
    TAB_TIPO_IQE = TAB_TIPO_IQE.merge(contagem_conf_exp_IQE, on='ANÁLISE', how='left')
    TAB_TIPO_IQE['EXPURGOS CONFORMES'] = TAB_TIPO_IQE['EXPURGOS CONFORMES'].fillna(0).astype(int)

    contagem_naoconf_exp_IQE = TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI_IQE.groupby('ANÁLISE').size().reset_index(name='EXPURGOS NAO CONFORMES')
    TAB_TIPO_IQE = TAB_TIPO_IQE.merge(contagem_naoconf_exp_IQE, on='ANÁLISE', how='left')
    TAB_TIPO_IQE['EXPURGOS NAO CONFORMES'] = TAB_TIPO_IQE['EXPURGOS NAO CONFORMES'].fillna(0).astype(int)

    TAB_TIPO_IQE['EXPURGOS TOTAIS'] = TAB_TIPO_IQE['EXPURGOS NAO CONFORMES'] + TAB_TIPO_IQE['EXPURGOS CONFORMES']

    TAB_TIPO_IQE["NAM REALIZ"] = (
      np.maximum(TAB_TIPO_IQE['PLANO DE AMOSTRAGEM AJUSTADO - ARSAL'], TAB_TIPO_IQE['ANALISES REALIZADAS'])
      - TAB_TIPO_IQE['EXPURGOS TOTAIS']
    ).astype(int)

    TAB_TIPO_IQE = TAB_TIPO_IQE.fillna(0)

    TAB_TIPO_IQE["NAM CONF"] = (
      TAB_TIPO_IQE['ANALISES CONFORMES']
      - TAB_TIPO_IQE['EXPURGOS CONFORMES']
      + TAB_TIPO_IQE['CONVERSÃO: NÃO CONFORME → CONFORME']
      - TAB_TIPO_IQE['CONVERSÃO: CONFORME → NÃO CONFORME']
    ).astype(int)

    TAB_TIPO_IQE["IQE (%)"] = (
      (TAB_TIPO_IQE["NAM CONF"] / TAB_TIPO_IQE["NAM REALIZ"]) * 100
    ).round(1)

    TAB_TIPO_FINAL_IQE = TAB_TIPO_IQE[[
      "ANÁLISE", "PLANO DE AMOSTRAGEM", "DESCONSIDERAÇÕES DO PLANO",
      "PLANO DE AMOSTRAGEM AJUSTADO - ARSAL",
      "ANALISES REALIZADAS", "ANALISES CONFORMES", "EXPURGOS CONFORMES",
      "CONVERSÃO: CONFORME → NÃO CONFORME", "CONVERSÃO: NÃO CONFORME → CONFORME", "NAM CONF"
    ]]
    TAB_TIPO_FINAL_IQE.loc['Total'] = TAB_TIPO_FINAL_IQE.sum()
    TAB_TIPO_FINAL_IQE.loc[TAB_TIPO_FINAL_IQE.index[-1], "ANÁLISE"] = "TOTAL"

    total_conf_vi_naoconf_arsal_IQE = int(TAB_TIPO_FINAL_IQE["CONVERSÃO: CONFORME → NÃO CONFORME"].iloc[-1])
    total_naoconf_vi_conf_arsal_IQE = int(TAB_TIPO_FINAL_IQE["CONVERSÃO: NÃO CONFORME → CONFORME"].iloc[-1])

    TAB_TIPO_FINAL_RESUMIDA_IQE = TAB_TIPO_FINAL_IQE[[
      "ANÁLISE", "ANALISES REALIZADAS", "ANALISES CONFORMES", "EXPURGOS CONFORMES",
      "CONVERSÃO: CONFORME → NÃO CONFORME", "CONVERSÃO: NÃO CONFORME → CONFORME", "NAM CONF"
    ]]

    contagem_tot_VI_IQE = TAB_REALIZADAS_ACEITO_IQE.groupby('ANÁLISE').size().reset_index(name='ANALISES REALIZADAS')
    TAB_TIPO_VI_IQE = TAB_TIPO_VI_IQE.merge(contagem_tot_VI_IQE, on='ANÁLISE', how='left')
    TAB_TIPO_VI_IQE['ANALISES REALIZADAS'] = TAB_TIPO_VI_IQE['ANALISES REALIZADAS'].astype(int)

    contagem_conf_VI_IQE = TAB_REALIZADAS_ACEITO_CONF_VI_IQE.groupby('ANÁLISE').size().reset_index(name='ANALISES CONFORMES')
    TAB_TIPO_VI_IQE = TAB_TIPO_VI_IQE.merge(contagem_conf_VI_IQE, on='ANÁLISE', how='left')

    contagem_naoconf_VI_IQE = TAB_REALIZADAS_ACEITO_NAOCONF_VI_IQE.groupby('ANÁLISE').size().reset_index(name='ANALISES NÃO CONFORMES')
    TAB_TIPO_VI_IQE = TAB_TIPO_VI_IQE.merge(contagem_naoconf_VI_IQE, on='ANÁLISE', how='left')
    TAB_TIPO_VI_IQE['ANALISES NÃO CONFORMES'] = TAB_TIPO_VI_IQE['ANALISES NÃO CONFORMES'].fillna(0).astype(int)

    contagem_conf_exp_VI_IQE = TAB_REALIZADAS_ACEITO_CONF_EXP_VI_IQE.groupby('ANÁLISE').size().reset_index(name='EXPURGOS CONFORMES')
    TAB_TIPO_VI_IQE = TAB_TIPO_VI_IQE.merge(contagem_conf_exp_VI_IQE, on='ANÁLISE', how='left')
    TAB_TIPO_VI_IQE['EXPURGOS CONFORMES'] = TAB_TIPO_VI_IQE['EXPURGOS CONFORMES'].fillna(0).astype(int)

    contagem_naoconf_exp_VI_IQE = TAB_REALIZADAS_ACEITO_NAOCONF_EXP_VI_IQE.groupby('ANÁLISE').size().reset_index(name='EXPURGOS NAO CONFORMES')
    TAB_TIPO_VI_IQE = TAB_TIPO_VI_IQE.merge(contagem_naoconf_exp_VI_IQE, on='ANÁLISE', how='left')
    TAB_TIPO_VI_IQE['EXPURGOS NAO CONFORMES'] = TAB_TIPO_VI_IQE['EXPURGOS NAO CONFORMES'].fillna(0).astype(int)

    TAB_TIPO_VI_IQE['EXPURGOS TOTAIS'] = TAB_TIPO_VI_IQE['EXPURGOS NAO CONFORMES'] + TAB_TIPO_VI_IQE['EXPURGOS CONFORMES']

    TAB_TIPO_VI_IQE = TAB_TIPO_VI_IQE.fillna(0)

    TAB_TIPO_VI_IQE["NAM REALIZ"] = (
      np.maximum(TAB_TIPO_VI_IQE['PLANO DE AMOSTRAGEM AJUSTADO - ARSAL'], TAB_TIPO_VI_IQE['ANALISES REALIZADAS'])
      - TAB_TIPO_VI_IQE['EXPURGOS TOTAIS']
    ).astype(int)

    TAB_TIPO_VI_IQE = TAB_TIPO_VI_IQE.fillna(0)

    TAB_TIPO_VI_IQE["NAM CONF"] = (
      TAB_TIPO_VI_IQE['ANALISES CONFORMES'] - TAB_TIPO_VI_IQE['EXPURGOS CONFORMES']
    ).astype(int)

    TAB_TIPO_VI_IQE = TAB_TIPO_VI_IQE.fillna(0)

    TAB_TIPO_VI_IQE["IQE (%)"] = (
      (TAB_TIPO_VI_IQE["NAM CONF"] / TAB_TIPO_VI_IQE["NAM REALIZ"]) * 100
    ).round(1)

    TAB_TIPO_FINAL_VI_IQE = TAB_TIPO_VI_IQE[[
      "ANÁLISE", "PLANO DE AMOSTRAGEM", "DESCONSIDERAÇÕES DO PLANO",
      "PLANO DE AMOSTRAGEM AJUSTADO - ARSAL",
      "ANALISES REALIZADAS", "ANALISES CONFORMES", "EXPURGOS CONFORMES", "NAM CONF"
    ]]
    TAB_TIPO_FINAL_VI_IQE.loc['Total'] = TAB_TIPO_FINAL_VI_IQE.sum()
    TAB_TIPO_FINAL_VI_IQE.loc[TAB_TIPO_FINAL_VI_IQE.index[-1], "ANÁLISE"] = "TOTAL"

    TAB_TIPO_RESUMIDA_VI_IQE = TAB_TIPO_FINAL_VI_IQE[["ANÁLISE", "ANALISES REALIZADAS", "ANALISES CONFORMES"]]
    TAB_TIPO_RESUMIDA_VI_IQE["NAM CONF (VI)"] = TAB_TIPO_RESUMIDA_VI_IQE["ANALISES CONFORMES"]

    return TAB_TIPO, TAB_TIPO_FINAL_RESUMIDA, TAB_TIPO_RESUMIDA_VI,  total_conf_vi_naoconf_arsal, total_naoconf_vi_conf_arsal, TAB_TIPO_IQE, TAB_TIPO_FINAL_RESUMIDA_IQE, TAB_TIPO_RESUMIDA_VI_IQE,  total_conf_vi_naoconf_arsal_IQE, total_naoconf_vi_conf_arsal_IQE



# ELABORAÇÃO DA TABELA DE PONDERAÇÕES

  def pond(TAB_POND, TAB_REALIZADAS, BASE_MUN_TIPO_completo, TAB_POND_IQE, TAB_REALIZADAS_IQE, BASE_MUN_TIPO_IQE_completo):

    # IQA

    TAB_REALIZADAS["ID_POND"] = TAB_REALIZADAS["ID_POND"].fillna(0)

    BASE_MUN_TIPO_completo["ID_POND"] = BASE_MUN_TIPO_completo["ID_POND"].fillna("").astype(str).str.strip()

    DESC_POR_ID = (
      BASE_MUN_TIPO_completo[BASE_MUN_TIPO_completo["ID_POND"] != ""]
      .groupby("ID_POND", as_index=False)["DESCONSIDERAÇÕES - ARSAL"]
      .sum()
    )
    DESC_POR_ID["ID_POND"] = DESC_POR_ID["ID_POND"].astype(float).astype("Int64")

    TAB_REALIZADAS['ANÁLISE'] = np.where(
        TAB_REALIZADAS['ANÁLISE'].isin(['TURBIDEZ', 'COR APARENTE', 'PH', 'COLIFORMES TOTAIS', 'ESCHERICHIA COLI']),
        TAB_REALIZADAS['ANÁLISE'],
        'DEMAIS PARAMETROS'
    )

    TAB_REALIZADAS_ACEITO = TAB_REALIZADAS[TAB_REALIZADAS["ACEITO?"] == "Validado"]

    TAB_EXPURGOS = TAB_REALIZADAS_ACEITO[
        TAB_REALIZADAS_ACEITO["EXPURGOS - ARSAL"] == "EXPURGAR"
    ]

    TAB_EXPURGOS_CONF = TAB_EXPURGOS[
        TAB_EXPURGOS["CONFORMIDADE - VI"] == "Conforme"
    ]

    TAB_SEM_EXPURGO = TAB_REALIZADAS_ACEITO[
        TAB_REALIZADAS_ACEITO["EXPURGOS - ARSAL"] != "EXPURGAR"
    ]

    TAB_CONF_VI_NAOCONF_ARSAL = TAB_SEM_EXPURGO[
        (TAB_SEM_EXPURGO["CONFORMIDADE - VI"] == "Conforme") &
        (TAB_SEM_EXPURGO["CONFORMIDADE - ARSAL"] != "Conforme") &
        (TAB_SEM_EXPURGO["CONFORMIDADE - ARSAL"].notna())
    ]

    TAB_NAOCONF_VI_CONF_ARSAL = TAB_SEM_EXPURGO[
        (TAB_SEM_EXPURGO["CONFORMIDADE - VI"] != "Conforme") &
        (TAB_SEM_EXPURGO["CONFORMIDADE - VI"].notna()) &
        (TAB_SEM_EXPURGO["CONFORMIDADE - ARSAL"] == "Conforme")
    ]

    QUADRO_POND_NCONF = TAB_POND.copy()
    QUADRO_POND_NCONF = QUADRO_POND_NCONF.rename(columns={
      QUADRO_POND_NCONF.columns[0]: "ID",
      QUADRO_POND_NCONF.columns[1]: "PONDERAÇÕES (PRESTADORA)",
      QUADRO_POND_NCONF.columns[2]: "PONDERAÇÕES (VI)",
      QUADRO_POND_NCONF.columns[3]: "PONDERAÇÕES (ARSAL)",
    })

    contagem_exp = TAB_EXPURGOS_CONF.groupby('ID_POND').size().reset_index(name='EXPURGOS')
    QUADRO_POND_NCONF = QUADRO_POND_NCONF.merge(contagem_exp, left_on='ID', right_on='ID_POND', how='left')
    QUADRO_POND_NCONF = QUADRO_POND_NCONF.drop(columns=["ID_POND"])
    QUADRO_POND_NCONF["EXPURGOS"] = QUADRO_POND_NCONF["EXPURGOS"].fillna(0).astype("Int64")
    QUADRO_POND_NCONF["EXPURGOS"] = QUADRO_POND_NCONF["EXPURGOS"].astype("object").replace(0, "")

    contagem_cv_na = TAB_CONF_VI_NAOCONF_ARSAL.groupby('ID_POND').size().reset_index(name='CONVERSÃO: CONFORME → NÃO CONFORME')
    QUADRO_POND_NCONF = QUADRO_POND_NCONF.merge(contagem_cv_na, left_on='ID', right_on='ID_POND', how='left')
    QUADRO_POND_NCONF = QUADRO_POND_NCONF.drop(columns=["ID_POND"], errors="ignore")
    QUADRO_POND_NCONF["CONVERSÃO: CONFORME → NÃO CONFORME"] = QUADRO_POND_NCONF["CONVERSÃO: CONFORME → NÃO CONFORME"].fillna(0).astype("Int64")
    QUADRO_POND_NCONF["CONVERSÃO: CONFORME → NÃO CONFORME"] = QUADRO_POND_NCONF["CONVERSÃO: CONFORME → NÃO CONFORME"].astype("object").replace(0, "")

    contagem_na_cv = TAB_NAOCONF_VI_CONF_ARSAL.groupby('ID_POND').size().reset_index(name='CONVERSÃO: NÃO CONFORME → CONFORME')
    QUADRO_POND_NCONF = QUADRO_POND_NCONF.merge(contagem_na_cv, left_on='ID', right_on='ID_POND', how='left')
    QUADRO_POND_NCONF = QUADRO_POND_NCONF.drop(columns=["ID_POND"], errors="ignore")
    QUADRO_POND_NCONF["CONVERSÃO: NÃO CONFORME → CONFORME"] = QUADRO_POND_NCONF["CONVERSÃO: NÃO CONFORME → CONFORME"].fillna(0).astype("Int64")
    QUADRO_POND_NCONF["CONVERSÃO: NÃO CONFORME → CONFORME"] = QUADRO_POND_NCONF["CONVERSÃO: NÃO CONFORME → CONFORME"].astype("object").replace(0, "")

    colunas_impacto = ['EXPURGOS', 'CONVERSÃO: CONFORME → NÃO CONFORME', 'CONVERSÃO: NÃO CONFORME → CONFORME']
    QUADRO_POND_NCONF = QUADRO_POND_NCONF[
        QUADRO_POND_NCONF[colunas_impacto].apply(lambda row: any(v != "" for v in row), axis=1)
    ]

    for col in colunas_impacto:
        if col in QUADRO_POND_NCONF.columns:
            if QUADRO_POND_NCONF[col].replace("", pd.NA).isna().all():
                QUADRO_POND_NCONF = QUADRO_POND_NCONF.drop(columns=[col])

    QUADRO_POND_REALIZ = TAB_POND.copy()

    QUADRO_POND_REALIZ = QUADRO_POND_REALIZ.rename(columns={
      QUADRO_POND_REALIZ.columns[0]: "ID",
      QUADRO_POND_REALIZ.columns[1]: "PONDERAÇÕES (PRESTADORA)",
      QUADRO_POND_REALIZ.columns[2]: "PONDERAÇÕES (VI)",
      QUADRO_POND_REALIZ.columns[3]: "PONDERAÇÕES (ARSAL)",
    })

    contagem = TAB_EXPURGOS.groupby('ID_POND').size().reset_index(name='EXPURGOS')

    QUADRO_POND_REALIZ = QUADRO_POND_REALIZ.merge(
        contagem, left_on='ID', right_on='ID_POND', how='left'
    ).drop(columns=["ID_POND"], errors="ignore")

    QUADRO_POND_REALIZ["EXPURGOS"] = (
        QUADRO_POND_REALIZ["EXPURGOS"]
        .fillna(0)
        .astype("Int64")
        .astype("object")
        .replace(0, "")
    )

    QUADRO_POND_REALIZ = QUADRO_POND_REALIZ.merge(
        DESC_POR_ID, left_on='ID', right_on='ID_POND', how='left'
    ).drop(columns=["ID_POND"], errors="ignore")

    QUADRO_POND_REALIZ = QUADRO_POND_REALIZ.rename(
        columns={"DESCONSIDERAÇÕES - ARSAL": "DESCONSIDERAÇÕES (PLANO)"}
    )

    QUADRO_POND_REALIZ["DESCONSIDERAÇÕES (PLANO)"] = (
        QUADRO_POND_REALIZ["DESCONSIDERAÇÕES (PLANO)"]
        .fillna(0)
        .astype("Int64")
        .astype("object")
        .replace(0, "")
    )

    if QUADRO_POND_REALIZ[["EXPURGOS", "DESCONSIDERAÇÕES (PLANO)"]].replace("", pd.NA).isna().all().all():
        QUADRO_POND_REALIZ = "SEM IMPACTOS"


    
    # IQE

    TAB_REALIZADAS_IQE["ID_POND"] = TAB_REALIZADAS_IQE["ID_POND"].fillna(0)

    BASE_MUN_TIPO_IQE_completo["ID_POND"] = BASE_MUN_TIPO_IQE_completo["ID_POND"].fillna("").astype(str).str.strip()

    DESC_POR_ID_IQE = (
      BASE_MUN_TIPO_IQE_completo[BASE_MUN_TIPO_IQE_completo["ID_POND"] != ""]
      .groupby("ID_POND", as_index=False)["DESCONSIDERAÇÕES - ARSAL"]
      .sum()
    )
    DESC_POR_ID_IQE["ID_POND"] = DESC_POR_ID_IQE["ID_POND"].astype(float).astype("Int64")

    TAB_REALIZADAS_IQE['ANÁLISE'] = np.where(
        TAB_REALIZADAS_IQE['ANÁLISE'].isin(['DBO', 'DQO', 'PH', 'MATERIAIS FLUTUANTES', 'OLEOS E GRAXAS', 'TEMPERATURA']),
        TAB_REALIZADAS_IQE['ANÁLISE'],
        'DEMAIS PARAMETROS'
    )

    TAB_REALIZADAS_ACEITO_IQE = TAB_REALIZADAS_IQE[TAB_REALIZADAS_IQE["ACEITO?"] == "Validado"]

    TAB_EXPURGOS_IQE = TAB_REALIZADAS_ACEITO_IQE[
        TAB_REALIZADAS_ACEITO_IQE["EXPURGOS - ARSAL"] == "EXPURGAR"
    ]

    TAB_EXPURGOS_CONF_IQE = TAB_EXPURGOS_IQE[
        TAB_EXPURGOS_IQE["CONFORMIDADE - VI"] == "Conforme"
    ]

    TAB_SEM_EXPURGO_IQE = TAB_REALIZADAS_ACEITO_IQE[
        TAB_REALIZADAS_ACEITO_IQE["EXPURGOS - ARSAL"] != "EXPURGAR"
    ]

    TAB_CONF_VI_NAOCONF_ARSAL_IQE = TAB_SEM_EXPURGO_IQE[
        (TAB_SEM_EXPURGO_IQE["CONFORMIDADE - VI"] == "Conforme") &
        (TAB_SEM_EXPURGO_IQE["CONFORMIDADE - ARSAL"] != "Conforme") &
        (TAB_SEM_EXPURGO_IQE["CONFORMIDADE - ARSAL"].notna())
    ]

    TAB_NAOCONF_VI_CONF_ARSAL_IQE = TAB_SEM_EXPURGO_IQE[
        (TAB_SEM_EXPURGO_IQE["CONFORMIDADE - VI"] != "Conforme") &
        (TAB_SEM_EXPURGO_IQE["CONFORMIDADE - VI"].notna()) &
        (TAB_SEM_EXPURGO_IQE["CONFORMIDADE - ARSAL"] == "Conforme")
    ]

    QUADRO_POND_NCONF_IQE = TAB_POND_IQE.copy()
    QUADRO_POND_NCONF_IQE = QUADRO_POND_NCONF_IQE.rename(columns={
      QUADRO_POND_NCONF_IQE.columns[0]: "ID",
      QUADRO_POND_NCONF_IQE.columns[1]: "PONDERAÇÕES (PRESTADORA)",
      QUADRO_POND_NCONF_IQE.columns[2]: "PONDERAÇÕES (VI)",
      QUADRO_POND_NCONF_IQE.columns[3]: "PONDERAÇÕES (ARSAL)",
    })

    contagem_exp_IQE = TAB_EXPURGOS_CONF_IQE.groupby('ID_POND').size().reset_index(name='EXPURGOS')
    QUADRO_POND_NCONF_IQE = QUADRO_POND_NCONF_IQE.merge(contagem_exp_IQE, left_on='ID', right_on='ID_POND', how='left')
    QUADRO_POND_NCONF_IQE = QUADRO_POND_NCONF_IQE.drop(columns=["ID_POND"])
    QUADRO_POND_NCONF_IQE["EXPURGOS"] = QUADRO_POND_NCONF_IQE["EXPURGOS"].fillna(0).astype("Int64")
    QUADRO_POND_NCONF_IQE["EXPURGOS"] = QUADRO_POND_NCONF_IQE["EXPURGOS"].astype("object").replace(0, "")

    contagem_cv_na_IQE = TAB_CONF_VI_NAOCONF_ARSAL_IQE.groupby('ID_POND').size().reset_index(name='CONVERSÃO: CONFORME → NÃO CONFORME')
    QUADRO_POND_NCONF_IQE = QUADRO_POND_NCONF_IQE.merge(contagem_cv_na_IQE, left_on='ID', right_on='ID_POND', how='left')
    QUADRO_POND_NCONF_IQE = QUADRO_POND_NCONF_IQE.drop(columns=["ID_POND"], errors="ignore")
    QUADRO_POND_NCONF_IQE["CONVERSÃO: CONFORME → NÃO CONFORME"] = QUADRO_POND_NCONF_IQE["CONVERSÃO: CONFORME → NÃO CONFORME"].fillna(0).astype("Int64")
    QUADRO_POND_NCONF_IQE["CONVERSÃO: CONFORME → NÃO CONFORME"] = QUADRO_POND_NCONF_IQE["CONVERSÃO: CONFORME → NÃO CONFORME"].astype("object").replace(0, "")

    contagem_na_cv_IQE = TAB_NAOCONF_VI_CONF_ARSAL_IQE.groupby('ID_POND').size().reset_index(name='CONVERSÃO: NÃO CONFORME → CONFORME')
    QUADRO_POND_NCONF_IQE = QUADRO_POND_NCONF_IQE.merge(contagem_na_cv_IQE, left_on='ID', right_on='ID_POND', how='left')
    QUADRO_POND_NCONF_IQE = QUADRO_POND_NCONF_IQE.drop(columns=["ID_POND"], errors="ignore")
    QUADRO_POND_NCONF_IQE["CONVERSÃO: NÃO CONFORME → CONFORME"] = QUADRO_POND_NCONF_IQE["CONVERSÃO: NÃO CONFORME → CONFORME"].fillna(0).astype("Int64")
    QUADRO_POND_NCONF_IQE["CONVERSÃO: NÃO CONFORME → CONFORME"] = QUADRO_POND_NCONF_IQE["CONVERSÃO: NÃO CONFORME → CONFORME"].astype("object").replace(0, "")

    colunas_impacto_IQE = ['EXPURGOS', 'CONVERSÃO: CONFORME → NÃO CONFORME', 'CONVERSÃO: NÃO CONFORME → CONFORME']
    QUADRO_POND_NCONF_IQE = QUADRO_POND_NCONF_IQE[
        QUADRO_POND_NCONF_IQE[colunas_impacto_IQE].apply(lambda row: any(v != "" for v in row), axis=1)
    ]

    for col in colunas_impacto_IQE:
        if col in QUADRO_POND_NCONF_IQE.columns:
            if QUADRO_POND_NCONF_IQE[col].replace("", pd.NA).isna().all():
                QUADRO_POND_NCONF_IQE = QUADRO_POND_NCONF_IQE.drop(columns=[col])

    QUADRO_POND_REALIZ_IQE = TAB_POND_IQE.copy()

    QUADRO_POND_REALIZ_IQE = QUADRO_POND_REALIZ_IQE.rename(columns={
      QUADRO_POND_REALIZ_IQE.columns[0]: "ID",
      QUADRO_POND_REALIZ_IQE.columns[1]: "PONDERAÇÕES (PRESTADORA)",
      QUADRO_POND_REALIZ_IQE.columns[2]: "PONDERAÇÕES (VI)",
      QUADRO_POND_REALIZ_IQE.columns[3]: "PONDERAÇÕES (ARSAL)",
    })

    contagem_IQE = TAB_EXPURGOS_IQE.groupby('ID_POND').size().reset_index(name='EXPURGOS')

    QUADRO_POND_REALIZ_IQE = QUADRO_POND_REALIZ_IQE.merge(
        contagem_IQE, left_on='ID', right_on='ID_POND', how='left'
    ).drop(columns=["ID_POND"], errors="ignore")

    QUADRO_POND_REALIZ_IQE["EXPURGOS"] = (
        QUADRO_POND_REALIZ_IQE["EXPURGOS"]
        .fillna(0)
        .astype("Int64")
        .astype("object")
        .replace(0, "")
    )

    QUADRO_POND_REALIZ_IQE = QUADRO_POND_REALIZ_IQE.merge(
        DESC_POR_ID_IQE, left_on='ID', right_on='ID_POND', how='left'
    ).drop(columns=["ID_POND"], errors="ignore")

    QUADRO_POND_REALIZ_IQE = QUADRO_POND_REALIZ_IQE.rename(
        columns={"DESCONSIDERAÇÕES - ARSAL": "DESCONSIDERAÇÕES (PLANO)"}
    )

    QUADRO_POND_REALIZ_IQE["DESCONSIDERAÇÕES (PLANO)"] = (
        QUADRO_POND_REALIZ_IQE["DESCONSIDERAÇÕES (PLANO)"]
        .fillna(0)
        .astype("Int64")
        .astype("object")
        .replace(0, "")
    )

    if QUADRO_POND_REALIZ_IQE[["EXPURGOS", "DESCONSIDERAÇÕES (PLANO)"]].replace("", pd.NA).isna().all().all():
      QUADRO_POND_REALIZ_IQE = "SEM IMPACTOS"

    return QUADRO_POND_NCONF, QUADRO_POND_REALIZ, QUADRO_POND_NCONF_IQE, QUADRO_POND_REALIZ_IQE



# CONSTRUÇÃO DOS GRÁFICOS

  def graf(TAB_MUN_FINAL, TAB_TIPO, TAB_MUN_FINAL_IQE, TAB_TIPO_IQE):

    # IQA

    TAB_MUN_FINAL = TAB_MUN_FINAL[:-1]
    
    def remove_acentos(txt):
      return ''.join(c for c in unicodedata.normalize('NFD', txt)
                    if unicodedata.category(c) != 'Mn')

    TAB_MUN_FINAL['MUNICÍPIO_SORT'] = TAB_MUN_FINAL['MUNICÍPIO'].apply(remove_acentos)

    TAB_MUN_FINAL = TAB_MUN_FINAL.sort_values(by='MUNICÍPIO_SORT', ascending=False)
    TAB_MUN_FINAL = TAB_MUN_FINAL.drop(columns=['MUNICÍPIO_SORT'])

    TAB_TIPO= TAB_TIPO.sort_values(by="ANÁLISE", ascending=False)


    
    # Plano de amostragem vs análises realizadas por município

    espacamento = 1.5
    y = np.arange(len(TAB_MUN_FINAL["MUNICÍPIO"])) * espacamento
    largura = 0.4

    fig, ax3 = plt.subplots(figsize=(12, len(TAB_MUN_FINAL["MUNICÍPIO"]) * 0.4 * espacamento))

    bars1 = ax3.barh(y + largura,   TAB_MUN_FINAL["ANALISES REALIZADAS"],                  height=largura, label="ANALISES REALIZADAS",                  color="green")
    bars2 = ax3.barh(y,             TAB_MUN_FINAL["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"], height=largura, label="PLANO DE AMOSTRAGEM AJUSTADO - ARSAL", color="skyblue")
    bars3 = ax3.barh(y - largura,   TAB_MUN_FINAL["EXPURGOS TOTAIS"],                      height=largura, label="EXPURGOS",                              color="orange")

    ax3.set_yticks(y)
    ax3.set_yticklabels(TAB_MUN_FINAL["MUNICÍPIO"])
    ax3.set_xlabel("Amostras")

    for bar in bars1:
      val = int(bar.get_width())
      if val > 0:
        ax3.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                f'{val}', va='center', fontsize=10)

    for bar in bars2:
      val = int(bar.get_width())
      if val > 0:
        ax3.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                f'{val}', va='center', fontsize=10)

    for bar in bars3:
      val = int(bar.get_width())
      if val > 0:
        ax3.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                f'{val}', va='center', fontsize=10)

    ax3.tick_params(axis='y', labelsize=11)
    ax3.grid(axis='x', linestyle='--', alpha=0.1)
    ax3.legend()

    plt.tight_layout()
    plt.savefig("grafico_plan_realiz_mun.png", dpi=300, bbox_inches="tight")
    plt.show()



    # Análises realizadas vs Análises conformes por município

    espacamento = 1.5
    y = np.arange(len(TAB_MUN_FINAL["MUNICÍPIO"])) * espacamento
    largura = 0.6

    fig, ax4 = plt.subplots(figsize=(12, len((TAB_MUN_FINAL["MUNICÍPIO"]))*0.4*espacamento))

    bars1 = ax4.barh(y + largura/2, TAB_MUN_FINAL["ANALISES CONFORMES"]-TAB_MUN_FINAL["EXPURGOS CONFORMES"], height=largura, label="ANALISES CONFORMES (SUBTRAINDO OS EXPURGOS CONFORMES)", color="green")
    bars2 = ax4.barh(y - largura/2, TAB_MUN_FINAL["ANALISES REALIZADAS"]-TAB_MUN_FINAL["EXPURGOS TOTAIS"], height=largura, label="ANALISES REALIZADAS (SUBTRAINDO OS EXPURGOS TOTAIS)", color="orange")

    ax4.set_yticks(y)
    ax4.set_yticklabels(TAB_MUN_FINAL["MUNICÍPIO"])
    ax4.set_xlabel("Amostras")

    for bar in bars1:
      ax4.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{int(bar.get_width())}', va='center', fontsize=10)

    for bar in bars2:
      ax4.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{int(bar.get_width())}', va='center', fontsize=10)

    ax4.tick_params(axis='y', labelsize=11)

    ax4.grid(axis='x', linestyle='--', alpha=0.1)
    ax4.legend()

    plt.savefig("grafico_conf_realiz_mun.png", dpi=300, bbox_inches="tight")

    plt.show()



    # IQA por município

    espacamento = 1.5
    y = np.arange(len(TAB_MUN_FINAL["MUNICÍPIO"])) * espacamento
    largura = 0.6

    fig, ax5 = plt.subplots(figsize=(12, len((TAB_MUN_FINAL["MUNICÍPIO"]))*0.4*espacamento))

    bars1 = ax5.barh(y + largura/2, TAB_MUN_FINAL["IQA (%)"], height=largura, label="IQA(%)", color="orange")

    ax5.set_yticks(y)
    ax5.set_yticklabels(TAB_MUN_FINAL["MUNICÍPIO"])
    ax5.set_xlabel("IQA (%)")

    for bar in bars1:
      ax5.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{bar.get_width():.1f}', va='center', fontsize=10)

    ax5.tick_params(axis='y', labelsize=11)

    ax5.grid(axis='x', linestyle='--', alpha=0.1)
    ax5.legend()

    plt.savefig("grafico_iqa_mun.png", dpi=300, bbox_inches="tight")

    plt.show()



    # Plano de amostragem vs análises realizadas por tipo de análise

    TAB_TIPO["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"] = TAB_TIPO["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"].astype(int)

    x = np.arange(len(TAB_TIPO["ANÁLISE"]))
    largura = 0.6

    fig, ax6 = plt.subplots(figsize=(10, 8))

    barras = ax6.bar(x, TAB_TIPO["ANALISES REALIZADAS"], width=largura, label="ANALISES REALIZADAS", color='skyblue')
    ax6.plot(x, TAB_TIPO["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"], label="PLANO DE AMOSTRAGEM AJUSTADO - ARSAL", color='green', marker='o', linewidth=2, linestyle='--')

    for barra in barras:
      altura = barra.get_height()
      ax6.text(barra.get_x() + barra.get_width()/2, altura/2, f'{altura}', ha='center', va='center', color='black', fontsize=10)

    for p, valor in enumerate(TAB_TIPO["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"]):
      ax6.text(x[p] + 0.1, valor + 1.5, f'{valor}', ha='left', va='bottom', color='green', fontsize=10)

    ax6.grid(True, linestyle='--', alpha=0.1)

    ax6.set_xticks(x)
    ax6.set_xticklabels(TAB_TIPO["ANÁLISE"], rotation=30)

    ax6.legend()

    ax6.set_ylabel('Amostras')

    plt.savefig("grafico_plan_realiz_tipo.png", dpi=300, bbox_inches="tight")

    plt.show()



    # Análises realizadas vs Análises conformes por tipo de análise

    espacamento = 1.5
    y = np.arange(len(TAB_TIPO["ANÁLISE"])) * espacamento
    largura = 0.6

    fig, ax7 = plt.subplots(figsize=(12, len((TAB_TIPO["ANÁLISE"]))*0.5*espacamento))

    bars1 = ax7.barh(y + largura/2, TAB_TIPO["ANALISES CONFORMES"] - TAB_TIPO["EXPURGOS CONFORMES"], height=largura, label="ANALISES CONFORMES (SUBTRAINDO OS EXPURGOS CONFORMES)", color="green")
    bars2 = ax7.barh(y - largura/2, TAB_TIPO["ANALISES REALIZADAS"] - TAB_TIPO["EXPURGOS TOTAIS"], height=largura, label="ANALISES REALIZADAS (SUBTRAINDO OS EXPURGOS TOTAIS)", color="orange")

    ax7.set_yticks(y)
    ax7.set_yticklabels(TAB_TIPO["ANÁLISE"])
    ax7.set_xlabel("Amostras")

    for bar in bars1:
      ax7.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{int(bar.get_width())}', va='center', fontsize=10)

    for bar in bars2:
      ax7.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{int(bar.get_width())}', va='center', fontsize=10)

    ax7.grid(axis='x', linestyle='--', alpha=0.1)
    ax7.legend()

    plt.savefig("grafico_conf_realiz_tipo.png", dpi=300, bbox_inches="tight")



    # IQE

    TAB_MUN_FINAL_IQE = TAB_MUN_FINAL_IQE[:-1]

    def remove_acentos(txt_IQE):
      return ''.join(c for c in unicodedata.normalize('NFD', txt_IQE)
                    if unicodedata.category(c) != 'Mn')

    TAB_MUN_FINAL_IQE['MUNICÍPIO_SORT'] = TAB_MUN_FINAL_IQE['MUNICÍPIO'].apply(remove_acentos)

    TAB_MUN_FINAL_IQE = TAB_MUN_FINAL_IQE.sort_values(by='MUNICÍPIO_SORT', ascending=False)
    TAB_MUN_FINAL_IQE = TAB_MUN_FINAL_IQE.drop(columns=['MUNICÍPIO_SORT'])

    TAB_TIPO_IQE = TAB_TIPO_IQE.sort_values(by="ANÁLISE", ascending=False)



    # Plano de amostragem vs análises realizadas por município

    espacamento = 1.5
    y = np.arange(len(TAB_MUN_FINAL_IQE["MUNICÍPIO"])) * espacamento
    largura = 0.4

    fig, ax3 = plt.subplots(figsize=(12, len(TAB_MUN_FINAL_IQE["MUNICÍPIO"]) * 0.4 * espacamento))

    bars1 = ax3.barh(y + largura,   TAB_MUN_FINAL_IQE["ANALISES REALIZADAS"],                  height=largura, label="ANALISES REALIZADAS",                  color="green")
    bars2 = ax3.barh(y,             TAB_MUN_FINAL_IQE["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"], height=largura, label="PLANO DE AMOSTRAGEM AJUSTADO - ARSAL", color="skyblue")
    bars3 = ax3.barh(y - largura,   TAB_MUN_FINAL_IQE["EXPURGOS TOTAIS"],                      height=largura, label="EXPURGOS",                              color="orange")

    ax3.set_yticks(y)
    ax3.set_yticklabels(TAB_MUN_FINAL_IQE["MUNICÍPIO"])
    ax3.set_xlabel("Amostras")

    for bar in bars1:
      val = int(bar.get_width())
      if val > 0:
        ax3.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                f'{val}', va='center', fontsize=10)

    for bar in bars2:
      val = int(bar.get_width())
      if val > 0:
        ax3.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                f'{val}', va='center', fontsize=10)

    for bar in bars3:
      val = int(bar.get_width())
      if val > 0:
        ax3.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                f'{val}', va='center', fontsize=10)

    ax3.tick_params(axis='y', labelsize=11)
    ax3.grid(axis='x', linestyle='--', alpha=0.1)
    ax3.legend()

    plt.tight_layout()
    plt.savefig("grafico_plan_realiz_mun_IQE.png", dpi=300, bbox_inches="tight")
    plt.show()



    # Análises realizadas vs Análises conformes por município

    espacamento = 1.5
    y = np.arange(len(TAB_MUN_FINAL_IQE["MUNICÍPIO"])) * espacamento
    largura = 0.6

    fig, ax4 = plt.subplots(figsize=(12, len((TAB_MUN_FINAL_IQE["MUNICÍPIO"]))*0.4*espacamento))

    bars1 = ax4.barh(y + largura/2, TAB_MUN_FINAL_IQE["ANALISES CONFORMES"]-TAB_MUN_FINAL_IQE["EXPURGOS CONFORMES"], height=largura, label="ANALISES CONFORMES (SUBTRAINDO OS EXPURGOS CONFORMES)", color="green")
    bars2 = ax4.barh(y - largura/2, TAB_MUN_FINAL_IQE["ANALISES REALIZADAS"]-TAB_MUN_FINAL_IQE["EXPURGOS TOTAIS"], height=largura, label="ANALISES REALIZADAS (SUBTRAINDO OS EXPURGOS TOTAIS)", color="orange")

    ax4.set_yticks(y)
    ax4.set_yticklabels(TAB_MUN_FINAL_IQE["MUNICÍPIO"])
    ax4.set_xlabel("Amostras")

    for bar in bars1:
      ax4.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{int(bar.get_width())}', va='center', fontsize=10)

    for bar in bars2:
      ax4.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{int(bar.get_width())}', va='center', fontsize=10)

    ax4.tick_params(axis='y', labelsize=11)

    ax4.grid(axis='x', linestyle='--', alpha=0.1)
    ax4.legend()

    plt.savefig("grafico_conf_realiz_mun_IQE.png", dpi=300, bbox_inches="tight")

    plt.show()



    # IQE por município

    espacamento = 1.5
    y = np.arange(len(TAB_MUN_FINAL_IQE["MUNICÍPIO"])) * espacamento
    largura = 0.6

    fig, ax5 = plt.subplots(figsize=(12, len((TAB_MUN_FINAL_IQE["MUNICÍPIO"]))*0.4*espacamento))

    bars1 = ax5.barh(y + largura/2, TAB_MUN_FINAL_IQE["IQE (%)"], height=largura, label="IQE(%)", color="orange")

    ax5.set_yticks(y)
    ax5.set_yticklabels(TAB_MUN_FINAL_IQE["MUNICÍPIO"])
    ax5.set_xlabel("IQE (%)")

    for bar in bars1:
      ax5.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{bar.get_width():.1f}', va='center', fontsize=10)

    ax5.tick_params(axis='y', labelsize=11)

    ax5.grid(axis='x', linestyle='--', alpha=0.1)
    ax5.legend()

    plt.savefig("grafico_iqa_mun_IQE.png", dpi=300, bbox_inches="tight")

    plt.show()



    # Plano de amostragem vs análises realizadas por tipo de análise

    TAB_TIPO_IQE["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"] = TAB_TIPO_IQE["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"].astype(int)

    x = np.arange(len(TAB_TIPO_IQE["ANÁLISE"]))
    largura = 0.6

    fig, ax6 = plt.subplots(figsize=(10, 8))

    barras = ax6.bar(x, TAB_TIPO_IQE["ANALISES REALIZADAS"], width=largura, label="ANALISES REALIZADAS", color='skyblue')
    ax6.plot(x, TAB_TIPO_IQE["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"], label="PLANO DE AMOSTRAGEM AJUSTADO - ARSAL", color='green', marker='o', linewidth=2, linestyle='--')

    for barra in barras:
      altura = barra.get_height()
      ax6.text(barra.get_x() + barra.get_width()/2, altura/2, f'{altura}', ha='center', va='center', color='black', fontsize=10)

    for p, valor in enumerate(TAB_TIPO_IQE["PLANO DE AMOSTRAGEM AJUSTADO - ARSAL"]):
      ax6.text(x[p] + 0.1, valor + 1.5, f'{valor}', ha='left', va='bottom', color='green', fontsize=10)

    ax6.grid(True, linestyle='--', alpha=0.1)

    ax6.set_xticks(x)
    ax6.set_xticklabels(TAB_TIPO_IQE["ANÁLISE"], rotation=30)

    ax6.legend()

    ax6.set_ylabel('Amostras')

    plt.savefig("grafico_plan_realiz_tipo_IQE.png", dpi=300, bbox_inches="tight")

    plt.show()



    # Análises realizadas vs Análises conformes por tipo de análise

    espacamento = 1.5
    y = np.arange(len(TAB_TIPO_IQE["ANÁLISE"])) * espacamento
    largura = 0.6

    fig, ax7 = plt.subplots(figsize=(12, len((TAB_TIPO_IQE["ANÁLISE"]))*0.5*espacamento))

    bars1 = ax7.barh(y + largura/2, TAB_TIPO_IQE["ANALISES CONFORMES"] - TAB_TIPO_IQE["EXPURGOS CONFORMES"], height=largura, label="ANALISES CONFORMES (SUBTRAINDO OS EXPURGOS CONFORMES)", color="green")
    bars2 = ax7.barh(y - largura/2, TAB_TIPO_IQE["ANALISES REALIZADAS"] - TAB_TIPO_IQE["EXPURGOS TOTAIS"], height=largura, label="ANALISES REALIZADAS (SUBTRAINDO OS EXPURGOS TOTAIS)", color="orange")

    ax7.set_yticks(y)
    ax7.set_yticklabels(TAB_TIPO_IQE["ANÁLISE"])
    ax7.set_xlabel("Amostras")

    for bar in bars1:
      ax7.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{int(bar.get_width())}', va='center', fontsize=10)

    for bar in bars2:
      ax7.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
              f'{int(bar.get_width())}', va='center', fontsize=10)

    ax7.grid(axis='x', linestyle='--', alpha=0.1)
    ax7.legend()

    plt.savefig("grafico_conf_realiz_tipo_IQE.png", dpi=300, bbox_inches="tight")

    plt.show()




# FUNÇÃO PRINCIPAL E EXECUÇÃO - IQA

  def MAIN_IQA_IQE():

    arquivo_gerais = pd.read_excel('ATT_SQL.xlsx', sheet_name='GERAIS')

    MES_ANO, TRIM_ANO, TRIMMIN_ANO, ANO_CONTRAT, PRESTADORA, META, NCOF_P, NREA_P, IQA_P, NCOF_VI, NREA_VI, IQA_VI, FONTE_NREALIZ, FONTE_PLANO, FONT_NCONF, ACREDITAÇÃO, META_IQE, NCOF_P_IQE, NREA_P_IQE, IQE_P, NCOF_VI_IQE, NREA_VI_IQE, IQE_VI, FONTE_NREALIZ_IQE, FONTE_PLANO_IQE, FONT_NCONF_IQE, ACREDITA_IQE, MESMIN_ANO = infos_gerais (arquivo_gerais)

    PLAN_MUN_TIPO = pd.read_excel ('ATT_SQL.xlsx', sheet_name='IQA_PLANO')

    PLAN_MUN_TIPO = PLAN_MUN_TIPO.rename(columns={'DESCONSIDERACOES - ARSAL' : 'DESCONSIDERAÇÕES - ARSAL'})

    AMOSTRAS_REALIZADAS = pd.read_excel ('ATT_SQL.xlsx', sheet_name='IQA_DETALHADO')

    AMOSTRAS_REALIZADAS = AMOSTRAS_REALIZADAS.rename(columns={
    'id_amostra': 'ID_AMOSTRA',
    'saa': 'CONTA',
    'analise': 'ANÁLISE',
    'municipio': 'CIDADE',
    'parecer': 'ACEITO?',
    'conformidade_vi': 'CONFORMIDADE - VI',
    'resultado': 'CONFORMIDADE - ARSAL',
    'expurgos': 'EXPURGOS - ARSAL',
    'id_pond': 'ID_POND'
    })

    PONDERACOES = pd.read_excel ('ATT_SQL.xlsx', sheet_name='IQA_PONDERACOES')

    PLAN_MUN_TIPO_IQE = pd.read_excel ('ATT_SQL.xlsx', sheet_name='IQE_PLANO')

    PLAN_MUN_TIPO_IQE = PLAN_MUN_TIPO_IQE.rename(columns={
    'desconsideracoes_arsal': 'DESCONSIDERAÇÕES - ARSAL',
    'id_pond': 'ID_POND',
    'plano': 'PLANO',
    'ETE': 'SES',
    'parametros': 'Parâmetros',
    'cidade': 'Cidade'
    })

    AMOSTRAS_REALIZADAS_IQE = pd.read_excel ('ATT_SQL.xlsx', sheet_name='IQE_DETALHADO')

    AMOSTRAS_REALIZADAS_IQE = AMOSTRAS_REALIZADAS_IQE.rename(columns={
    'id_amostra': 'ID_AMOSTRA',
    'ETE': 'CONTA',
    'analise': 'ANÁLISE',
    'cidade': 'CIDADE',
    'aceito?': 'ACEITO?',
    'conformidade_vi': 'CONFORMIDADE - VI',
    'resultado': 'CONFORMIDADE - ARSAL',
    'expurgos': 'EXPURGOS - ARSAL',
    'id_pond': 'ID_POND'
    })

    PONDERACOES_IQE = pd.read_excel('ATT_SQL.xlsx', sheet_name='IQE_PONDERACOES')

    PLAN_MUN_TIPO = PLAN_MUN_TIPO.fillna(0)
    AMOSTRAS_REALIZADAS = AMOSTRAS_REALIZADAS.fillna(0)
    PONDERACOES = PONDERACOES.fillna(0)

    PLAN_MUN_TIPO_IQE = PLAN_MUN_TIPO_IQE.fillna(0)
    AMOSTRAS_REALIZADAS_IQE = AMOSTRAS_REALIZADAS_IQE.fillna(0)
    PONDERACOES_IQE = PONDERACOES_IQE.fillna(0)

    TAB_MUN, TAB_TIPO, TAB_REALIZADAS, TAB_POND, BASE_MUN_TIPO_completo = base(PLAN_MUN_TIPO, AMOSTRAS_REALIZADAS, PONDERACOES, MES_ANO)

    TAB_MUN_IQE, TAB_TIPO_IQE, TAB_REALIZADAS_IQE, TAB_POND_IQE, BASE_MUN_TIPO_IQE_completo = base_IQE(PLAN_MUN_TIPO_IQE, AMOSTRAS_REALIZADAS_IQE, PONDERACOES_IQE, TRIM_ANO)

    TAB_MUN_FINAL, TAB_MUN_ARSAL, TAB_MUN_VI, desc_plano, TAB_MUN_NCONF, TAB_MUN_NCONF_VI, TAB_MUN_FINAL_IQE, TAB_MUN_ARSAL_IQE, TAB_MUN_VI_IQE, desc_plano_IQE, TAB_MUN_NCONF_IQE, TAB_MUN_NCONF_VI_IQE = analise_mun(TAB_MUN, TAB_REALIZADAS, BASE_MUN_TIPO_completo, TAB_MUN_IQE, TAB_REALIZADAS_IQE, BASE_MUN_TIPO_IQE_completo)

    AN_TIPO, TAB_TIPO_FINAL, TAB_TIPO_RESUMIDA_VI, total_conf_vi_naoconf_arsal, total_naoconf_vi_conf_arsal, AN_TIPO_IQE, TAB_TIPO_FINAL_IQE, TAB_TIPO_RESUMIDA_VI_IQE, total_conf_vi_naoconf_arsal_IQE, total_naoconf_vi_conf_arsal_IQE  = analise_tipo(TAB_TIPO, TAB_REALIZADAS, BASE_MUN_TIPO_completo, TAB_TIPO_IQE, TAB_REALIZADAS_IQE, BASE_MUN_TIPO_IQE_completo)

    QUADRO_POND_NCONF, QUADRO_POND_REALIZ, QUADRO_POND_NCONF_IQE, QUADRO_POND_REALIZ_IQE = pond(TAB_POND, TAB_REALIZADAS, BASE_MUN_TIPO_completo, TAB_POND_IQE, TAB_REALIZADAS_IQE, BASE_MUN_TIPO_IQE_completo)

    graf(TAB_MUN_FINAL, AN_TIPO, TAB_MUN_FINAL_IQE, AN_TIPO_IQE)

    dic_info = var_doc (MES_ANO, TRIM_ANO, TRIMMIN_ANO, ANO_CONTRAT, PRESTADORA, META, NCOF_P, NREA_P, IQA_P, NCOF_VI, NREA_VI, IQA_VI, FONTE_NREALIZ, FONTE_PLANO, FONT_NCONF, ACREDITAÇÃO, TAB_MUN_FINAL, TAB_MUN_FINAL_IQE, QUADRO_POND_NCONF, QUADRO_POND_REALIZ, TAB_MUN_ARSAL, TAB_MUN_VI, TAB_TIPO_FINAL, TAB_TIPO_RESUMIDA_VI, MESMIN_ANO, total_conf_vi_naoconf_arsal, total_naoconf_vi_conf_arsal, desc_plano, TAB_MUN_NCONF, TAB_MUN_NCONF_VI, META_IQE, NCOF_P_IQE, NREA_P_IQE, IQE_P, NCOF_VI_IQE, NREA_VI_IQE, IQE_VI, FONTE_NREALIZ_IQE, FONTE_PLANO_IQE, FONT_NCONF_IQE, ACREDITA_IQE, QUADRO_POND_REALIZ_IQE, QUADRO_POND_NCONF_IQE, TAB_MUN_VI_IQE, TAB_MUN_ARSAL_IQE, TAB_TIPO_RESUMIDA_VI_IQE, TAB_TIPO_FINAL_IQE, total_conf_vi_naoconf_arsal_IQE, total_naoconf_vi_conf_arsal_IQE, desc_plano_IQE, TAB_MUN_NCONF_IQE, TAB_MUN_NCONF_VI_IQE)

    arquivo_entrada = "modelo_iqa_iqe_automatico.docx"

    doc_modificado = substituir_var(arquivo_entrada, dic_info)

    arquivo_saida = f'INDICADORES- ' + str(PRESTADORA) + ' - ' + str(MES_ANO) + '.docx'

    doc_modificado.save(arquivo_saida)

    return arquivo_saida

  
  url = "https://github.com/carlosfarsal/DASH_ATT/raw/main/modelo_iqa_iqe_automatico.docx"

  output = "modelo_iqa_iqe_automatico.docx"

  gdown.download(url, output, quiet=False)

  nome_arq = MAIN_IQA_IQE()

  return nome_arq


#############################################################################################################################################################################################################################

# FUNÇÃO PRINCIPAL DO APP

def MAIN ():

    
  prestadora_sigla, ano, ano_contratual, mes, iqa_prest, iqa_vi, iqa_meta, trimestre, iqa_detalhado, plano_dados_tratados, plano_totais_municipio, plano_totais_parametros, iqe_prest, iqe_vi, iqe_meta, iqe_detalhado, plano_iqe_completo, plano_iqe_totais_municipio, plano_iqe_totais_parametros, iqe_detalhado2 = ler_ajustar_arquivo ()
    
  
  print (plano_iqe_totais_parametros)

  if mes == "November" or mes == "February" or mes == "May" or mes == "August":

    realizados, parametros, mun, plano_final, iqa_final, realizados_iqe, parametros_iqe, mun_iqe, plano_final_iqe, iqe_final = criar_tabelas (prestadora_sigla, ano, ano_contratual, mes, iqa_prest, iqa_vi, iqa_meta, trimestre, iqa_detalhado, plano_dados_tratados, plano_totais_municipio, plano_totais_parametros, iqe_prest, iqe_vi, iqe_meta, iqe_detalhado, plano_iqe_completo, plano_iqe_totais_municipio, plano_iqe_totais_parametros, iqe_detalhado2)

    print (realizados)
    print (parametros)
    print (mun)
    print (plano_final)
    print (iqa_final)

    alimentar_bd (prestadora_sigla, ano, mes, iqa_detalhado, realizados, parametros, mun, plano_final, iqa_final, iqe_detalhado, realizados_iqe, parametros_iqe, mun_iqe, plano_final_iqe, iqe_final, trimestre, ano_contratual)
    backup_bd ()
    graf_iqa_iqe (prestadora_sigla, mes)
    nc_nr (prestadora_sigla, mes)
    nome_arq = rel_iqa_iqe()

    return nome_arq


  else:

    iqe_prest = 0
    iqe_vi = 0
    iqe_meta = 0
    iqe_detalhado = 0
    plano_iqe_completo = 0
    plano_iqe_totais_municipio = 0
    plano_iqe_totais_parametros = 0
    realizados_iqe = 0
    parametros_iqe = 0
    mun_iqe = 0
    plano_final_iqe = 0
    iqe_final = 0

    realizados, parametros, mun, plano_final, iqa_final = criar_tabelas (prestadora_sigla, ano, ano_contratual, mes, iqa_prest, iqa_vi, iqa_meta, trimestre, iqa_detalhado, plano_dados_tratados, plano_totais_municipio, plano_totais_parametros, iqe_prest, iqe_vi, iqe_meta, iqe_detalhado, plano_iqe_completo, plano_iqe_totais_municipio, plano_iqe_totais_parametros)
    alimentar_bd (prestadora_sigla, ano, mes, iqa_detalhado, realizados, parametros, mun, plano_final, iqa_final, iqe_detalhado, realizados_iqe, parametros_iqe, mun_iqe, plano_final_iqe, iqe_final, trimestre, ano_contratual)
    backup_bd ()
    graf_iqa_iqe (prestadora_sigla)
    nc_nr (prestadora_sigla)
    nome_arq = rel_iqa()

    return nome_arq


#############################################################################################################################################################################################################################

# INTERFACE

st.set_page_config(
    page_title="ELABORAÇÃO DE RELATÓRIOS + ATUALIZAÇÃO DO BANCO DE DADOS - IQA/IQE",
    page_icon="💧",
    layout="wide"
)

st.markdown("""
<style>

/* Remove elementos padrão do Streamlit */

#MainMenu {
    visibility: hidden;
}

footer {
    visibility: hidden;
}

header {
    visibility: hidden;
}

/* Fundo geral */

.stApp {
    background-color: #E8F4F8;
}

/* CABEÇALHO - Faixa branca com logo e título */

.header-container {
    background-color: white;
    padding: 20px 40px;
    display: flex;
    align-items: center;
    justify-content: flex-start;
    gap: 30px;
    box-shadow: 0px 2px 8px rgba(0,0,0,0.1);
    margin: -16px -16px 0 -16px;
}

/* Cor preta do texto do spinner */
.stSpinner > div {
    color: #000000 !important;
}

/* Cor preta do ícone girando */
.stSpinner svg {
    stroke: #28A745 !important;
}

.header-logo {
    flex-shrink: 0;
}

.header-content {
    flex-grow: 1;
}

.header-content h1 {
    color: #17365D;
    margin: 0;
    font-size: 32px;
    font-weight: 700;
}

.header-content p {
    color: #555;
    margin: 8px 0 0 0;
    font-size: 14px;
}

/* Container principal com fundo azul claro */

.main-content {
    background-color: #E8F4F8;
    padding: 30px;
    margin-top: 0;
}

#/* Título */

# h1 {
#     color: #17365D;
# }

# h3 {
#     color: #17365D;
# }

/* Título do cabeçalho */

.header-title {
    color: #000000 !important;
    margin: 0;
    font-size: 42px;
    font-weight: 700;
}

/* Texto abaixo do título */

.header-subtitle {
    color: #000000 !important;
    margin-top: 12px;
    font-size: 20px;
}

/* Botões */

.stButton button {
    background-color: #17365D;
    color: white;
    border-radius: 10px;
    height: 45px;
    padding: 0 20px;
    font-size: 16px;
    font-weight: bold;
    border: none;
    cursor: pointer;
    transition: background-color 0.3s;
}

.stButton button:hover {
    background-color: #0F243E;
}

/* Caixa de conteúdo */

.card {
    background-color: white;
    padding: 25px;
    border-radius: 15px;
    box-shadow: 0px 3px 12px rgba(0,0,0,0.12);
    margin-bottom: 20px;
}

/* Adicione isto: */
.card h3 {
    color: #17365D;
    font-size: 20px;
    margin: 0 0 12px 0;
}

.card p {
    color: #333;
    font-size: 16px;
    margin: 0;
}

/* Texto */

.texto {
    text-align: center;
    color: #555;
    font-size: 17px;
}

/* RODAPÉ - Faixa azul escuro */

.footer-container {
    background-color: #0F243E;
    color: white;
    padding: 25px 40px;
    text-align: center;
    margin-top: 40px;
    margin-left: -16px;
    margin-right: -16px;
    margin-bottom: -16px;
}

.footer-container p {
    margin: 8px 0;
    font-size: 14px;
    line-height: 1.6;
}

.footer-title {
    font-weight: 700;
    font-size: 15px;
    color: #E8F4F8;
}

.footer-address {
    font-size: 13px;
    color: #BDD9E8;
}

</style>

""", unsafe_allow_html=True)

CAMINHO_IMAGEM = os.path.join(
    os.path.dirname(__file__),
    "arsal_cover.png"
)

header_style = """
<div style="
    background-color: white;
    padding: 20px 40px;
    display: flex;
    align-items: center;
    gap: 30px;
    box-shadow: 0px 2px 8px rgba(0,0,0,0.1);
    margin: -16px -16px 20px -16px;
    border-radius: 0;
">
    <div style="flex-shrink: 0;">
"""

if os.path.exists(CAMINHO_IMAGEM):
    header_style += f'<img src="file://{CAMINHO_IMAGEM}" width="140" style="object-fit: contain; height: auto;">'
else:
    header_style += '<div style="width: 140px; height: 120px; background-color: #ccc; display: flex; align-items: center; justify-content: center; border-radius: 8px; color: #666; font-size: 12px;">Logo</div>'

st.markdown("""
<style>
.header-box{
    background:white;
    padding:20px 30px;
    border-radius:12px;
    box-shadow:0 2px 8px rgba(0,0,0,.12);
    margin-bottom:25px;
}
</style>
""", unsafe_allow_html=True)

col_logo, col_texto = st.columns([0.8, 2.2], gap="small", vertical_alignment="center")

with col_logo:
    st.image(CAMINHO_IMAGEM, width=320)

with col_texto:
    st.markdown("""
    <h1 class="header-title">
    ELABORAÇÃO DE RELATÓRIOS + ATUALIZAÇÃO DO BANCO DE DADOS - IQA/IQE
    </h1>

    <p class="header-subtitle">
    Sistema para elaboração de relatórios de acompanhamento dos indicadores de qualidade da água e esgoto (IQA/IQE) e atualização do banco de dados.
    </p>
    """, unsafe_allow_html=True)

st.markdown('<div class="main-content">', unsafe_allow_html=True)

st.markdown(
"""
<div class="card">

<h3>
📂 Upload da Planilha - IQA/IQE
<a href="https://raw.githubusercontent.com/carlosfarsal/DASH_ATT/main/ATT_SQL.xlsx" target="_blank">(TEMPLATE)</a>
</h3>

<p>
Selecione a planilha mensal para atualização do banco de dados e elaboração de relatório.
</p>

</div>

""",
unsafe_allow_html=True
)

arquivo = st.file_uploader(
    "Selecione a planilha IQA/IQE",
    type=["xlsx", "xls"],
    label_visibility="collapsed"
)

if arquivo is not None:

    st.success(
        f"Arquivo selecionado: {arquivo.name}"
    )

    st.divider()

    if "confirmar_envio" not in st.session_state:
        st.session_state.confirmar_envio = False

    if st.button("📤 Enviar arquivo"):

        st.session_state.confirmar_envio = True

    if st.session_state.confirmar_envio:

        st.markdown("""
        <div style="
            background-color: #E8F4F8;
            color: #000000;
            padding: 15px;
            border-radius: 10px;
            border-left: 6px solid #17365D;
            font-size: 16px;
            font-weight: 600;
        ">
        Confirma o envio e atualização do banco de dados?
        </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns(2)

        with col1:

            if st.button("✅ Confirmar"):

                novo_nome = "ATT_SQL.xlsx"

                caminho = os.path.join(
                    os.path.dirname(__file__),
                    novo_nome
                )

                with open(caminho, "wb") as f:

                    f.write(
                        arquivo.getbuffer()
                    )

                if os.path.exists(caminho):

                    try:

                        with st.spinner("Processando dados e atualizando banco..."):

                            nome_arq = MAIN()

                        st.success("Arquivo processado com sucesso!")
                        st.divider()

                        if os.path.exists(nome_arq):
                            with open(nome_arq, "rb") as file:
                                st.download_button(
                                    label="📥 Baixar Arquivo",
                                    data=file.read(),
                                    file_name=nome_arq,
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    use_container_width=True
                                )

                        st.session_state.confirmar_envio = False

                    except Exception as e:

                        st.error(
                            "Erro durante atualização:"
                        )

                        st.exception(e)

                else:

                    st.error(
                        "Arquivo ATT_SQL.xlsx não foi encontrado."
                    )

        with col2:

            if st.button("❌ Cancelar"):

                st.session_state.confirmar_envio = False

                st.info(
                    "Envio cancelado."
                )

st.markdown('</div>', unsafe_allow_html=True)

footer_html = """
<div class="footer-container">
    <p class="footer-title">ARSAL - Agência Reguladora de Serviços Públicos do Estado de Alagoas</p>
    <p class="footer-address">Rua Engenheiro Roberto Gonçalves Menezes, 149, 1º andar, Centro, Maceió - AL, CEP 57020-680</p>
</div>
"""

st.markdown(footer_html, unsafe_allow_html=True)
